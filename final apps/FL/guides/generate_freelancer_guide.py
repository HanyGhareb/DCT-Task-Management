# =============================================================================
# Freelancers Module (App 203) — Freelancer End-User Guide Generator
# Produces:  FL_USER_GUIDE_FREELANCER_V1.docx   (this folder)
#            assets/*.png                        (flowcharts + cropped screens)
# Written for freelancers in plain language — no role codes, no system jargon.
# Business rules sourced from final apps/FL/docs/BRD.md and DCT_FL_PKG.
# Regenerate after any rule change:  python generate_freelancer_guide.py
# =============================================================================
import os
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
EVIDENCE = os.path.join(HERE, "..", "UAT", "evidence_12-Jun-2026-04")
os.makedirs(ASSETS, exist_ok=True)

TODAY = "13 June 2026"
VERSION = "1.0"
OUT_NAME = "FL_USER_GUIDE_FREELANCER_V1.docx"

# ── palette ──────────────────────────────────────────────────────────────────
PURPLE = "#7C4DBE"; PURPLE_DARK = "#5B378E"; PURPLE_SOFT = "#EFE9F8"
GREEN = "#2E7D32"; RED = "#C62828"; AMBER = "#C77700"; GREY = "#58606B"
W_PURPLE = RGBColor(0x7C, 0x4D, 0xBE)
W_PURPLE_DARK = RGBColor(0x5B, 0x37, 0x8E)
W_GREY = RGBColor(0x58, 0x60, 0x6B)
W_FIELD_BLUE = RGBColor(0x0B, 0x53, 0x94)        # bold-blue field names
W_GREEN = RGBColor(0x2E, 0x7D, 0x32)
W_RED = RGBColor(0xC6, 0x28, 0x28)
W_AMBER = RGBColor(0xC7, 0x77, 0x00)

STATUS_COLORS = {
    "approved": W_GREEN, "active": W_GREEN, "paid": W_GREEN, "accepted": W_GREEN,
    "valid": W_GREEN,
    "rejected": W_RED, "cancelled": W_RED, "expired": W_RED,
    "returned": W_AMBER, "expiring soon": W_AMBER, "pending": W_AMBER,
    "under review": W_AMBER, "submitted": W_AMBER,
    "draft": W_GREY, "completed": W_GREY, "skipped": W_GREY,
    "voucher raised": W_PURPLE_DARK,
}


# =============================================================================
# Diagrams (matplotlib)
# =============================================================================
def _fig(w, h):
    fig, ax = plt.subplots(figsize=(w, h), dpi=210)
    ax.set_xlim(0, 10); ax.set_ylim(0, 10)
    ax.axis("off")
    return fig, ax


def _box(ax, cx, cy, w, h, title, sub=None, fc=PURPLE, tc="white",
         fs=10.5, sub_fs=8.2):
    ax.add_patch(FancyBboxPatch((cx - w / 2, cy - h / 2), w, h,
                 boxstyle="round,pad=0.02,rounding_size=0.14",
                 fc=fc, ec=fc, lw=1.0))
    if sub:
        ax.text(cx, cy + h * 0.24, title, ha="center", va="center",
                color=tc, fontsize=fs, fontweight="bold")
        ax.text(cx, cy - h * 0.26, sub, ha="center", va="center",
                color=tc, fontsize=sub_fs, linespacing=1.25)
    else:
        ax.text(cx, cy, title, ha="center", va="center", color=tc,
                fontsize=fs, fontweight="bold", linespacing=1.3)


def _arrow(ax, p1, p2, color="#6B7280", style="solid", label=None,
           lab_dx=0, lab_dy=0.28, lab_color=None, rad=0.0):
    ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="-|>", mutation_scale=15,
                 lw=1.7, color=color, linestyle=style,
                 connectionstyle=f"arc3,rad={rad}", shrinkA=2, shrinkB=2))
    if label:
        mx, my = (p1[0] + p2[0]) / 2 + lab_dx, (p1[1] + p2[1]) / 2 + lab_dy
        ax.text(mx, my, label, ha="center", va="center", fontsize=8.2,
                color=lab_color or color, fontweight="bold")


def _save(fig, name):
    path = os.path.join(ASSETS, name)
    fig.savefig(path, bbox_inches="tight", facecolor="white", pad_inches=0.12)
    plt.close(fig)
    print("  diagram:", name)
    return path


def diagram_journey():
    fig, ax = _fig(10.5, 3.6)
    bw, bh = 2.15, 1.85
    top = [("1 · Register", "Fill and send the\nregistration form"),
           ("2 · Approval", "The Freelancers team\nreviews your request"),
           ("3 · Your account", "Login and vendor\nnumber are issued"),
           ("4 · Contract", "Contract activated,\npayment plan created")]
    bot = [("7 · Renew", "Continue with a\nnew contract"),
           ("6 · Get paid", "One payment per\nperiod, to your bank"),
           ("5 · Deliver work", "Submit your work\nfor acceptance")]
    xs = [1.35, 3.85, 6.35, 8.85]
    for i, (t, s) in enumerate(top):
        _box(ax, xs[i], 7.6, bw, bh, t, s)
        if i:
            _arrow(ax, (xs[i - 1] + bw / 2, 7.6), (xs[i] - bw / 2, 7.6))
    for i, (t, s) in enumerate(bot):
        _box(ax, xs[3 - i], 2.6, bw, bh, t, s,
             fc=GREEN if t.startswith("6") else PURPLE)
        if i:
            _arrow(ax, (xs[4 - i] - bw / 2, 2.6), (xs[3 - i] + bw / 2, 2.6))
    _arrow(ax, (xs[3], 7.6 - bh / 2), (xs[3], 2.6 + bh / 2))
    return _save(fig, "journey.png")


def diagram_registration():
    fig, ax = _fig(8.6, 6.0)
    cx = 3.1; bw, bh = 4.3, 1.15
    _box(ax, cx, 9.0, bw, bh, "You complete the registration form",
         "Saved as Draft — you can still edit it", fc=GREY)
    _box(ax, cx, 6.6, bw, bh, "Review — step 1",
         "A Freelancers manager checks your request")
    _box(ax, cx, 4.2, bw, bh, "Review — step 2",
         "The Freelancers team gives the final approval")
    _box(ax, cx, 1.6, bw + 0.6, bh, "Approved",
         "Your profile, vendor number and login are created", fc=GREEN)
    _arrow(ax, (cx, 9.0 - bh / 2), (cx, 6.6 + bh / 2),
           label="you submit", lab_dx=1.15, lab_dy=0)
    _arrow(ax, (cx, 6.6 - bh / 2), (cx, 4.2 + bh / 2),
           label="approved", lab_dx=1.1, lab_dy=0, lab_color=GREEN)
    _arrow(ax, (cx, 4.2 - bh / 2), (cx, 1.6 + bh / 2),
           label="approved", lab_dx=1.1, lab_dy=0, lab_color=GREEN)
    # side outcomes
    _box(ax, 7.9, 6.6, 3.4, 1.15, "Returned to you",
         "Something needs fixing —\ncorrect it and submit again", fc=AMBER)
    _box(ax, 7.9, 3.4, 3.4, 1.0, "Rejected",
         "Closed, with the reason recorded", fc=RED)
    _arrow(ax, (cx + bw / 2, 6.6), (7.9 - 1.7, 6.6), color=AMBER)
    _arrow(ax, (cx + bw / 2, 4.2), (7.9 - 1.7, 6.2), color=AMBER, rad=-0.15)
    _arrow(ax, (cx + bw / 2, 6.35), (7.9 - 1.7, 3.6), color=RED, rad=0.18)
    _arrow(ax, (cx + bw / 2, 3.95), (7.9 - 1.7, 3.35), color=RED, rad=0.1)
    _arrow(ax, (7.9, 6.6 + 0.58), (cx + 0.6, 9.0 - bh / 2), color=AMBER,
           style="dashed", rad=0.25, label="fix and resubmit", lab_dy=0.45)
    return _save(fig, "registration_flow.png")


def diagram_payment():
    fig, ax = _fig(11.0, 2.5)
    bw, bh = 1.78, 2.6
    steps = [("Payment\nperiod due", "From your contract’s\npayment plan", GREY),
             ("Voucher\nprepared", "Your invoice number\nand date are added", PURPLE),
             ("Approval", "Your manager contact —\nplus a second approval\nfor 50,000 AED or more", PURPLE),
             ("Payment\nprocessed", "Sent to your primary\nbank account", PURPLE),
             ("Paid", "Payment date and bank\nreference recorded", GREEN)]
    xs = [1.05, 3.05, 5.05, 7.05, 9.05]
    for (t, s, c), x in zip(steps, xs):
        _box(ax, x, 5.0, bw, bh + 2.2, t, s, fc=c, fs=10, sub_fs=7.8)
    for a, b in zip(xs, xs[1:]):
        _arrow(ax, (a + bw / 2, 5.0), (b - bw / 2, 5.0))
    return _save(fig, "payment_flow.png")


def diagram_deliverable():
    fig, ax = _fig(9.6, 3.8)
    bw, bh = 2.6, 2.6
    _box(ax, 1.5, 5.0, bw, bh, "You deliver\nthe work", fc=GREY)
    _box(ax, 4.7, 5.0, bw, bh + 0.8, "Recorded as\na deliverable",
         "title, date, quantity", fc=PURPLE)
    _box(ax, 7.9, 7.4, bw + 0.5, bh, "Accepted",
         "Counts towards\nyour payment", fc=GREEN)
    _box(ax, 7.9, 2.6, bw + 0.5, bh, "Rejected",
         "Reason recorded — submit\ncorrected work as a new item", fc=RED)
    _arrow(ax, (1.5 + bw / 2, 5.0), (4.7 - bw / 2, 5.0))
    _arrow(ax, (4.7 + bw / 2, 5.7), (7.9 - (bw + 0.5) / 2, 7.2), color=GREEN)
    _arrow(ax, (4.7 + bw / 2, 4.3), (7.9 - (bw + 0.5) / 2, 2.8), color=RED)
    return _save(fig, "deliverable_flow.png")


def diagram_profile_change():
    fig, ax = _fig(9.6, 3.8)
    bw, bh = 3.0, 3.2
    _box(ax, 1.75, 5.0, bw, bh, "You request\na change",
         "bank account · e-mail · phone", fc=GREY)
    _box(ax, 5.05, 5.0, bw, bh - 0.6, "Freelancers team\nreview", fc=PURPLE)
    _box(ax, 8.4, 7.4, bw, bh - 0.6, "Approved",
         "Applied automatically", fc=GREEN)
    _box(ax, 8.4, 2.6, bw, bh - 0.4, "Rejected /\nReturned",
         "Reason or correction\nrequest recorded", fc=RED, fs=10)
    _arrow(ax, (1.75 + bw / 2, 5.0), (5.05 - bw / 2, 5.0))
    _arrow(ax, (5.05 + bw / 2, 5.6), (8.4 - bw / 2, 7.2), color=GREEN)
    _arrow(ax, (5.05 + bw / 2, 4.4), (8.4 - bw / 2, 2.8), color=RED)
    return _save(fig, "profile_change_flow.png")


def diagram_documents():
    fig, ax = _fig(10.2, 3.0)
    bh = 4.2
    _box(ax, 1.95, 5.4, 3.3, bh, "Valid",
         "Nothing to do", fc=GREEN)
    _box(ax, 5.45, 5.4, 3.2, bh, "Expiring soon",
         "Reminders start 30 days\nbefore the expiry date", fc=AMBER)
    _box(ax, 8.8, 5.4, 2.9, bh, "Expired",
         "New contracts for you\nare blocked", fc=RED)
    _arrow(ax, (3.68, 5.4), (3.78, 5.4))
    _arrow(ax, (7.13, 5.4), (7.28, 5.4), label="expiry date",
           lab_dy=2.7, lab_color=GREY)
    ax.text(5.35, 1.6, "Send your renewed document to the Freelancers team at any point to return to Valid.",
            ha="center", fontsize=8.6, color=GREY, style="italic")
    return _save(fig, "documents_timeline.png")


def crop_registration_screenshot():
    """Form-only crop of the UAT capture — staff navigation and top bar removed."""
    src = os.path.join(EVIDENCE, "FL-REG-02.png")
    out = os.path.join(ASSETS, "registration_form.png")
    Image.open(src).crop((262, 120, 1428, 830)).save(out)
    print("  screenshot crop:", os.path.basename(out))
    return out


# =============================================================================
# Word building blocks
# =============================================================================
TOKEN = re.compile(r"(\[\[.+?\]\]|\*\*.+?\*\*)")


def rich(p, text, size=None, color=None, italic=False):
    """[[Field Name]] -> bold blue; **text** -> bold."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("[["):
            r = p.add_run(part[2:-2])
            r.font.bold = True
            r.font.color.rgb = W_FIELD_BLUE
        elif part.startswith("**"):
            r = p.add_run(part[2:-2])
            r.font.bold = True
            if color:
                r.font.color.rgb = color
        else:
            r = p.add_run(part)
            if color:
                r.font.color.rgb = color
        if size:
            r.font.size = Pt(size)
        if italic:
            r.font.italic = True
    return p


def _set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_margins(table, top=60, bottom=60, left=110, right=110):
    tblPr = table._tbl.tblPr
    m = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom),
                     ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tblPr.append(m)


def _field(run, instruction):
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve")
    i.text = instruction
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Heading 1", 17, W_PURPLE_DARK),
                              ("Heading 2", 13, W_PURPLE),
                              ("Heading 3", 11.5, W_GREY)):
        st = doc.styles[name]
        st.font.name = "Calibri"; st.font.size = Pt(size)
        st.font.bold = True; st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        rf.set(qn("w:ascii"), "Calibri"); rf.set(qn("w:hAnsi"), "Calibri")


def footer(doc):
    f = doc.sections[0].footer.paragraphs[0]
    f.text = ""
    r = f.add_run("Freelancer Guide  ·  i-Finance Freelancers  ·  Page ")
    r.font.size = Pt(8); r.font.color.rgb = W_GREY
    pg = f.add_run(); pg.font.size = Pt(8); pg.font.color.rgb = W_GREY
    _field(pg, "PAGE")
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER


def banner(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Cm(16.6)
    _set_cell_bg(cell, "7C4DBE")
    _cell_margins(t, top=260, bottom=260, left=280, right=280)
    cell.text = ""
    for i, (txt, size, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(size); r.font.bold = bold
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    banner(doc, [("i-Finance  ·  Freelancers", 13, False),
                 ("Freelancer Guide", 30, True),
                 ("Everything you need to register, work and get paid", 12, False)])
    doc.add_paragraph(); doc.add_paragraph()
    rows = [("For", "Freelancers working with the Finance Division"),
            ("Version", VERSION), ("Date", TODAY), ("Status", "Issued"),
            ("Issued by", "Finance Division — Freelancers Team")]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_margins(t)
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(c0, "EFE9F8")
        c0.width = Cm(4.5); c1.width = Cm(11.5)
    doc.add_page_break()


_bookmark_id = [100]


def h1(doc, text, bookmark=None):
    h = doc.add_heading(text, level=1)
    pPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "7C4DBE")
    pBdr.append(bottom); pPr.append(pBdr)
    if bookmark:
        _bookmark_id[0] += 1
        bid = str(_bookmark_id[0])
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), bid); bs.set(qn("w:name"), bookmark)
        be = OxmlElement("w:bookmarkEnd"); be.set(qn("w:id"), bid)
        h._p.insert(0, bs); h._p.append(be)
    return h


def contents_page(doc, sections):
    h1(doc, "What’s in this guide")
    p = doc.add_paragraph()
    rich(p, "Click any line to jump straight to that part.", italic=True,
         color=W_GREY)
    for i, (title, bm) in enumerate(sections, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.3)
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), bm)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), "0B5394")
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
        b = OxmlElement("w:b")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "23")
        rPr.append(col); rPr.append(u); rPr.append(b); rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
        t.text = f"{i}.  {title}"
        r.append(t)
        hl.append(r)
        para._p.append(hl)
    doc.add_page_break()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    _cell_margins(t)
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "7C4DBE")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            sval = str(val)
            key = sval.lower()
            if key in STATUS_COLORS:
                r = p.add_run(sval)
                r.font.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = STATUS_COLORS[key]
            else:
                rich(p, sval, size=10)
        if i % 2 == 1:
            for cell in t.rows[i + 1].cells:
                _set_cell_bg(cell, "F5F1FB")
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


CALLOUTS = {
    "note":      ("Good to know", "EAF2FB", W_FIELD_BLUE),
    "important": ("Important",    "FDECEA", W_RED),
    "tip":       ("Tip",          "EAF7EC", W_GREEN),
}


def callout(doc, kind, text):
    label, fill, color = CALLOUTS[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Cm(16.2)
    _set_cell_bg(cell, fill)
    _cell_margins(t, top=110, bottom=110, left=180, right=180)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label + "   ")
    r.font.bold = True; r.font.color.rgb = color; r.font.size = Pt(10.5)
    rich(p, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


_fig_no = [0]


def picture(doc, path, caption, width_cm=15.8):
    _fig_no[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {_fig_no[0]} — {caption}")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = W_GREY
    cap.paragraph_format.space_after = Pt(10)


def bullets(doc, items):
    for it in items:
        rich(doc.add_paragraph(style="List Bullet"), it)


def steps(doc, items):
    for it in items:
        rich(doc.add_paragraph(style="List Number"), it)


def para(doc, text):
    rich(doc.add_paragraph(), text)


# =============================================================================
# Build the document
# =============================================================================
def build():
    print("Generating diagrams…")
    d_journey = diagram_journey()
    d_reg = diagram_registration()
    d_pay = diagram_payment()
    d_del = diagram_deliverable()
    d_pcr = diagram_profile_change()
    d_doc = diagram_documents()
    shot_form = crop_registration_screenshot()

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    setup_styles(doc)
    footer(doc)
    cover(doc)

    SECTIONS = [
        ("Welcome", "sec_welcome"),
        ("Quick Reference — Keep This Handy", "sec_quickref"),
        ("How It All Works", "sec_journey"),
        ("Registering as a Freelancer", "sec_register"),
        ("Your Account and Vendor Number", "sec_account"),
        ("Keep Your Documents Valid", "sec_documents"),
        ("Your Contract and Payment Plan", "sec_contract"),
        ("Submitting Your Work", "sec_work"),
        ("Getting Paid", "sec_paid"),
        ("Updating Your Personal Details", "sec_details"),
        ("Renewing Your Contract", "sec_renewal"),
        ("Frequently Asked Questions", "sec_faq"),
        ("Words You Will Come Across", "sec_words"),
        ("Need Help?", "sec_help"),
    ]
    contents_page(doc, SECTIONS)

    # ── 1 Welcome ────────────────────────────────────────────────────────────
    h1(doc, "Welcome", "sec_welcome")
    para(doc, "Thank you for working with the Finance Division. This guide walks you, step by step, "
              "through everything you will do as a freelancer: registering, keeping your documents "
              "up to date, understanding your contract, submitting your work, getting paid, and "
              "changing your personal details. No technical knowledge is needed.")
    para(doc, "Throughout the guide:")
    bullets(doc, [
        "Names of boxes you fill in are shown in [[bold blue]] — for example [[Email]].",
        "**Bold words** are statuses or terms you will see on screen or hear from the team.",
        "Shaded boxes highlight things that are good to know, important warnings, or tips.",
    ])

    # ── 2 Quick reference ───────────────────────────────────────────────────
    h1(doc, "Quick Reference — Keep This Handy", "sec_quickref")
    table(doc,
          ["What", "Your details"],
          [["Portal address (URL)",
            "Provided in your welcome e-mail — write it here:  _________________________________"],
           ["Your username", "Your registered e-mail address, in lowercase"],
           ["Your password", "Set by the Freelancers team and sent to you separately. Contact them to reset it."],
           ["Your vendor number", "Looks like FL-VND-000123 — issued when your registration is approved. Write it here:  ______________"],
           ["Who to contact", "The Freelancers team at the Finance Division (see “Need Help?” at the end of this guide)"]],
          [5.2, 10.8])
    callout(doc, "note",
            "Every request you make gets a permanent reference number — for example FL-REG-000123 for "
            "a registration or FL-CON-000045 for a contract. Quote it whenever you contact the team; "
            "it is the fastest way for them to find your record.")

    # ── 3 Journey ───────────────────────────────────────────────────────────
    h1(doc, "How It All Works", "sec_journey")
    para(doc, "Your whole journey, from first registration to renewal, looks like this:")
    picture(doc, d_journey, "Your journey with the Finance Division, in seven steps.")
    para(doc, "The rest of this guide takes each step in turn.")

    # ── 4 Register ──────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Registering as a Freelancer", "sec_register")
    para(doc, "Registration is a one-time step. You can register yourself through the registration page, "
              "or a member of the Finance staff can register you — the form and the rules are the same.")

    doc.add_heading("What the form asks for", level=2)
    table(doc,
          ["Box on the form", "Required?", "Notes"],
          [["[[First Name (English)]] and [[Last Name (English)]]", "Yes",
            "Exactly as written in your passport or Emirates ID — these names go on your contract and payments."],
           ["[[First Name (Arabic)]] and [[Last Name (Arabic)]]", "Optional", "Recommended."],
           ["[[Date of Birth]]", "Yes", ""],
           ["[[Nationality]]", "Yes", "Choose from the list."],
           ["[[Emirates ID]]", "Yes for UAE nationals",
            "If your nationality is UAE, the form will not let you submit without it."],
           ["[[Passport Number]]", "Recommended", ""],
           ["[[Email]]", "Yes",
            "This becomes your **username** — use an address you check regularly."],
           ["[[Mobile]]", "Recommended", "Include the country code, e.g. +971…"],
           ["[[Specialization]]", "Recommended", "Your line of work, e.g. translator, photographer."],
           ["[[Photo]]", "Yes", "A recent photo — you cannot submit the form without one."],
           ["[[Notes]]", "Optional", "Anything else the team should know."]],
          [5.8, 3.2, 7.0])
    picture(doc, shot_form,
            "The registration form. Boxes marked * are required; the photo must be attached before "
            "the Submit button will work.", width_cm=15.5)
    callout(doc, "important",
            "Two things stop a registration from being submitted: a missing [[Photo]], and — for UAE "
            "nationals — a missing [[Emirates ID]]. The screen will tell you exactly which one is missing.")

    doc.add_heading("What happens after you submit", level=2)
    para(doc, "Your request is reviewed in two steps — first by a Freelancers manager, then the "
              "Freelancers team gives the final approval:")
    picture(doc, d_reg, "The registration journey, including what happens if something needs fixing.",
            width_cm=14.5)
    table(doc,
          ["Status you will see", "What it means", "What to do"],
          [["Draft", "Saved but not yet sent.", "Edit freely, then submit when ready."],
           ["Submitted", "Being reviewed.", "Nothing — you will be informed of the outcome."],
           ["Returned", "Something needs fixing; the reviewer’s comment says exactly what.",
            "Make the correction and submit again. Your reference number stays the same."],
           ["Approved", "You’re in! Your profile, vendor number and login are created automatically.",
            "Wait for your login details, then continue to your contract."],
           ["Rejected", "Declined, with the reason recorded. This is final for this request.",
            "Contact the team if you believe circumstances have changed — a new registration would be needed."]],
          [3.2, 6.6, 6.2])

    # ── 5 Account ───────────────────────────────────────────────────────────
    h1(doc, "Your Account and Vendor Number", "sec_account")
    bullets(doc, [
        "When your registration is approved, your **freelancer profile** is created automatically, "
        "together with your permanent **vendor number** (it looks like FL-VND-000123). The vendor "
        "number identifies you on every contract and payment — quote it when you contact the team.",
        "Your **login** is created at the same time: the username is your e-mail address in lowercase. "
        "The Freelancers team sets your first password and sends it to you separately — contact them "
        "if it has not arrived.",
        "If you have worked with the organisation before and register again, your old account is "
        "switched back on rather than duplicated — your history stays in one place.",
    ])
    callout(doc, "tip",
            "If your e-mail address changes later, don’t re-register. Ask for an e-mail change (see "
            "“Updating Your Personal Details”) — your username follows the new address automatically.")

    # ── 6 Documents ─────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Keep Your Documents Valid", "sec_documents")
    para(doc, "The team keeps copies of your identity and qualification documents — Emirates ID, "
              "passport, visa or residency, trade licence, certificates. Most of these carry an "
              "expiry date, and the system tracks it:")
    picture(doc, d_doc, "What happens as a document approaches and passes its expiry date.")
    table(doc,
          ["Document state", "What it means for you"],
          [["Valid", "Nothing to do."],
           ["Expiring soon", "The team starts receiving daily reminders 30 days before the expiry date — "
                             "expect them to ask you for the renewed document."],
           ["Expired", "If the document is one the organisation must hold (for example your Emirates ID), "
                       "**no new contract or renewal can be started for you** until the renewed document is provided."]],
          [3.6, 12.4])
    callout(doc, "important",
            "An expired required document is the most common reason a new contract gets stuck. Send "
            "renewed documents to the team **before** the old ones expire and you will never hit this block.")

    # ── 7 Contract ──────────────────────────────────────────────────────────
    h1(doc, "Your Contract and Payment Plan", "sec_contract")
    para(doc, "Your engagement is set out in a contract: what the work is, the start and end dates, "
              "the total amount (in AED), and **how you are paid** — monthly, weekly, or per piece of work. "
              "The moment your contract is fully approved it becomes **Active**, and a payment plan is "
              "created from it automatically: one line for each payment you will receive.")
    table(doc,
          ["How you are paid", "What your payment plan looks like"],
          [["Monthly", "One payment per calendar month, due at the end of each month."],
           ["Weekly", "One payment per week (Monday to Sunday), due on the Sunday."],
           ["Per piece of work", "No fixed plan — a payment line is added for each piece of work "
                                 "(report, visit, deliverable, hour or unit) once it is accepted."]],
          [4.2, 11.8]),
    para(doc, "Each payment is either the agreed amount per period, or — if only a total was agreed — "
              "the total split equally across the periods.")
    table(doc,
          ["Contract status", "What it means"],
          [["Draft", "Being prepared — not yet in force."],
           ["Submitted", "Going through approval."],
           ["Active", "In force. Work and payments happen only under an Active contract."],
           ["Completed", "Finished — naturally, or because a renewal replaced it."],
           ["Cancelled", "Will not go ahead."]],
          [3.4, 12.6])
    doc.add_heading("If the contract needs to change", level=2)
    bullets(doc, [
        "Changes to the amount, the dates or the payment pattern are made through a formal "
        "**amendment**, which goes through the same kind of approval as the contract itself.",
        "Payments already made to you are never touched — only the remaining, unpaid part of the "
        "plan is recalculated.",
        "The contract total can never be reduced below what you have already been paid.",
    ])

    # ── 8 Work ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Submitting Your Work", "sec_work")
    para(doc, "Each piece of work you hand over — a report, a site visit, a translation — is recorded "
              "as a **deliverable** with a title, the date, and the quantity. The responsible person "
              "then reviews it:")
    picture(doc, d_del, "What happens to each piece of work you deliver.")
    callout(doc, "note",
            "If your contract pays per piece of work, accepted deliverables are exactly what creates "
            "your payment lines. And depending on the agreed policy, a payment may need an accepted "
            "deliverable before it can go ahead — so the sooner your work is reviewed, the sooner you are paid.")

    # ── 9 Paid ──────────────────────────────────────────────────────────────
    h1(doc, "Getting Paid", "sec_paid")
    para(doc, "For every line on your payment plan, a **payment voucher** is prepared — that is the "
              "instruction that actually sends money to your bank account. Here is the full path of a payment:")
    picture(doc, d_pay, "The path of one payment, from due date to money in your account.")
    steps(doc, [
        "When a payment period is due, the team prepares the voucher. The amount comes straight from your payment plan.",
        "You provide an invoice — the [[Invoice Number]] and [[Invoice Date]] must be on the voucher "
        "before it can move forward. Send your invoice promptly when asked.",
        "The voucher is approved (payments of 50,000 AED or more get a second approval — this is normal "
        "and just takes one extra step).",
        "The payment is made to your **primary bank account**, and the voucher records the payment date "
        "and the bank reference.",
    ])
    table(doc,
          ["What you may hear / see", "What it means"],
          [["Draft", "The voucher is being prepared."],
           ["Submitted", "The voucher is in approval."],
           ["Approved", "Approved and queued for payment."],
           ["Paid", "Money sent — the payment date and bank reference are recorded."],
           ["Rejected", "The voucher will not be paid in its current form; the team will follow up."]],
          [4.2, 11.8])
    bullets(doc, [
        "One payment period = one voucher. The system itself prevents duplicates.",
        "All your payments together can never exceed your contract total — also enforced automatically.",
        "Payments go to your registered **primary** bank account. To change it, see the next section — "
        "never just e-mail new bank details.",
    ])

    # ── 10 Details ──────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Updating Your Personal Details", "sec_details")
    para(doc, "You never need to re-register to change your details. Instead you make a **change request** "
              "saying what should change and why; the Freelancers team reviews and approves it:")
    picture(doc, d_pcr, "How a change to your details is reviewed and applied.")
    table(doc,
          ["What you want to change", "What happens once approved"],
          [["Bank account", "Your new account becomes the one your payments go to; the old one is kept "
                            "on file but no longer used. A payment already on its way is not redirected."],
           ["E-mail address", "Your profile **and your username** change to the new address — sign in "
                              "with the new e-mail from then on."],
           ["Phone number", "Your mobile number is updated."],
           ["Anything else", "The team reviews your description and applies the change by hand."]],
          [4.5, 11.5])
    callout(doc, "important",
            "Bank account changes only take effect after approval. This protects you — money cannot be "
            "redirected on the strength of a single e-mail.")

    # ── 11 Renewal ──────────────────────────────────────────────────────────
    h1(doc, "Renewing Your Contract", "sec_renewal")
    bullets(doc, [
        "When your contract is coming to an end, the team can raise a **renewal** — a request to "
        "continue with new dates and a new amount.",
        "The new period always starts **after** the current contract ends — there is never a gap "
        "in between paid twice, and never an overlap.",
        "Once approved, a brand-new contract is created (already Active, with its own payment plan) "
        "and the old one is closed as **Completed**. The new contract remembers which one it continues.",
        "For a renewal to go ahead, your profile must be in good standing and your required documents "
        "valid — one more reason to renew documents early.",
    ])

    # ── 12 FAQ ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Frequently Asked Questions", "sec_faq")
    table(doc,
          ["Question", "Answer"],
          [["I haven’t received my login details.",
            "Logins are created the moment your registration is approved, but your first password is "
            "sent separately by the Freelancers team — contact them if it has not arrived."],
           ["The form won’t let me submit my registration.",
            "Check the [[Photo]] is attached and, if you are a UAE national, that the [[Emirates ID]] "
            "is filled in. The message on screen says which one is missing."],
           ["My registration came back as Returned. Is that bad?",
            "Not at all — it just needs a correction. Read the reviewer’s comment, fix that one thing, "
            "and submit again. Your reference number stays the same."],
           ["When exactly will I be paid?",
            "Each line of your payment plan has a due date. Around that date a voucher is prepared, "
            "approved and paid. You can always ask the team for the status of a specific period."],
           ["My payment seems delayed. What’s the usual cause?",
            "In order of likelihood: your invoice (number and date) hasn’t been provided yet; the work "
            "for that period hasn’t been accepted yet; or a required document of yours has expired. "
            "All three are quick to fix."],
           ["My Emirates ID / visa is about to expire.",
            "Send the renewed document to the team as soon as you have it. Expired required documents "
            "block new contracts and renewals — but never payments already approved."],
           ["Can I have more than one contract at the same time?",
            "Yes. Each contract is independent, with its own payment plan and payments."],
           ["Who sees my information?",
            "Only the Finance Division staff who run the freelancers process. Your bank details are "
            "used solely to pay you."]],
          [5.6, 10.4])

    # ── 13 Glossary ─────────────────────────────────────────────────────────
    h1(doc, "Words You Will Come Across", "sec_words")
    table(doc,
          ["Word", "Plain meaning"],
          [["Vendor number", "Your permanent supplier reference (FL-VND-…). Think of it as your account number with Finance."],
           ["Payment plan / schedule", "The list of payments your contract will produce — one line per period."],
           ["Payment voucher", "The instruction that sends one of those payments to your bank."],
           ["Deliverable", "One piece of work you hand over, recorded so it can be accepted."],
           ["Amendment", "An approved change to your contract (amount, dates or payment pattern)."],
           ["Renewal", "A new contract that continues an ending one."],
           ["Change request", "Your formal request to update bank, e-mail or phone details."],
           ["Required document", "A document Finance must hold in valid form (e.g. Emirates ID). If it expires, new contracts wait until it is renewed."],
           ["Primary bank account", "The single account your payments are sent to."]],
          [4.4, 11.6])

    # ── 14 Help ─────────────────────────────────────────────────────────────
    h1(doc, "Need Help?", "sec_help")
    bullets(doc, [
        "**Anything about your registration, contract, documents or payments** — contact the "
        "Freelancers team at the Finance Division.",
        "**Login or password problems** — the same team will reset your access.",
        "Always mention your **vendor number** (FL-VND-…) or the reference number of the request "
        "you are asking about — it gets you an answer much faster.",
    ])
    para(doc, "")
    callout(doc, "note",
            "This guide describes the standard policies in force at the date on the cover. Individual "
            "contracts may carry additional agreed terms — your contract document always prevails.")

    # ── save ────────────────────────────────────────────────────────────────
    out = os.path.join(HERE, OUT_NAME)
    try:
        if os.path.exists(out):
            os.remove(out)
        doc.save(out)
        print("Created:", out)
    except PermissionError:
        alt = os.path.join(HERE, OUT_NAME.replace(".docx", ".new.docx"))
        doc.save(alt)
        print("! %s is open in Word — saved as %s instead. "
              "Close the document and rename the file." % (OUT_NAME, os.path.basename(alt)))


if __name__ == "__main__":
    build()
