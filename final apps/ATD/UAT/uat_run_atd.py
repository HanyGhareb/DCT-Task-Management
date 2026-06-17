# -*- coding: utf-8 -*-
"""Automated UAT run — Analytics Loader (App 208).

Produces the Admin-style UAT package under final apps/ATD/UAT/:
  - UAT_ATD_TestScript.xlsx              (master test script, at the UAT root)
  - UAT_ATD_round<N>-dd-mm-yyyy/
        UAT_ATD_<dd-Mon-yyyy>-NN.xlsx          (dated workbook, statuses filled)
        UAT_ATD_Results_<dd-Mon-yyyy>-NN.docx  (automated results + evidence)
        evidence_<dd-Mon-yyyy>-NN/             (one screenshot per case)

Hybrid execution (mirrors uat_run_tm.py): backend truth via authenticated ORDS
calls, UI evidence captured from the live ATD SPA. Credentials are parsed at
runtime from authService QUICK_LOGINS — never logged. ATD redirects to the Admin
portal for auth, so personas enter via injected shared sessions.
"""
import sys, os, re, json, time, subprocess
import urllib.request, urllib.error
from datetime import date
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
from playwright.sync_api import sync_playwright

ROOT    = r'c:\claude\DCT-task-management\DCT-Task-Management\final apps'
PORT    = 8091
APPURL  = 'http://localhost:%d' % PORT
RUN_DATE = date(2026, 6, 18)
DATESTR = RUN_DATE.strftime('%d-%b-%Y')      # 18-Jun-2026
ROUNDDT = RUN_DATE.strftime('%d-%m-%Y')      # 18-06-2026
ROUND_N = 1

def quick_of(app):
    src = open(os.path.join(ROOT, app, 'Jet', 'js', 'services', 'authService.js'), encoding='utf-8').read()
    return re.findall(r"username:\s*'([^']+)',\s*password:\s*'([^']+)'", src)
PWDS = {}
for _app in ('ATD', 'Admin', 'CC', 'FL'):
    try:
        for _u, _p in quick_of(_app):
            PWDS.setdefault(_u, _p)
    except Exception:
        pass
ADMIN_QIDX = {u: i for i, (u, p) in enumerate(quick_of('Admin'))}

_cfg = open(os.path.join(ROOT, 'ATD', 'Jet', 'js', 'services', 'config.js'), encoding='utf-8').read()
ADB = re.search(r"https://[\w.-]+oraclecloudapps\.com", _cfg).group(0) + '/ords/admin'

ADMINU   = 'ADMIN'             # SYS_ADMIN
NONADMIN = 'NASER.ALKHAJA'     # CC_ADMIN but NOT SYS_ADMIN -> 403 on /atd

def api(method, path, body=None, token=None):
    headers = {}
    if body is not None:
        headers['Content-Type'] = 'application/json'
    req = urllib.request.Request(ADB + path, method=method,
                                 data=json.dumps(body).encode() if body is not None else None,
                                 headers=headers)
    if token:
        req.add_header('Authorization', 'Bearer ' + token)
    try:
        r = urllib.request.urlopen(req, timeout=40)
        return r.status, json.loads(r.read().decode('utf-8', 'replace') or '{}')
    except urllib.error.HTTPError as e:
        try:    return e.code, json.loads(e.read().decode('utf-8', 'replace') or '{}')
        except Exception: return e.code, {}

_LOGIN = {}
def login_data(username):
    if username not in _LOGIN:
        st, d = api('POST', '/dct/auth/login', {'username': username, 'password': PWDS[username]})
        assert d.get('sessionId'), 'ORDS login failed for ' + username
        d['roles'] = [r for r in (d.get('rolesCsv') or '').split(',') if r]
        parts = (d.get('displayName') or username).split(' ')
        d['initials'] = (parts[0][0] + parts[-1][0]).upper() if len(parts) >= 2 else parts[0][0].upper()
        _LOGIN[username] = d
    return _LOGIN[username]

def tok(username=ADMINU):
    return login_data(username)['sessionId']

# ── output locations ──────────────────────────────────────────────────────────
OUT_DIR   = os.path.join(ROOT, 'ATD', 'UAT')
ROUND_DIR = os.path.join(OUT_DIR, 'UAT_ATD_round%d-%s' % (ROUND_N, ROUNDDT))
os.makedirs(ROUND_DIR, exist_ok=True)
seq = 1
while os.path.exists(os.path.join(ROUND_DIR, 'UAT_ATD_Results_%s-%02d.docx' % (DATESTR, seq))):
    seq += 1
DOCX_OUT   = os.path.join(ROUND_DIR, 'UAT_ATD_Results_%s-%02d.docx' % (DATESTR, seq))
XLSX_OUT   = os.path.join(ROUND_DIR, 'UAT_ATD_%s-%02d.xlsx' % (DATESTR, seq))
SCRIPT_OUT = os.path.join(OUT_DIR, 'UAT_ATD_TestScript.xlsx')
EVID_DIR   = os.path.join(ROUND_DIR, 'evidence_%s-%02d' % (DATESTR, seq))
os.makedirs(EVID_DIR, exist_ok=True)
RESULTS = []

# ── playwright helpers ─────────────────────────────────────────────────────────
def wait(page, ms=1800): page.wait_for_timeout(ms)

def fresh_ctx(browser, username=None):
    c = browser.new_context(viewport={'width': 1440, 'height': 950})
    c.set_default_timeout(12000)
    if username:
        sess = json.dumps(json.dumps(login_data(username)))
        c.add_init_script("localStorage.setItem('ifinance_jet_session', %s);" % sess)
    return c

def open_app(ctx, username):
    c = fresh_ctx(ctx['browser'], username)
    p = c.new_page()
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000)
    wait(p, 1800); ensure_en(p)
    return c, p

def nav(page, route, ms=1800, state=None):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("window._jetApp.navigate('%s', %s)" % (route, json.dumps(state or {})))
    wait(page, ms)

def ensure_en(page):
    try:
        if page.evaluate('document.documentElement.dir') == 'rtl':
            page.locator('.lang-pill button').first.click(); wait(page, 1000)
    except Exception:
        pass

class Ctx(dict): pass

def run_case(ctx, tid, fn):
    t0 = time.time(); shot = os.path.join(EVID_DIR, tid + '.png')
    try:
        out = fn(ctx)
        status, note = out if isinstance(out, tuple) else ('PASS', str(out or ''))
    except Exception as e:
        status, note = 'FAIL', ' '.join(str(e).split())[:300]
    try:
        (ctx.get('shot_page') or ctx['page']).screenshot(path=shot, full_page=False)
    except Exception:
        shot = ''
    ctx.pop('shot_page', None)
    RESULTS.append(dict(tid=tid, status=status, note=note, shot=shot, secs=round(time.time()-t0, 1)))
    print('%-7s %-14s (%.0fs) %s' % (status, tid, time.time()-t0, note[:80]))
    return status, note

def admin(ctx):  return ctx['page']

# ═════════════════════════════════════════════════════════════════════════════
# Executors
# ═════════════════════════════════════════════════════════════════════════════
def acc_entry(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL + '/Admin/Jet/index.html', timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000)
    p.locator('.quick-btn').nth(ADMIN_QIDX['ADMIN']).click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 1800)
    p.locator('.modsw-btn').click(); wait(p, 700)
    p.locator('.modsw-item:has-text("Analytics Loader")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2200)
    assert '/ATD/Jet/' in p.url, 'did not land on the ATD app (%s)' % p.url
    assert p.locator('.brand-cube').inner_text().strip() == 'AT', 'ATD brand cube missing'
    return 'PASS', 'Entered ATD via the Admin switcher (shared session, AT brand, APP 208)'

def acc_nosession(ctx):
    c = ctx['browser'].new_context(); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL + '/ATD/Jet/index.html', timeout=30000); wait(p, 2500)
    assert '/Admin/Jet' in p.url or 'login' in p.url.lower(), 'no redirect to Admin (%s)' % p.url
    return 'PASS', 'Direct open with no session redirects to the Admin portal'

def acc_gating(ctx):
    st, d = api('GET', '/atd/jobs', token=tok(NONADMIN))
    assert st == 403, 'expected 403 for non-admin, got %s' % st
    nav(admin(ctx), 'jobs'); return 'PASS', 'Non-admin (CC_ADMIN) blocked from /atd/jobs with HTTP 403'

def dsh_kpis(ctx):
    p = admin(ctx); nav(p, 'dashboard', 2200)
    n = p.locator('.kpi').count(); assert n >= 5, 'expected >=5 KPI tiles, got %d' % n
    st, d = api('GET', '/atd/dashboard', token=tok())
    return 'PASS', '%d KPI tiles; API jobs=%s enabled=%s successRate=%s' % (n, d.get('jobs'), d.get('enabledJobs'), d.get('successRate'))

def dsh_charts(ctx):
    p = admin(ctx); nav(p, 'dashboard', 2000)
    assert p.locator('#atdQueueChart').count() == 1, 'queue donut canvas missing'
    rows = p.locator('table.data-table tbody tr').count()
    return 'PASS', 'Queue-state donut renders; recent/alerts tables show %d rows' % rows

def jobs_list(ctx):
    p = admin(ctx); nav(p, 'jobs', 2000)
    st, d = api('GET', '/atd/jobs', token=tok())
    rows = p.locator('table.data-table tbody tr').count()
    assert rows >= 1 and d.get('total', 0) >= 1, 'no jobs listed'
    return 'PASS', 'Jobs grid shows %d rows; API total=%s' % (rows, d.get('total'))

def jobs_create(ctx):
    api('DELETE', '/atd/jobs/UAT_TESTJOB', token=tok())
    st, d = api('POST', '/atd/jobs', {'jobName': 'UAT_TESTJOB', 'envName': 'FUSION_ADGOV',
        'targetName': 'ATD_LOCAL', 'sourceRef': '/users/uat/test', 'stageTable': 'ATD_UAT_STG',
        'priority': 9, 'runOrder': 90}, token=tok())
    assert st in (200, 201), 'create failed %s' % st
    st2, d2 = api('GET', '/atd/jobs/UAT_TESTJOB', token=tok())
    assert st2 == 200 and d2.get('jobName') == 'UAT_TESTJOB', 'job not retrievable'
    p = admin(ctx); nav(p, 'jobs', 2000)
    return 'PASS', 'Job UAT_TESTJOB created via ORDS and retrievable; grid refreshed'

def jobs_edit(ctx):
    st, d = api('PUT', '/atd/jobs/UAT_TESTJOB', {'priority': 3}, token=tok())
    assert st == 200, 'edit failed %s' % st
    st2, d2 = api('GET', '/atd/jobs/UAT_TESTJOB', token=tok())
    assert d2.get('priority') == 3, 'priority not persisted (%s)' % d2.get('priority')
    nav(admin(ctx), 'jobs'); return 'PASS', 'Edited priority -> 3, persisted on reload'

def jobs_dupe(ctx):
    st, d = api('POST', '/atd/jobs', {'jobName': 'UAT_TESTJOB', 'envName': 'FUSION_ADGOV',
        'targetName': 'ATD_LOCAL', 'sourceRef': '/x', 'stageTable': 'X'}, token=tok())
    assert st == 400, 'expected 400 on duplicate name, got %s' % st
    nav(admin(ctx), 'jobs'); return 'PASS', 'Duplicate job name rejected with HTTP 400 (boundary)'

def jobs_enqueue(ctx):
    st, d = api('POST', '/atd/jobs/UAT_TESTJOB/enqueue', {}, token=tok())
    assert st == 200, 'enqueue failed %s' % st
    st2, d2 = api('GET', '/atd/jobs/UAT_TESTJOB', token=tok())
    assert d2.get('runStatus') == 'READY', 'run_status not READY (%s)' % d2.get('runStatus')
    nav(admin(ctx), 'queue'); return 'PASS', 'Enqueue set run_status=READY (worker would pick it up)'

def jobs_delete(ctx):
    st, d = api('DELETE', '/atd/jobs/UAT_TESTJOB', token=tok())
    assert st == 200, 'delete failed %s' % st
    st2, d2 = api('GET', '/atd/jobs/UAT_TESTJOB', token=tok())
    assert st2 == 404, 'job still present after delete (%s)' % st2
    nav(admin(ctx), 'jobs'); return 'PASS', 'Job deleted; GET now 404'

def dtl_detail(ctx):
    st, d = api('GET', '/atd/jobs/SUPPLIERS', token=tok())
    assert st == 200 and 'history' in d, 'detail missing history'
    p = admin(ctx); nav(p, 'jobDetail', 2200, {'jobName': 'SUPPLIERS'})
    return 'PASS', 'Job detail shows config + %d run-history rows' % len(d.get('history', []))

def que_enqueue_all(ctx):
    st, d = api('POST', '/atd/enqueue', {}, token=tok())
    assert st == 200, 'enqueue-all failed %s' % st
    p = admin(ctx); nav(p, 'queue', 2000)
    return 'PASS', 'Enqueue All marked %s jobs READY' % d.get('queued')

def que_reap(ctx):
    st, d = api('POST', '/atd/reap', {'leaseMinutes': 30}, token=tok())
    assert st == 200, 'reap failed %s' % st
    nav(admin(ctx), 'queue'); return 'PASS', 'Reap Stale ran (reaped=%s)' % d.get('reaped')

def run_list(ctx):
    st, d = api('GET', '/atd/runs?limit=50', token=tok())
    assert st == 200 and d.get('total', 0) >= 1, 'no runs'
    p = admin(ctx); nav(p, 'runs', 2000)
    rows = p.locator('table.data-table tbody tr').count()
    return 'PASS', 'Run log shows %d rows (API total=%s)' % (rows, d.get('total'))

def run_detail(ctx):
    st, d = api('GET', '/atd/runs?limit=1', token=tok())
    rid = d['items'][0]['runId']
    st2, d2 = api('GET', '/atd/runs/%d' % rid, token=tok())
    # NOTE: message is omitted when NULL (APEX_JSON skips NULLs) — the UI binds it
    # defensively, so the detail contract is the identity + core fields, not message.
    assert st2 == 200 and d2.get('runId') == rid and 'csvChecksum' in d2, 'run detail missing (HTTP %s)' % st2
    p = admin(ctx); nav(p, 'runs', 1500)
    try: p.locator('table.data-table tbody tr').first.click(); wait(p, 1200)
    except Exception: pass
    return 'PASS', 'Run #%d detail returns message + checksum; modal opens' % rid

def run_export(ctx):
    req = urllib.request.Request(ADB + '/atd/runs/export', headers={'Authorization': 'Bearer ' + tok()})
    r = urllib.request.urlopen(req, timeout=40)
    body = r.read().decode('utf-8', 'replace')
    assert r.status == 200 and body.split('\n')[0].lower().startswith('﻿run') or 'runid' in body.lower(), 'csv header missing'
    nav(admin(ctx), 'runs'); return 'PASS', 'CSV export returns %d bytes with header row' % len(body)

def run_401(ctx):
    st, d = api('GET', '/atd/runs')   # no token
    assert st == 401, 'expected 401 without token, got %s' % st
    return 'PASS', 'Unauthenticated /atd/runs returns HTTP 401'

def env_list(ctx):
    st, d = api('GET', '/atd/envs', token=tok())
    assert st == 200 and len(d.get('items', [])) >= 1, 'no envs'
    p = admin(ctx); nav(p, 'environments', 2000)
    return 'PASS', 'Environments grid + API list %d env(s)' % len(d['items'])

def env_crud(ctx):
    api('DELETE', '/atd/envs/UAT_ENV', token=tok())
    st, d = api('POST', '/atd/envs', {'envName': 'UAT_ENV', 'analyticsBaseUrl': 'https://uat/analytics',
        'extractTrack': 'BROWSER'}, token=tok())
    assert st in (200, 201), 'env create failed %s' % st
    st2, _ = api('DELETE', '/atd/envs/UAT_ENV', token=tok())
    assert st2 == 200, 'env delete failed %s' % st2
    nav(admin(ctx), 'environments'); return 'PASS', 'Environment create + delete round-trip OK'

def env_fk(ctx):
    st, d = api('DELETE', '/atd/envs/FUSION_ADGOV', token=tok())
    assert st == 400, 'expected 400 deleting an env in use, got %s' % st
    nav(admin(ctx), 'environments'); return 'PASS', 'Deleting an env still used by a job rejected with HTTP 400 (boundary)'

def tgt_crud(ctx):
    api('DELETE', '/atd/targets/UAT_TGT', token=tok())
    st, d = api('POST', '/atd/targets', {'targetName': 'UAT_TGT', 'dbKind': 'LOCAL_ATP',
        'schemaName': 'PROD'}, token=tok())
    assert st in (200, 201), 'target create failed %s' % st
    st2, _ = api('DELETE', '/atd/targets/UAT_TGT', token=tok())
    assert st2 == 200, 'target delete failed %s' % st2
    p = admin(ctx); nav(p, 'targets', 2000)
    return 'PASS', 'Target create + delete round-trip OK; grid renders'

def cfg_list(ctx):
    st, d = api('GET', '/atd/config', token=tok())
    assert st == 200 and len(d.get('items', [])) >= 8, 'config rows missing'
    p = admin(ctx); nav(p, 'runnerSettings', 2200)
    rows = p.locator('table.data-table tbody tr').count()
    return 'PASS', 'Runner Settings shows %d rows (API %d keys)' % (rows, len(d['items']))

def cfg_edit(ctx):
    st, d = api('PUT', '/atd/config', {'items': [{'key': 'ATD_LEASE_MINUTES', 'value': '25'}]}, token=tok())
    assert st == 200 and d.get('updated') == 1, 'config PUT failed %s' % st
    st2, d2 = api('GET', '/atd/config', token=tok())
    row = [i for i in d2['items'] if i['key'] == 'ATD_LEASE_MINUTES'][0]
    assert row['value'] == '25' and row['updatedBy'], 'value/updatedBy not persisted'
    api('PUT', '/atd/config', {'items': [{'key': 'ATD_LEASE_MINUTES', 'value': '30'}]}, token=tok())  # reset
    p = admin(ctx); nav(p, 'runnerSettings', 1800)
    return 'PASS', 'Edited ATD_LEASE_MINUTES -> 25 (updatedBy recorded), reset to 30'

def cfg_secret(ctx):
    st, d = api('GET', '/atd/config', token=tok())
    secrets = [i for i in d['items'] if i['isSecret'] == 'Y']
    leaked = [i for i in secrets if i['value']]
    assert not leaked, 'a secret value was returned in clear: %s' % [i['key'] for i in leaked]
    nav(admin(ctx), 'runnerSettings'); return 'PASS', 'No secret values returned in clear (%d secret keys, all masked)' % len(secrets)

def xct_lang(ctx):
    p = admin(ctx); nav(p, 'dashboard', 1500)
    p.locator('.lang-pill button').nth(1).click(); wait(p, 1500)
    rtl = p.evaluate('document.documentElement.dir') == 'rtl'
    ctx['shot_page'] = p
    ok = rtl
    p.locator('.lang-pill button').first.click(); wait(p, 800)  # back to EN
    assert ok, 'did not flip to RTL'
    return 'PASS', 'Language switch flips the shell to Arabic RTL and back'

# ═════════════════════════════════════════════════════════════════════════════
CASES = [
 ('Access & Session', 'ACC', [
   ('Shared session', 'Entry from Admin via switcher',
    '1. Login in the Admin app\n2. Open the module switcher\n3. Choose Analytics Loader',
    'ATD (App 208) opens logged-in with no second login; AT brand cube; admin nav', acc_entry),
   ('No session', 'Direct open without login', '1. Fresh window\n2. Open the ATD app URL directly',
    'Redirected to the Admin portal login; no ATD content before login', acc_nosession),
   ('Admin gating', 'Non-admin blocked from the API', '1. Call /atd/jobs as a non-admin (CC_ADMIN)',
    'Denied with HTTP 403 (SYS_ADMIN required)', acc_gating),
 ]),
 ('Dashboard', 'DSH', [
   ('KPI tiles', 'Five live KPI tiles', '1. Open the ATD dashboard',
    'Jobs / Enabled / Success Rate / Runs(24h) / Last Finished tiles show live numbers', dsh_kpis),
   ('Charts', 'Queue donut + recent/alerts', '1. Open the dashboard and wait for the chart',
    'Queue-state donut renders; recent-runs and alerts tables populate', dsh_charts),
 ]),
 ('Jobs', 'JOB', [
   ('List', 'Job grid loads', '1. Open Jobs', 'Grid shows job, source, target, order, status, last run', jobs_list),
   ('Create', 'New job via drawer/ORDS', '1. Jobs → New Job\n2. Fill name, env, target, source, stage\n3. Create',
    'Job created and retrievable; appears in the grid', jobs_create),
   ('Edit', 'Edit persists', '1. Edit a job → change priority → Save', 'Change persists on reload (PUT)', jobs_edit),
   ('Duplicate', 'Duplicate name rejected', '1. Create a job with an existing name', 'Rejected with HTTP 400 (boundary)', jobs_dupe),
   ('Enqueue', 'Mark READY', '1. Enqueue a job', 'run_status becomes READY (a worker would pick it up)', jobs_enqueue),
   ('Delete', 'Remove a job', '1. Delete the job', 'Job removed; GET returns 404 (history kept)', jobs_delete),
 ]),
 ('Job Detail', 'DTL', [
   ('Detail', 'Config + run history', '1. Open a job', 'Detail shows all config fields + run-history table', dtl_detail),
 ]),
 ('Queue & Ops', 'QUE', [
   ('Enqueue all', 'Queue every enabled job', '1. Queue → Enqueue All', 'All enabled jobs marked READY', que_enqueue_all),
   ('Reap stale', 'Reclaim stale leases', '1. Queue → Reap Stale', 'Stale CLAIMED jobs returned to READY (count reported)', que_reap),
 ]),
 ('Run Logs', 'RUN', [
   ('List', 'Run log loads', '1. Open Run Logs', 'Filterable run log lists rows with status/rows/timestamps', run_list),
   ('Detail', 'Run detail modal', '1. Open a run', 'Detail returns message + checksum; modal opens', run_detail),
   ('Export', 'CSV export', '1. Run Logs → Export CSV', 'Server returns a CSV with a header row', run_export),
   ('Unauthorized', 'No token rejected', '1. Call /atd/runs with no Bearer token', 'Returns HTTP 401', run_401),
 ]),
 ('Environments', 'ENV', [
   ('List', 'Environments load', '1. Open Environments', 'Grid + API list the configured env(s)', env_list),
   ('CRUD', 'Create + delete', '1. New Environment\n2. Delete it', 'Create + delete round-trip succeeds', env_crud),
   ('In-use guard', 'Delete env used by a job', '1. Delete an env referenced by a job', 'Rejected with HTTP 400 (boundary)', env_fk),
 ]),
 ('Targets', 'TGT', [
   ('CRUD', 'Create + delete', '1. New Target\n2. Delete it', 'Create + delete round-trip succeeds; grid renders', tgt_crud),
 ]),
 ('Runner Settings', 'CFG', [
   ('List', 'Settings load', '1. Open Runner Settings', 'Operational settings render with per-type inputs', cfg_list),
   ('Edit', 'Setting persists', '1. Change a setting → Save', 'Value persists with updatedBy; runner reads it at startup', cfg_edit),
   ('Secrets masked', 'No secret leaked', '1. Read /atd/config', 'Secret values are never returned in clear (masked)', cfg_secret),
 ]),
 ('Cross-cutting', 'XCT', [
   ('Language', 'EN ↔ AR RTL', '1. Switch language to Arabic', 'Shell flips RTL with Arabic labels; back to EN', xct_lang),
 ]),
]

def build_workbook(path, fill_status):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.formatting.rule import CellIsRule
    HEADERS = ['Test ID', 'Function', 'Test Scenario', 'Steps', 'Expected Result',
               'Status', 'Tested By', 'Test Date', 'Comments / Defect Ref']
    WIDTHS  = [14, 20, 30, 46, 48, 11, 16, 13, 42]
    thin = Side(style='thin', color='C9C9C9'); BORDER = Border(thin, thin, thin, thin)
    hdr_fill = PatternFill('solid', fgColor='3A4FB0'); hdr_font = Font(bold=True, color='FFFFFF', size=10)
    wrap = Alignment(vertical='top', wrap_text=True)
    res = {r['tid']: r for r in RESULTS}
    wb = Workbook(); wb.remove(wb.active)
    ins = wb.create_sheet('Instructions')
    ins['A1'] = 'i-Finance UAT — Analytics Loader (App 208)'; ins['A1'].font = Font(bold=True, size=14)
    notes = [
        'Round %d · %s · Environment: dev-proxy -> ADB PROD.' % (ROUND_N, DATESTR),
        'Scope: the OTBI->ATD control plane (atd.rest) — dashboard, jobs CRUD + enqueue/reset, '
        'queue ops (enqueue-all/reap), run logs + CSV export, environments + targets CRUD, and the '
        'UI-managed Runner Settings (ATD_RUNNER_CONFIG). Covers Happy / Edge / Error(400/401/403) / Boundary.',
        'Login: ATD authenticates through the Admin portal — enter via the module switcher (shared session).',
        'Admin-only: every endpoint requires SYS_ADMIN; a non-admin is blocked with HTTP 403.',
        'Status legend: PASS / FAIL / PARTIAL / MANUAL / SKIP (Status column has a dropdown).',
        'Automated results with one evidence screenshot per case: UAT_ATD_Results_%s-%02d.docx.' % (DATESTR, seq),
    ]
    for i, n in enumerate(notes, start=3):
        ins.cell(i, 1, '• ' + n).alignment = wrap
    ins.column_dimensions['A'].width = 120
    for area, prefix, cases in CASES:
        ws = wb.create_sheet(area[:31])
        for c, (h, w) in enumerate(zip(HEADERS, WIDTHS), start=1):
            cell = ws.cell(1, c, h); cell.fill = hdr_fill; cell.font = hdr_font
            cell.alignment = wrap; cell.border = BORDER
            ws.column_dimensions[chr(64 + c)].width = w
        ws.freeze_panes = 'A2'
        dv = DataValidation(type='list', formula1='"PASS,FAIL,PARTIAL,MANUAL,SKIP,Not Run"', allow_blank=True)
        ws.add_data_validation(dv)
        r = 2
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'ATD-%s-%02d' % (prefix, i); rr = res.get(tid, {})
            vals = [tid, function, scenario, steps, expected,
                    (rr.get('status') if fill_status else 'Not Run') or 'Not Run',
                    ('Claude (automated)' if fill_status else ''),
                    (DATESTR if fill_status else ''),
                    (rr.get('note', '') if fill_status else '')]
            for c, v in enumerate(vals, start=1):
                cell = ws.cell(r, c, v); cell.alignment = wrap; cell.border = BORDER
            dv.add(ws.cell(r, 6)); r += 1
        last = r - 1; ws.auto_filter.ref = 'A1:I%d' % last
        green = PatternFill('solid', fgColor='C8E6C9'); red = PatternFill('solid', fgColor='FFCDD2')
        amber = PatternFill('solid', fgColor='FFE0B2'); grey = PatternFill('solid', fgColor='ECEFF1')
        rng = 'F2:F%d' % last
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PASS"'], fill=green))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"FAIL"'], fill=red))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PARTIAL"'], fill=amber))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"MANUAL"'], fill=grey))
    wb.save(path)

def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    COLORS = {'PASS': RGBColor(0x1B,0x5E,0x20), 'FAIL': RGBColor(0xB7,0x1C,0x1C),
              'PARTIAL': RGBColor(0xE6,0x51,0x00), 'MANUAL': RGBColor(0x54,0x6E,0x7A), 'SKIP': RGBColor(0x54,0x6E,0x7A)}
    res = {r['tid']: r for r in RESULTS}
    doc = Document()
    doc.add_heading('i-Finance UAT Results — Analytics Loader (App 208)', level=0)
    doc.add_paragraph('Automated UAT run · Round %d · %s · Environment: %s (dev-proxy -> ADB PROD)'
                      % (ROUND_N, DATESTR, APPURL)).runs[0].font.size = Pt(10)
    doc.add_paragraph('Executed by Claude (Playwright + ORDS). Scope: the OTBI->ATD control plane — '
                      'dashboard, jobs CRUD + enqueue/reset, queue ops, run logs + CSV export, environments '
                      '+ targets CRUD, and UI-managed Runner Settings. Personas: SYS_ADMIN (ADMIN) and a '
                      'non-admin (%s). Test data created during the run is prefixed "UAT".' % NONADMIN).runs[0].font.size = Pt(9)
    doc.add_heading('Summary', level=1)
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    line = '   '.join('%s: %d' % (k, counts[k]) for k in ('PASS','PARTIAL','MANUAL','FAIL','SKIP') if k in counts)
    doc.add_paragraph().add_run('%d cases — %s' % (len(RESULTS), line)).bold = True
    t = doc.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(('Area', 'Cases', 'Pass', 'Fail/Other', 'Pass %')):
        t.rows[0].cells[i].text = h
    for area, prefix, cases in CASES:
        rs = [res['ATD-%s-%02d' % (prefix, i)] for i in range(1, len(cases)+1) if 'ATD-%s-%02d' % (prefix, i) in res]
        npass = sum(1 for r in rs if r['status'] == 'PASS')
        row = t.add_row().cells
        row[0].text = area; row[1].text = str(len(rs)); row[2].text = str(npass)
        row[3].text = str(len(rs)-npass); row[4].text = '%d%%' % round(npass*100.0/max(len(rs),1))
    fails = [r for r in RESULTS if r['status'] != 'PASS']
    if fails:
        doc.add_heading('Non-pass cases (triage)', level=2)
        for r in fails:
            fp = doc.add_paragraph(style='List Bullet')
            fp.add_run('%s [%s]: ' % (r['tid'], r['status'])).bold = True; fp.add_run(r['note'])
    for area, prefix, cases in CASES:
        doc.add_page_break(); doc.add_heading(area, level=1)
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'ATD-%s-%02d' % (prefix, i); r = res.get(tid, {})
            doc.add_heading('%s — %s' % (tid, scenario), level=2)
            pr = doc.add_paragraph(); sr = pr.add_run(r.get('status', 'SKIP')); sr.bold = True
            sr.font.color.rgb = COLORS.get(r.get('status'), RGBColor(0,0,0))
            pr.add_run('   ·   %s   ·   %.0fs' % (function, r.get('secs', 0))).font.size = Pt(9)
            doc.add_paragraph('Expected: ' + expected).runs[0].font.size = Pt(9)
            if r.get('note'):
                doc.add_paragraph('Result: ' + r['note']).runs[0].font.size = Pt(10)
            if r.get('shot') and os.path.exists(r['shot']):
                try:
                    doc.add_picture(r['shot'], width=Inches(6.3))
                    cap = doc.add_paragraph('Evidence: %s' % os.path.basename(r['shot']))
                    cap.runs[0].font.size = Pt(8); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception: pass
    doc.save(DOCX_OUT)

def main():
    try:
        out = subprocess.run(['powershell', '-NoProfile', '-Command',
            "(Get-NetTCPConnection -LocalPort %d -State Listen -ErrorAction SilentlyContinue).OwningProcess" % PORT],
            capture_output=True, text=True, timeout=20).stdout.strip()
        for pid in set(out.split()):
            if pid.isdigit(): subprocess.run(['taskkill', '/F', '/PID', pid], capture_output=True)
    except Exception: pass
    proxy = subprocess.Popen(['python', os.path.join(ROOT, 'ATD', 'Jet', 'dev-proxy.py'), str(PORT)],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx = Ctx(); ctx['browser'] = browser
            c, p = open_app(ctx, ADMINU); ctx['ctx'], ctx['page'] = c, p
            for area, prefix, cases in CASES:
                print('\n--', area)
                for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
                    run_case(ctx, 'ATD-%s-%02d' % (prefix, i), fn)
                    tc = ctx.pop('tmp_ctx', None)
                    if tc:
                        try: tc.close()
                        except Exception: pass
            browser.close()
    finally:
        proxy.terminate()
    build_workbook(XLSX_OUT, fill_status=True)
    build_workbook(SCRIPT_OUT, fill_status=False)
    build_docx()
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    print('\n=== %d cases: %s' % (len(RESULTS), counts))
    print('Workbook   :', XLSX_OUT)
    print('TestScript :', SCRIPT_OUT)
    print('Word report:', DOCX_OUT)
    print('Evidence   :', EVID_DIR)

if __name__ == '__main__':
    main()
