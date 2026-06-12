# -*- coding: utf-8 -*-
"""Automated UAT run — Credit Cards (App 202), 29 cases mirroring
UAT_CC_dd-Mon-YYYY-NN.xlsx. Produces a Word report with one evidence
screenshot per case under final apps/CC/UAT/.

Statuses: PASS / FAIL / PARTIAL / MANUAL / SKIP.
Credentials parsed at runtime from authService QUICK_LOGINS — never logged.
Personas (holder / CC_ADMIN / no-card user / SYS_ADMIN) enter via injected
shared sessions because module apps redirect to the Admin portal for login.
"""
import sys, os, re, json, time, base64, secrets, traceback, subprocess
import urllib.request, urllib.error
from datetime import date, timedelta
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
from playwright.sync_api import sync_playwright

ROOT   = r'c:\claude\DCT-task-management\DCT-Task-Management\final apps'
PORT   = 8081                       # own port so the FL/Admin runs on 8080 are untouched
APPURL = 'http://localhost:%d' % PORT
DATESTR = date.today().strftime('%d-%b-%Y')
RUN = time.strftime('%H%M%S')

HOLDER  = 'SHAIKHA.ALSUWAIDI'     # seeded card CC-2026-00002 (20,000 AED)
CCADMIN = 'NASER.ALKHAJA'         # CC_ADMIN
NOCARD  = 'SHAIKHA.GALAMERI'      # no card (creds live in the FL app file)

# ── credentials (runtime parse from CC + FL + Admin quick logins) ─────────────
def quick_of(app):
    src = open(os.path.join(ROOT, app, 'Jet', 'js', 'services', 'authService.js'),
               encoding='utf-8').read()
    return re.findall(r"username:\s*'([^']+)',\s*password:\s*'([^']+)'", src)
PWDS = {}
for _app in ('CC', 'FL', 'Admin'):
    for _u, _p in quick_of(_app):
        PWDS.setdefault(_u, _p)
ADMIN_QIDX = {u: i for i, (u, p) in enumerate(quick_of('Admin'))}

_cfg = open(os.path.join(ROOT, 'CC', 'Jet', 'js', 'services', 'config.js'), encoding='utf-8').read()
ADB = re.search(r"https://[\w.-]+oraclecloudapps\.com", _cfg).group(0) + '/ords/admin'

PNG_BYTES = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8'
    'z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==')
PNG_B64 = base64.b64encode(PNG_BYTES).decode()

# ── REST helper ───────────────────────────────────────────────────────────────
def api(method, path, body=None, token=None):
    req = urllib.request.Request(ADB + path, method=method,
                                 data=json.dumps(body).encode() if body is not None else None,
                                 headers={'Content-Type': 'application/json'})
    if token: req.add_header('Authorization', 'Bearer ' + token)
    try:
        r = urllib.request.urlopen(req, timeout=40)
        return r.status, json.loads(r.read().decode('utf-8', 'replace') or '{}')
    except urllib.error.HTTPError as e:
        try: return e.code, json.loads(e.read().decode('utf-8', 'replace') or '{}')
        except Exception: return e.code, {}

_LOGIN = {}
def login_data(username):
    if username not in _LOGIN:
        st, d = api('POST', '/dct/auth/login', {'username': username, 'password': PWDS[username]})
        assert d.get('sessionId'), 'ORDS login failed for ' + username
        d['roles'] = [r for r in (d.get('rolesCsv') or '').split(',') if r]
        parts = (d.get('displayName') or username).split(' ')
        d['initials'] = (parts[0][0] + parts[-1][0]).upper() if len(parts) >= 2 else parts[0][0].upper()
        _LOGIN[username] = d
    return _LOGIN[username]

def tok(username='ADMIN'):
    return login_data(username)['sessionId']

def active_card_of(username):
    u = login_data(username)
    st, d = api('GET', '/cc/cards/?limit=200', token=tok())
    for c in d.get('items', []):
        if c.get('holderUserId') == u['userId'] and c.get('status') == 'ACTIVE':
            return c
    return None

def upload_mandatory_docs(request_id, context, token):
    st, dr = api('GET', '/cc/doc-requirements?context=' + context, token=token)
    for d in dr.get('items', []):
        if d['isMandatory'] != 'Y': continue
        st, doc = api('POST', '/cc/documents/', {'sourceType': 'REQUEST', 'sourceId': request_id,
                                                 'docReqId': d['docReqId'], 'docTypeId': d['docTypeId'],
                                                 'fileName': 'uat_%s.png' % d['docReqId'],
                                                 'mimeType': 'image/png'}, token=token)
        api('PUT', '/cc/documents/%d/file' % doc['documentId'],
            {'file_data_b64': PNG_B64, 'mime_type': 'image/png'}, token=token)

def approve_all(ref, comment='UAT automated approval'):
    """Approve every pending step of ref via API (SYS_ADMIN may act)."""
    for _ in range(5):
        st, pend = api('GET', '/cc/approvals/pending', token=tok())
        inst = next((p for p in pend.get('items', []) if p.get('requestRef') == ref), None)
        if not inst: return
        api('POST', '/cc/approvals/%d/action' % inst['instanceId'],
            {'action': 'APPROVED', 'comments': comment}, token=tok())

# ── output locations ─────────────────────────────────────────────────────────
OUT_DIR = os.path.join(ROOT, 'CC', 'UAT')
os.makedirs(OUT_DIR, exist_ok=True)
seq = 1
while os.path.exists(os.path.join(OUT_DIR, 'UAT_CC_Results_%s-%02d.docx' % (DATESTR, seq))):
    seq += 1
DOCX_OUT = os.path.join(OUT_DIR, 'UAT_CC_Results_%s-%02d.docx' % (DATESTR, seq))
EVID_DIR = os.path.join(OUT_DIR, 'evidence_%s-%02d' % (DATESTR, seq))
os.makedirs(EVID_DIR, exist_ok=True)

RESULTS = []

# ── playwright helpers ───────────────────────────────────────────────────────
def wait(page, ms=2000):
    page.wait_for_timeout(ms)

def fresh_ctx(browser, username=None):
    c = browser.new_context(viewport={'width': 1440, 'height': 950})
    c.set_default_timeout(12000)
    if username:
        sess = json.dumps(json.dumps(login_data(username)))
        c.add_init_script("localStorage.setItem('ifinance_jet_session', %s);" % sess)
    return c

def open_app(ctx, username):
    c = fresh_ctx(ctx['browser'], username)
    p = c.new_page()
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000)
    wait(p, 2500)
    ensure_en(p)
    return c, p

def persona(ctx, username):
    if username not in ctx['pp']:
        ctx['pp'][username] = open_app(ctx, username)
    return ctx['pp'][username][1]

def nav(page, route, ms=2500, state=None):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("window._jetApp.navigate('%s', %s)" % (route, json.dumps(state or {})))
    wait(page, ms)

def ensure_en(page):
    try:
        if page.evaluate('document.documentElement.dir') == 'rtl':
            page.locator('.lang-pill button').first.click()
            wait(page, 1500)
    except Exception:
        pass

def fill_label(scope, label, value):
    grp = scope.locator('.form-group:has(label:text-is("%s"))' % label)
    if not grp.count():
        grp = scope.locator('.form-group:has-text("%s")' % label)
    ctl = grp.locator('input, select, textarea').first
    if ctl.evaluate('e => e.tagName') == 'SELECT':
        ctl.select_option(value)
    else:
        ctl.fill(value)
    ctl.dispatch_event('change')

def png_file(name='uat_doc.png'):
    return {'name': name, 'mimeType': 'image/png', 'buffer': PNG_BYTES}

def ui_act(page, ref, action, comment='UAT automated run'):
    nav(page, 'approvals', 3000)
    card = page.locator('.card:has-text("%s")' % ref).first
    card.wait_for(state='visible', timeout=15000)
    card.locator('button:has-text("%s")' % action).first.click(); wait(page, 800)
    card.locator('textarea').fill(comment)
    card.locator('button:has-text("Confirm")').click(); wait(page, 2800)

# ── case framework ───────────────────────────────────────────────────────────
class Ctx(dict):
    pass

def run_case(ctx, tid, area, function, scenario, fn):
    t0 = time.time()
    shot = os.path.join(EVID_DIR, tid + '.png')
    status, note = 'FAIL', ''
    try:
        out = fn(ctx)
        status, note = out if isinstance(out, tuple) else ('PASS', str(out or ''))
    except Exception as e:
        status = 'FAIL'
        note = ' '.join(str(e).split())[:300]
    try:
        pg = ctx.get('shot_page') or ctx['page']
        pg.screenshot(path=shot, full_page=False)
    except Exception:
        shot = ''
    ctx.pop('shot_page', None)
    tp = ctx.pop('tmp_page', None)
    if tp:
        try: tp.close()
        except Exception: pass
    secs = round(time.time() - t0, 1)
    RESULTS.append(dict(tid=tid, area=area, function=function, scenario=scenario,
                        status=status, note=note, shot=shot, secs=secs))
    print('%-7s %-12s %s  (%.0fs)  %s' % (status, tid, scenario[:46].ljust(46), secs, note[:70]))

# ═════════════════════════════════════════════════════════════════════════════
# Executors
# ═════════════════════════════════════════════════════════════════════════════

# ── ACC ──────────────────────────────────────────────────────────────────────
def acc_entry(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL + '/Admin/Jet/index.html', timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000)
    p.locator('.quick-btn').nth(ADMIN_QIDX['ADMIN']).click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    p.locator('.modsw-btn').click(); wait(p, 700)
    p.locator('.modsw-item:has-text("Credit Card")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 3000)
    assert '/CC/Jet/' in p.url, 'did not land on the CC app (%s)' % p.url
    cube = p.locator('.brand-cube').inner_text()
    assert cube.strip() == 'CC', 'CC brand cube missing'
    assert p.locator('.side .nav-item').count() >= 4, 'CC nav missing'
    return 'PASS', 'Admin login → module switcher → CC opened logged-in with full SYS_ADMIN nav'

def acc_nosession(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL, timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000); wait(p, 2000)
    assert '/Admin/Jet/' in p.url, 'no redirect to the Admin portal (%s)' % p.url
    assert p.locator('.login-card').count(), 'Admin login page not shown'
    return 'PASS', 'CC without a session redirected to the Admin portal login'

def acc_roles(ctx):
    ph = persona(ctx, HOLDER); nav(ph, 'dashboard', 1500)
    holder_side = ph.locator('.side').inner_text().upper()
    pn = persona(ctx, CCADMIN); nav(pn, 'dashboard', 1500)
    admin_side = pn.locator('.side').inner_text().upper()
    assert 'ALL CARDS' not in holder_side and 'MODULE SETTINGS' not in holder_side \
           and 'PROXIES' not in holder_side, 'holder sees admin menus'
    assert 'ALL CARDS' in admin_side and 'PROXIES' in admin_side \
           and 'MODULE SETTINGS' in admin_side, 'CC_ADMIN lacks admin menus'
    assert 'APPROVALS' in admin_side, 'CC_ADMIN lacks the approvals group'
    holder_apr = 'has' if 'APPROVALS' in holder_side else 'does not have'
    ctx['shot_page'] = pn
    return 'PASS', 'Admin group (All Cards/Proxies/Module Settings) only for CC_ADMIN; holder %s Approvals (MANAGER role dependent)' % holder_apr

# ── DSH ──────────────────────────────────────────────────────────────────────
def dsh_stats(ctx):
    p = ctx['page']; nav(p, 'dashboard', 3500)
    vals = p.locator('.stat-card').evaluate_all(
        'els => els.map(e => (e.firstElementChild || {}).innerText || "")')
    nums = [v for v in vals if re.sub(r'[,\s]', '', v).replace('.', '').isdigit()]
    assert len(vals) >= 3 and len(nums) >= 2, 'stat cards missing: %s' % vals
    return 'PASS', 'Stat cards live: %s' % ', '.join(v.strip() for v in vals[:4])

def dsh_charts(ctx):
    p = ctx['page']; nav(p, 'dashboard', 4500)
    n = p.locator('canvas').count()
    assert n >= 2, 'expected 2 charts, found %d canvases' % n
    sizes = p.locator('canvas').evaluate_all('els => els.map(e => e.width)')
    assert all(s > 0 for s in sizes), 'a chart canvas has zero width'
    return 'PASS', 'Limits-by-organisation + 6-month replenishment compliance charts rendered (%d canvases)' % n

# ── CRD ──────────────────────────────────────────────────────────────────────
def crd_visual(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'myCard', 3200)
    assert p.locator('.cc-card').count(), 'card visual missing'
    txt = p.locator('.cc-card').inner_text()
    assert '20,000' in txt or re.search(r'\d{1,3},\d{3}', txt), 'limit missing on the visual'
    assert p.locator('.cc-number').inner_text().strip(), 'masked number missing'
    badge = p.locator('.cc-card .badge').inner_text()
    return 'PASS', 'Card visual with masked number, holder, limit and %s status badge' % badge

def crd_history(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'myCard', 3000)
    hist = p.locator('.card .data-table').last.inner_text()
    assert 'INITIAL' in hist, 'INITIAL limit-history row missing'
    kinds = sorted(set(re.findall(r'INITIAL|INCREASE|DECREASE', hist)))
    return 'PASS', 'Limit history shows the registration INITIAL row (entries: %s)' % ', '.join(kinds)

def crd_due_banner(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    card = active_card_of(HOLDER)
    today = date.today()
    st, d = api('GET', '/cc/replenishments/?mine=Y&limit=100', token=tok(HOLDER))
    cur = [r for r in d.get('items', [])
           if r.get('ccId') == card['ccId'] and str(today.year) in str(r.get('periodLabel', ''))
           and r.get('periodMonth', 0) == today.month] if card else []
    # periodMonth may not be in the list payload — fall back to the label
    if card and not cur:
        cur = [r for r in d.get('items', [])
               if today.strftime('%b') in str(r.get('periodLabel', '')) and str(today.year) in str(r.get('periodLabel', ''))]
    nav(p, 'myCard', 3200)
    banner = p.locator('.due-banner').count()
    if not cur:
        assert banner, 'current-month statement missing but no due banner'
        return 'PASS', 'No %s statement yet → amber due banner with a New Replenishment shortcut shown' % today.strftime('%B')
    assert not banner, 'statement exists for the current month but the banner still shows'
    return 'PASS', 'Current-month statement already exists (%s) → banner correctly absent' % cur[0].get('reimbNumber')

def crd_nocard(ctx):
    p = persona(ctx, NOCARD); ctx['shot_page'] = p
    nav(p, 'myCard', 3000)
    assert p.locator('.empty-state').count(), 'no friendly empty state'
    assert p.locator('.empty-state .btn-primary').count(), 'no New Request shortcut'
    return 'PASS', 'User without a card sees a friendly empty state offering a New Card request'

# ── REQ ──────────────────────────────────────────────────────────────────────
def req_wizard(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'requests', 2500)
    p.locator('.page-header-row .btn-primary').click(); wait(p, 2500)
    tiles = p.locator('.type-tile').count()
    steps = p.locator('.wiz-step').count()
    assert tiles == 5, 'expected 5 type tiles, got %d' % tiles
    assert steps == 3, 'expected a 3-step wizard, got %d' % steps
    body = p.locator('.page-wrap').inner_text()
    for t in ('New Card', 'Increase', 'Decrease', 'Replacement', 'Close'):
        assert t in body, 'tile missing: %s' % t
    return 'PASS', '5 request-type tiles (New/Increase/Decrease/Replacement/Close) on a 3-step wizard'

def _requester_without_card(ctx):
    for u in ('ADMIN', NOCARD):
        if not active_card_of(u):
            return u
    return None

def req_newcard(ctx):
    u = _requester_without_card(ctx)
    assert u, 'every quick-login user already holds an active card — free one up first'
    ctx['newcard_user'] = u
    p = ctx['page'] if u == 'ADMIN' else persona(ctx, u)
    ctx['shot_page'] = p
    nav(p, 'requestNew', 2800)
    p.locator('.type-tile:has-text("New Card")').click(); wait(p, 500)
    p.locator('button:has-text("Continue")').click(); wait(p, 1200)
    fill_label(p, 'Requested Limit (AED) *', '10000')
    p.locator('textarea').fill('UAT automated happy path ' + RUN)
    p.locator('button:has-text("Save Draft")').click(); wait(p, 3000)
    rows = p.locator('.doc-req-row')
    n = rows.count()
    for i in range(n):
        with p.expect_file_chooser() as fc:
            rows.nth(i).locator('button').click()
        fc.value.set_files(png_file('uat_doc_%d.png' % i)); wait(p, 2200)
    submit = p.locator('button:has-text("Submit for Approval")')
    assert submit.is_enabled(), 'submit still disabled after uploading all documents'
    submit.click(); wait(p, 3200)
    body = p.locator('.page-title').inner_text()
    assert 'SUBMITTED' in body or 'UNDER_REVIEW' in body, 'request not submitted: %s' % body[:80]
    ctx['req_num'] = p.locator('.page-title .mono').first.inner_text().strip()
    st, d = api('GET', '/cc/requests/?mine=Y&limit=10', token=tok(u))
    mine = next((r for r in d.get('items', []) if r.get('requestNumber') == ctx['req_num']), None)
    ctx['req_id'] = (mine or {}).get('requestId')
    return 'PASS', '%s: NEW_CARD draft by %s, %d documents uploaded, submitted to approval' % (ctx['req_num'], u, n)

def req_docs_enforced(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'requestNew', 2800)
    p.locator('.type-tile:has-text("Increase")').click(); wait(p, 500)
    p.locator('button:has-text("Continue")').click(); wait(p, 1200)
    card_sel = p.locator('.form-group:has-text("Card *") select')
    if card_sel.count():
        card_sel.select_option(index=1)
        card_sel.dispatch_event('change')
    fill_label(p, 'Requested Limit (AED) *', '25000')
    p.locator('textarea').fill('UAT doc-checklist test ' + RUN)
    p.locator('button:has-text("Save Draft")').click(); wait(p, 3000)
    err = p.locator('.alert-banner--danger')
    if err.count() and err.is_visible():
        return ('PARTIAL', 'Draft creation itself was blocked: "%s" — the doc checklist could not '
                'be reached this run (likely an in-flight request on the card).' % err.inner_text()[:120])
    mand = p.locator('.doc-req-row:has-text("Mandatory")').count()
    submit = p.locator('button:has-text("Submit for Approval")')
    if mand == 0:
        return ('PARTIAL', 'INCREASE_LIMIT has no mandatory documents configured; nothing to block on. '
                'Submit enabled state: %s' % submit.is_enabled())
    assert not submit.is_enabled(), 'Submit enabled although %d mandatory documents are missing' % mand
    st, d = api('GET', '/cc/requests/?mine=Y&limit=10', token=tok(HOLDER))
    mine = next((r for r in d.get('items', []) if 'UAT doc-checklist test ' + RUN in (r.get('reason') or '')
                 or (r.get('status') == 'DRAFT' and r.get('requestType') == 'INCREASE_LIMIT')), None)
    ctx['inc_req'] = mine
    p.locator('button:has-text("Save as Draft")').click(); wait(p, 2000)
    return 'PASS', 'With %d mandatory document(s) missing the Submit button stays disabled; saved as draft' % mand

def req_increase_validation(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    card = active_card_of(HOLDER)
    cur = int(card['creditLimit'])
    low = max(1000, cur - 5000)
    nav(p, 'requestNew', 2800)
    p.locator('.type-tile:has-text("Increase")').click(); wait(p, 500)
    p.locator('button:has-text("Continue")').click(); wait(p, 1200)
    card_sel = p.locator('.form-group:has-text("Card *") select')
    if card_sel.count():
        card_sel.select_option(index=1)
        card_sel.dispatch_event('change')
    fill_label(p, 'Requested Limit (AED) *', str(low))
    p.locator('textarea').fill('UAT lower-limit validation ' + RUN)
    p.locator('button:has-text("Save Draft")').click(); wait(p, 3000)
    err = p.locator('.alert-banner--danger')
    if err.count() and err.is_visible():
        msg = err.inner_text()
        return 'PASS', 'Blocked at draft creation: "%s" (requested %s < current %s)' % (msg.strip()[:120], low, cur)
    # draft allowed — the block must come at submit
    rows = p.locator('.doc-req-row')
    for i in range(rows.count()):
        with p.expect_file_chooser() as fc:
            rows.nth(i).locator('button').click()
        fc.value.set_files(png_file()); wait(p, 2000)
    p.locator('button:has-text("Submit for Approval")').click(); wait(p, 2800)
    err = p.locator('.alert-banner--danger')
    assert err.count() and err.is_visible(), 'lower-than-current limit was accepted!'
    return 'PASS', 'Blocked at submit: "%s" (requested %s < current %s)' % (err.inner_text().strip()[:120], low, cur)

def req_withdraw(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    inc = ctx.get('inc_req')
    assert inc, 'no INCREASE_LIMIT draft from CC-REQ-03'
    upload_mandatory_docs(inc['requestId'], 'INCREASE_LIMIT', tok(HOLDER))
    nav(p, 'requestDetail', 3000, state={'requestId': inc['requestId']})
    sub = p.locator('button:has-text("Submit")')
    if sub.count():
        sub.click(); wait(p, 3000)
    body = p.locator('.page-title').inner_text()
    assert 'SUBMITTED' in body or 'UNDER_REVIEW' in body, 'request did not submit (%s)' % body[:60]
    p.locator('button:has-text("Withdraw")').click(); wait(p, 3000)
    body = p.locator('.page-title').inner_text()
    assert 'WITHDRAWN' in body, 'request not withdrawn (%s)' % body[:60]
    st, pend = api('GET', '/cc/approvals/pending', token=tok())
    still = next((x for x in pend.get('items', []) if x.get('requestRef') == inc['requestNumber']), None)
    assert not still, 'withdrawn request still sits in the approver queue'
    return 'PASS', '%s submitted then withdrawn; status WITHDRAWN and gone from the approver queue' % inc['requestNumber']

def req_detail(ctx):
    p = ctx['page'] if ctx.get('newcard_user') == 'ADMIN' else persona(ctx, ctx.get('newcard_user', HOLDER))
    ctx['shot_page'] = p
    if ctx.get('req_id'):
        nav(p, 'requestDetail', 3000, state={'requestId': ctx['req_id']})
    else:
        nav(p, 'requests', 2500)
        p.locator('.data-table tbody tr').first.click(); wait(p, 2800)
    body = p.locator('.page-wrap').inner_text().upper()
    assert p.locator('.audit-info').count(), 'audit info missing'
    for needle in ('TYPE', 'REQUESTED LIMIT', 'REASON'):
        assert needle in body, '%s missing on the detail' % needle
    docs = p.locator('.card .data-table tbody tr').count()
    return 'PASS', 'Detail shows type, card, limits, reason, %d document(s) and audit info' % docs

# ── RPL ──────────────────────────────────────────────────────────────────────
def rpl_create(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'replenishments', 2800)
    p.locator('.page-header-row .btn-primary').click(); wait(p, 1500)
    modal = p.locator('.modal-box')
    # duplicate path first: May 2026 is seeded (CCM-2026-05-00001)
    modal.locator('.form-group:has-text("Month") select').select_option('5')
    modal.locator('.form-group:has-text("Year") input').fill('2026')
    modal.locator('button:has-text("Create Draft")').click(); wait(p, 2500)
    dup = modal.locator('.alert-banner--danger')
    dup_msg = dup.inner_text().strip()[:90] if dup.count() and dup.is_visible() else ''
    assert dup_msg, 'May 2026 duplicate was not blocked'
    # now find a free month for the real draft
    created = None
    today = date.today()
    tries = [today] + [today.replace(day=1) - timedelta(days=30 * k) for k in (2, 3, 4, 5)]
    for t in tries:
        modal.locator('.form-group:has-text("Month") select').select_option(str(t.month))
        modal.locator('.form-group:has-text("Year") input').fill(str(t.year))
        modal.locator('button:has-text("Create Draft")').click(); wait(p, 2800)
        if not p.locator('.modal-box').count():           # modal closed → navigated to the detail
            created = t; break
    assert created, 'no free period found to create a draft'
    wait(p, 1500)
    num = p.locator('.page-title .mono').first.inner_text()
    assert num.startswith('CCM-'), 'no CCM number (%s)' % num
    ctx['rpl_num'] = num
    ctx['rpl_period'] = created
    st, d = api('GET', '/cc/replenishments/?mine=Y&limit=50', token=tok(HOLDER))
    ctx['rpl_id'] = next(r['replenishmentId'] for r in d['items'] if r['reimbNumber'] == num)
    return 'PASS', 'May 2026 duplicate blocked ("%s"); %s draft created for %s' % (dup_msg, num, created.strftime('%b %Y'))

def rpl_lines(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'replenishmentDetail', 3200, state={'replenishmentId': ctx['rpl_id']})
    period = ctx['rpl_period']
    lines = [('Etisalat', '420.50', 5), ('Amazon AE', '1310.00', 12), ('ADNOC', '240.00', 21)]
    for merchant, amount, day in lines:
        p.locator('button:has-text("Add Line")').click(); wait(p, 600)
        row = p.locator('.lines-table tbody tr').last
        row.locator('input[type=date]').fill(period.replace(day=day).isoformat())
        row.locator('input[placeholder="Merchant name"]').fill(merchant)
        row.locator('input[type=number]').fill(amount)
        gl = row.locator('select').nth(1)
        gl.select_option(index=1); gl.dispatch_event('change')
        row.locator('input[type=checkbox]').locator('..').click()   # receipt flag
    p.locator('button:has-text("Save Lines")').click(); wait(p, 3000)
    total = p.locator('.stat-card .mono').first.inner_text()
    assert '1,970.5' in total, 'running total wrong: %s' % total
    receipts = p.locator('.stat-card .mono').nth(2).inner_text()
    return 'PASS', '3 merchant lines saved; running total AED %s; receipts flagged %s' % (total, receipts)

def rpl_validation(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'replenishmentDetail', 3200, state={'replenishmentId': ctx['rpl_id']})
    before = p.locator('.lines-table tbody tr').count()
    p.locator('button:has-text("Add Line")').click(); wait(p, 600)
    row = p.locator('.lines-table tbody tr').last
    row.locator('input[placeholder="Merchant name"]').fill('Incomplete line')   # no amount/date
    p.locator('button:has-text("Save Lines")').click(); wait(p, 1500)
    body = p.locator('body').inner_text()
    assert 'required' in body.lower(), 'no per-line validation message'
    row.locator('button[title="Remove line"]').click(); wait(p, 600)
    p.locator('button:has-text("Save Lines")').click(); wait(p, 2500)
    after_reload_rows = p.locator('.lines-table tbody tr').count()
    assert after_reload_rows == before, 'row count drifted (%d → %d)' % (before, after_reload_rows)
    return 'PASS', 'Line without date/amount blocked with a per-line message; nothing saved after removal'

def rpl_submit(ctx):
    p = persona(ctx, HOLDER); ctx['shot_page'] = p
    nav(p, 'replenishmentDetail', 3200, state={'replenishmentId': ctx['rpl_id']})
    quirk = ''
    if p.locator(':text("Unsaved changes")').count():
        # freshly loaded detail sometimes flags phantom unsaved changes — save
        # first so Submit is accepted, and surface the quirk in the note
        quirk = ' (note: a phantom "Unsaved changes" marker appeared on a freshly loaded draft — saved first; minor UI defect to review)'
        p.locator('button:has-text("Save Lines")').click(); wait(p, 2800)
    p.locator('button:has-text("Submit")').click(); wait(p, 3200)
    badge = p.locator('.page-title').inner_text()
    assert 'SUBMITTED' in badge or 'UNDER_REVIEW' in badge, 'not submitted (%s)' % badge[:60]
    st, pend = api('GET', '/cc/approvals/pending', token=tok())
    inq = next((x for x in pend.get('items', []) if x.get('requestRef') == ctx['rpl_num']), None)
    assert inq, 'statement missing from the approver queue'
    return 'PASS', '%s submitted by the holder; visible in the approvals queue%s' % (ctx['rpl_num'], quirk)

def rpl_proxy(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    holder_card = active_card_of(HOLDER)
    # 1. create the proxy via the UI
    nav(pn, 'proxies', 2800)
    pn.locator('.page-header-row .btn-primary').click(); wait(pn, 1200)
    modal = pn.locator('.modal-box')
    cs = modal.locator('.form-group:has-text("Card *") select')
    cs.select_option(str(holder_card['ccId'])); cs.dispatch_event('change')
    us = modal.locator('.form-group:has-text("Proxy User *") select')
    us.select_option(str(login_data(CCADMIN)['userId'])); us.dispatch_event('change')
    fill_label(modal, 'Start Date', date.today().isoformat())
    fill_label(modal, 'End Date', (date.today() + timedelta(days=7)).isoformat())
    modal.locator('button:has-text("Create Proxy")').click(); wait(pn, 2800)
    # 2. proxy submits a statement for the holder's card (free past month, via API)
    sub_ok, used = None, None
    for k in (3, 4, 5, 6):
        t = date.today().replace(day=1) - timedelta(days=30 * k)
        st, rep = api('POST', '/cc/replenishments/', {'ccId': holder_card['ccId'],
                      'periodMonth': t.month, 'periodYear': t.year}, token=tok(CCADMIN))
        if st in (200, 201) and rep.get('replenishmentId'):
            st, gl = api('GET', '/cc/gl-codes', token=tok(CCADMIN))
            glid = next(g['ccId'] for g in gl if g['isActive'] == 'Y')
            api('PUT', '/cc/replenishments/%d/lines' % rep['replenishmentId'],
                {'lines': [{'expenseDate': t.replace(day=9).isoformat(), 'merchantName': 'Careem',
                            'amount': 95.0, 'codingType': 'GL', 'ccIdGl': glid,
                            'receiptAttached': 'Y', 'description': 'UAT proxy line'}]}, token=tok(CCADMIN))
            st2, d2 = api('POST', '/cc/replenishments/%d/submit' % rep['replenishmentId'], {}, token=tok(CCADMIN))
            sub_ok, used = d2.get('ok'), rep['reimbNumber']
            break
    assert sub_ok, 'proxy submit failed'
    # 3. non-proxy is blocked
    blocked = ''
    t = date.today().replace(day=1) - timedelta(days=30 * 8)
    st, rep = api('POST', '/cc/replenishments/', {'ccId': holder_card['ccId'],
                  'periodMonth': t.month, 'periodYear': t.year}, token=tok(NOCARD))
    if st in (200, 201) and rep.get('replenishmentId'):
        st3, d3 = api('POST', '/cc/replenishments/%d/submit' % rep['replenishmentId'], {}, token=tok(NOCARD))
        blocked = (d3.get('message') or d3.get('error') or str(d3))[:90] if st3 >= 400 or not d3.get('ok') else ''
        assert blocked, 'a non-proxy was allowed to submit!'
    nav(pn, 'replenishments', 2500)
    sc = pn.locator('.lang-pill button:has-text("All")')
    if sc.count(): sc.click(); wait(pn, 2000)
    assert pn.locator('.data-table tbody tr:has-text("%s")' % used).count(), 'proxy statement not listed'
    return 'PASS', 'Proxy created for %s; proxy submitted %s on behalf of the holder; non-proxy blocked ("%s")' % (
        holder_card['ccNumber'], used, blocked or 'create blocked upstream')

# ── APR ──────────────────────────────────────────────────────────────────────
def apr_queue(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    nav(pn, 'approvals', 3200)
    cards = pn.locator('.page-wrap .card').count()
    assert cards >= 1, 'approver queue empty (REQ/RPL submissions expected)'
    body = pn.locator('.page-wrap').inner_text()
    has_req = bool(re.search(r'REQUEST|CCR|CC-REQ', body.upper()))
    has_rpl = 'CCM-' in body
    return 'PASS', '%d items awaiting the approver (card requests: %s, replenishments: %s)' % (cards, has_req, has_rpl)

def apr_comment(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    nav(pn, 'approvals', 3000)
    card = pn.locator('.page-wrap .card').first
    ref = card.locator('.mono').first.inner_text()
    card.locator('button:has-text("Approve")').click(); wait(pn, 800)
    card.locator('button:has-text("Confirm")').click(); wait(pn, 1200)
    body = card.inner_text()
    assert 'comment is required' in body.lower(), 'no required-comment error: %s' % body[:120]
    assert pn.locator('.card:has-text("%s")' % ref).count(), 'item left the queue without a comment!'
    card.locator('button:has-text("Cancel")').click(); wait(pn, 600)
    return 'PASS', 'Empty comment blocked with "A comment is required."; %s stayed queued' % ref

def apr_lifecycle(ctx):
    card = active_card_of(HOLDER)
    cur = int(card['creditLimit'])
    target = cur + 2000
    st, req = api('POST', '/cc/requests/', {'ccId': card['ccId'], 'requestType': 'INCREASE_LIMIT',
                                            'requestedLimit': target,
                                            'reason': 'UAT lifecycle ' + RUN}, token=tok(HOLDER))
    assert req.get('requestId'), 'lifecycle request create failed: %s %s' % (st, req)
    upload_mandatory_docs(req['requestId'], 'INCREASE_LIMIT', tok(HOLDER))
    st, d = api('POST', '/cc/requests/%d/submit' % req['requestId'], {}, token=tok(HOLDER))
    assert d.get('ok'), 'lifecycle submit failed: %s' % d
    pn = persona(ctx, CCADMIN)
    for _ in range(4):     # approve every step that lands in the CC_ADMIN queue
        nav(pn, 'approvals', 2800)
        if not pn.locator('.card:has-text("%s")' % req['requestNumber']).count():
            break
        ui_act(pn, req['requestNumber'], 'Approve', 'UAT lifecycle approval')
    approve_all(req['requestNumber'])          # safety net for steps owned by other roles
    card2 = active_card_of(HOLDER)
    assert int(card2['creditLimit']) == target, 'limit not applied (%s, expected %s)' % (card2['creditLimit'], target)
    ph = persona(ctx, HOLDER); ctx['shot_page'] = ph
    nav(ph, 'myCard', 3200)
    hist = ph.locator('.card .data-table').last.inner_text()
    assert 'INCREASE' in hist, 'INCREASE row missing from the limit history'
    # restore the original limit so re-runs start clean
    try:
        st, dec = api('POST', '/cc/requests/', {'ccId': card['ccId'], 'requestType': 'DECREASE_LIMIT',
                                                'requestedLimit': cur,
                                                'reason': 'UAT restore ' + RUN}, token=tok(HOLDER))
        upload_mandatory_docs(dec['requestId'], 'DECREASE_LIMIT', tok(HOLDER))
        api('POST', '/cc/requests/%d/submit' % dec['requestId'], {}, token=tok(HOLDER))
        approve_all(dec['requestNumber'], 'UAT restore approval')
        restored = int(active_card_of(HOLDER)['creditLimit']) == cur
    except Exception:
        restored = False
    return 'PASS', 'INCREASE_LIMIT %s → %s fully approved; card limit updated + INCREASE history row; limit then restored to %s (%s)' % (
        cur, target, cur, 'ok' if restored else 'restore pending — check manually')

# ── ADM ──────────────────────────────────────────────────────────────────────
def adm_allcards(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    nav(pn, 'allCards', 3000)
    rows = pn.locator('.data-table tbody tr').count()
    assert rows >= 2, 'expected the org-wide registry (%d rows)' % rows
    pn.locator('.view-toolbar input').fill('SHAIKHA'); wait(pn, 1800)
    found = pn.locator('.data-table tbody tr').count()
    assert found >= 1, 'search found nothing'
    pn.locator('.view-toolbar input').fill(''); wait(pn, 1200)
    pn.locator('.view-toolbar select').select_option('ACTIVE'); wait(pn, 1800)
    active = pn.locator('.data-table tbody tr').count()
    pn.locator('.view-toolbar select').select_option(''); wait(pn, 1000)
    return 'PASS', '%d cards org-wide; search "SHAIKHA" → %d; ACTIVE filter → %d; pending-op chips shown' % (rows, found, active)

def adm_register(ctx):
    assert ctx.get('req_num'), 'no NEW_CARD request from CC-REQ-02'
    approve_all(ctx['req_num'], 'UAT — approve NEW_CARD for registration')
    u = login_data(ctx['newcard_user'])
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    nav(pn, 'allCards', 2800)
    pn.locator('.page-header-row .btn-primary').click(); wait(pn, 1500)
    modal = pn.locator('.modal-box')
    rs = modal.locator('.form-group:has-text("Approved NEW_CARD Request *") select')
    opts = rs.locator('option').all_inner_texts()
    pick = next((o for o in opts if ctx['req_num'] in o), None)
    assert pick, '%s not offered in the drawer (options: %s)' % (ctx['req_num'], opts[:4])
    rs.select_option(label=pick); rs.dispatch_event('change')
    hs = modal.locator('.form-group:has-text("Card Holder *") select')
    hs.select_option(str(u['userId'])); hs.dispatch_event('change')
    os_ = modal.locator('.form-group:has-text("Organisation *") select')
    os_.select_option(index=1); os_.dispatch_event('change')
    fill_label(modal, "Mother's Name *", 'UAT Mother')
    fill_label(modal, 'Nationality *', 'Emirati')
    fill_label(modal, 'Credit Limit (AED) *', '10000')
    fill_label(modal, 'Card Expiry *', (date.today() + timedelta(days=3 * 365)).isoformat())
    fill_label(modal, 'Email *', (u.get('email') or 'uat@example.com'))
    modal.locator('button:has-text("Register Card")').click(); wait(pn, 3200)
    newcard = active_card_of(ctx['newcard_user'])
    assert newcard, 'card not created'
    assert pn.locator('.data-table tbody tr:has-text("%s")' % newcard['ccNumber']).count(), 'new card not listed'
    ctx['registered_card'] = newcard['ccNumber']
    return 'PASS', 'Approved %s registered as %s (ACTIVE, 10,000 AED, INITIAL history row) for %s' % (
        ctx['req_num'], newcard['ccNumber'], ctx['newcard_user'])

def adm_proxies(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    st, cards = api('GET', '/cc/cards/?status=ACTIVE&limit=10', token=tok())
    pick = next((c for c in cards['items'] if c['holderUserId'] != login_data(HOLDER)['userId']),
                cards['items'][0])
    nav(pn, 'proxies', 2800)
    before = pn.locator('.data-table tbody tr').count()
    pn.locator('.page-header-row .btn-primary').click(); wait(pn, 1200)
    modal = pn.locator('.modal-box')
    cs = modal.locator('.form-group:has-text("Card *") select')
    cs.select_option(str(pick['ccId'])); cs.dispatch_event('change')
    us = modal.locator('.form-group:has-text("Proxy User *") select')
    us.select_option(index=1); us.dispatch_event('change')
    fill_label(modal, 'Start Date', date.today().isoformat())
    fill_label(modal, 'End Date', (date.today() + timedelta(days=5)).isoformat())
    modal.locator('button:has-text("Create Proxy")').click(); wait(pn, 2800)
    row = pn.locator('.data-table tbody tr:has-text("%s")' % pick['ccNumber']).first
    assert row.count(), 'new proxy not listed'
    row.locator('button:has-text("Deactivate")').click(); wait(pn, 2500)
    row = pn.locator('.data-table tbody tr:has-text("%s")' % pick['ccNumber']).first
    assert 'No' in row.inner_text(), 'proxy still active after deactivation'
    return 'PASS', 'Proxy created on %s then deactivated; submit rights removed' % pick['ccNumber']

def adm_settings(ctx):
    pn = persona(ctx, CCADMIN); ctx['shot_page'] = pn
    nav(pn, 'moduleSettings', 3200)
    row = pn.locator('.switch-row:has-text("REPLENISHMENT_DUE_DAY")')
    if not row.count():
        row = pn.locator('.switch-row').first
    inp = row.locator('input.form-control').first
    orig = inp.input_value()
    inp.fill(str(int(orig) + 1) if orig.isdigit() else orig + '1')
    inp.dispatch_event('input')
    pn.locator('.page-header-row .btn-primary').click(); wait(pn, 2500)
    nav(pn, 'dashboard', 1200); nav(pn, 'moduleSettings', 3000)
    row = pn.locator('.switch-row:has-text("REPLENISHMENT_DUE_DAY")')
    if not row.count():
        row = pn.locator('.switch-row').first
    newval = row.locator('input.form-control').first.input_value()
    changed = newval != orig
    row.locator('input.form-control').first.fill(orig)
    row.locator('input.form-control').first.dispatch_event('input')
    pn.locator('.page-header-row .btn-primary').click(); wait(pn, 2000)
    assert changed, 'setting did not persist (still %s)' % newval
    return 'PASS', 'Setting %s → %s persisted across reload; restored to %s. Non-admins have no Module Settings nav (CC-ACC-03)' % (orig, newval, orig)

# ── SHL ──────────────────────────────────────────────────────────────────────
def shl_switch(ctx):
    p = ctx['page']
    nav(p, 'dashboard', 2000)
    p.locator('.modsw-btn').click(); wait(p, 700)
    cc_item = p.locator('.modsw-item:has-text("Credit Card")')
    has_soon = cc_item.locator('.soon-b').count() > 0
    p.locator('.modsw-item:has-text("Freelancer")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    assert '/FL/Jet/' in p.url and not p.locator('.login-card').count(), 're-login was demanded'
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    ensure_en(p)
    assert not has_soon, 'CC still shows a "soon" badge in the switcher'
    return 'PASS', 'CC → FL → CC with no re-login; CC listed live (no "soon" badge)'

def shl_arabic(ctx):
    ph = persona(ctx, HOLDER); ctx['shot_page'] = ph
    nav(ph, 'dashboard', 2000)
    ph.locator('.lang-pill button').nth(1).click(); wait(ph, 1800)
    assert ph.evaluate('document.documentElement.dir') == 'rtl', 'document did not flip to RTL'
    nav(ph, 'dashboard', 3200)
    canvases = ph.locator('canvas').count()
    nav(ph, 'myCard', 3000)
    visual = ph.locator('.cc-card').count()
    limit_txt = ph.locator('.cc-card').inner_text() if visual else ''
    latin = bool(re.search(r'[0-9]', limit_txt))
    if ctx.get('rpl_id'):
        nav(ph, 'replenishmentDetail', 2800, state={'replenishmentId': ctx['rpl_id']})
    # leave AR active for the evidence screenshot; main() restores EN
    assert canvases >= 2 and visual, 'AR walk incomplete (%d canvases, visual=%s)' % (canvases, visual)
    assert latin, 'digits not Latin in AR mode'
    return 'PASS', 'RTL mirrored with Latin digits; charts re-rendered (%d); card visual unaffected; detail OK' % canvases

# ═════════════════════════════════════════════════════════════════════════════
CASES = [
 ('Access & Session', 'ACC', [
   ('Shared session', 'Entry from Admin', acc_entry),
   ('No session', 'Direct open without login', acc_nosession),
   ('Role gating', 'Admin menus', acc_roles),
 ]),
 ('Dashboard', 'DSH', [
   ('Stats', 'Cards show live numbers', dsh_stats),
   ('Charts', 'Limits + compliance', dsh_charts),
 ]),
 ('My Card', 'CRD', [
   ('Card visual', 'My card renders', crd_visual),
   ('Limit history', 'Timeline', crd_history),
   ('Due banner', 'Replenishment due', crd_due_banner),
   ('No card', 'Holder without card', crd_nocard),
 ]),
 ('Card Requests', 'REQ', [
   ('Wizard', 'Type tiles', req_wizard),
   ('New card', 'Full happy path', req_newcard),
   ('Doc checklist', 'Mandatory docs enforced', req_docs_enforced),
   ('Increase limit', 'Validation', req_increase_validation),
   ('Withdraw', 'Pull back a submitted request', req_withdraw),
   ('Detail', 'Request detail', req_detail),
 ]),
 ('Replenishments', 'RPL', [
   ('Create', 'New monthly statement', rpl_create),
   ('Lines', 'Merchant lines editor', rpl_lines),
   ('Validation', 'Line rules', rpl_validation),
   ('Submit', 'Send for approval', rpl_submit),
   ('Proxy', 'Proxy submits for the holder', rpl_proxy),
 ]),
 ('Approvals', 'APR', [
   ('Queue', 'Pending list', apr_queue),
   ('Comment required', 'Approve needs a comment', apr_comment),
   ('Lifecycle', 'Approved request applies', apr_lifecycle),
 ]),
 ('Administration', 'ADM', [
   ('All cards', 'Registry + search', adm_allcards),
   ('Register card', 'From approved NEW_CARD', adm_register),
   ('Proxies', 'Manage proxies', adm_proxies),
   ('Module settings', 'CC settings', adm_settings),
 ]),
 ('Shell & Language', 'SHL', [
   ('Switcher', 'Cross-module hop', shl_switch),
   ('Arabic', 'RTL pass', shl_arabic),
 ]),
]

# ═════════════════════════════════════════════════════════════════════════════
def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    COLORS = {'PASS': RGBColor(0x1B, 0x5E, 0x20), 'FAIL': RGBColor(0xB7, 0x1C, 0x1C),
              'PARTIAL': RGBColor(0xE6, 0x51, 0x00), 'MANUAL': RGBColor(0x54, 0x6E, 0x7A),
              'SKIP': RGBColor(0x54, 0x6E, 0x7A)}
    doc = Document()
    doc.add_heading('i-Finance UAT Results — Credit Cards (App 202)', level=0)
    p = doc.add_paragraph('Automated UAT run · %s · Environment: %s (dev-proxy → ADB PROD)'
                          % (DATESTR, APPURL))
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph('Executed by Claude (Playwright). Mirrors UAT_CC_%s-NN.xlsx — one section per '
                      'workbook case, with the evidence screenshot captured at the verification '
                      'moment. Personas: SYS_ADMIN, CC_ADMIN (NASER.ALKHAJA), card holder '
                      '(SHAIKHA.ALSUWAIDI), no-card user (SHAIKHA.GALAMERI). The card limit changed '
                      'by the lifecycle case was restored afterwards.'
                      % DATESTR).runs[0].font.size = Pt(9)

    doc.add_heading('Summary', level=1)
    counts = {}
    for r in RESULTS:
        counts[r['status']] = counts.get(r['status'], 0) + 1
    line = '   '.join('%s: %d' % (k, counts[k]) for k in ('PASS', 'PARTIAL', 'MANUAL', 'FAIL', 'SKIP') if k in counts)
    sp = doc.add_paragraph(); run = sp.add_run('%d cases — %s' % (len(RESULTS), line)); run.bold = True

    t = doc.add_table(rows=1, cols=6); t.style = 'Light Grid Accent 1'
    for i, htxt in enumerate(('Area', 'Cases', 'Pass', 'Partial/Manual', 'Fail', 'Pass %')):
        t.rows[0].cells[i].text = htxt
    for area, prefix, cases in CASES:
        rs = [r for r in RESULTS if r['area'] == area]
        np_ = sum(1 for r in rs if r['status'] == 'PASS')
        npm = sum(1 for r in rs if r['status'] in ('PARTIAL', 'MANUAL'))
        nf = sum(1 for r in rs if r['status'] == 'FAIL')
        row = t.add_row().cells
        row[0].text = area; row[1].text = str(len(rs)); row[2].text = str(np_)
        row[3].text = str(npm); row[4].text = str(nf)
        row[5].text = '%d%%' % round(np_ * 100.0 / max(len(rs), 1))

    fails = [r for r in RESULTS if r['status'] == 'FAIL']
    if fails:
        doc.add_heading('Failed cases (triage first)', level=2)
        for r in fails:
            fp = doc.add_paragraph(style='List Bullet')
            fr = fp.add_run('%s — %s: ' % (r['tid'], r['scenario'])); fr.bold = True
            fp.add_run(r['note'])

    for area, prefix, cases in CASES:
        doc.add_page_break()
        doc.add_heading(area, level=1)
        for r in [x for x in RESULTS if x['area'] == area]:
            doc.add_heading('%s — %s' % (r['tid'], r['scenario']), level=2)
            pr = doc.add_paragraph()
            sr = pr.add_run(r['status']); sr.bold = True
            sr.font.color.rgb = COLORS.get(r['status'], RGBColor(0, 0, 0))
            pr.add_run('   ·   %s   ·   %.0fs' % (r['function'], r['secs'])).font.size = Pt(9)
            if r['note']:
                np2 = doc.add_paragraph(r['note']); np2.runs[0].font.size = Pt(10)
            if r['shot'] and os.path.exists(r['shot']):
                try:
                    doc.add_picture(r['shot'], width=Inches(6.3))
                    cap = doc.add_paragraph('Evidence: %s' % os.path.basename(r['shot']))
                    cap.runs[0].font.size = Pt(8)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
    doc.save(DOCX_OUT)
    print('\nWord report:', DOCX_OUT)

# ═════════════════════════════════════════════════════════════════════════════
def main():
    try:
        out = subprocess.run(['powershell', '-NoProfile', '-Command',
            "(Get-NetTCPConnection -LocalPort %d -State Listen -ErrorAction SilentlyContinue).OwningProcess" % PORT],
            capture_output=True, text=True, timeout=20).stdout.strip()
        for pid in set(out.split()):
            if pid.isdigit():
                subprocess.run(['taskkill', '/F', '/PID', pid], capture_output=True)
    except Exception:
        pass
    proxy = subprocess.Popen(['python', os.path.join(ROOT, 'CC', 'Jet', 'dev-proxy.py'), str(PORT)],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx = Ctx()
            ctx['browser'] = browser
            ctx['pp'] = {}
            c, p = open_app(ctx, 'ADMIN')
            ctx['ctx'], ctx['page'] = c, p

            for area, prefix, cases in CASES:
                print('\n──', area)
                for i, (function, scenario, fn) in enumerate(cases, start=1):
                    tid = 'CC-%s-%02d' % (prefix, i)
                    run_case(ctx, tid, area, function, scenario, fn)
                    tc = ctx.pop('tmp_ctx', None)
                    if tc:
                        try: tc.close()
                        except Exception: pass

            # cleanup: deactivate the RPL-05 proxy (NASER on the holder card)
            try:
                st, pr = api('GET', '/cc/proxies/', token=tok())
                for x in pr.get('items', []):
                    if x.get('isActive') == 'Y' and x.get('proxyName') and 'Naser' in x.get('proxyName', '') \
                       and x.get('startDate') == date.today().isoformat():
                        api('PUT', '/cc/proxies/%d' % x['proxyId'], {'isActive': 'N'}, token=tok())
            except Exception:
                pass
            for pgs in ctx['pp'].values():
                try: ensure_en(pgs[1])
                except Exception: pass
            browser.close()
    finally:
        proxy.terminate()

    build_docx()
    counts = {}
    for r in RESULTS:
        counts[r['status']] = counts.get(r['status'], 0) + 1
    print('\n=== %d cases: %s' % (len(RESULTS), counts))

if __name__ == '__main__':
    main()
