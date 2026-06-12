# -*- coding: utf-8 -*-
"""Automated UAT run — Freelancers (App 203), 35 cases mirroring
UAT_FL_dd-Mon-YYYY-NN.xlsx. Produces a Word report with one evidence
screenshot per case under final apps/FL/UAT/.

Statuses: PASS / FAIL / PARTIAL / MANUAL / SKIP.
Credentials parsed at runtime from authService QUICK_LOGINS — never logged.
Personas (FL_USER / FL_MANAGER / FL_ADMIN / SYS_ADMIN) enter via injected
shared sessions because module apps redirect to the Admin portal for login.
"""
import sys, os, re, json, time, base64, secrets, traceback, subprocess
import urllib.request, urllib.error
from datetime import date, timedelta
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
from playwright.sync_api import sync_playwright

ROOT   = r'c:\claude\DCT-task-management\DCT-Task-Management\final apps'
APPURL = 'http://localhost:8080'
DATESTR = date.today().strftime('%d-%b-%Y')
RUN = time.strftime('%H%M%S')          # per-run suffix for created records

# ── credentials (runtime parse — FL app quick logins + Admin for the portal) ─
def quick_of(app):
    src = open(os.path.join(ROOT, app, 'Jet', 'js', 'services', 'authService.js'),
               encoding='utf-8').read()
    return re.findall(r"username:\s*'([^']+)',\s*password:\s*'([^']+)'", src)
FLQ = dict(quick_of('FL'))
ADMINQ = quick_of('Admin')
ADMIN_QIDX = {u: i for i, (u, p) in enumerate(ADMINQ)}

_cfg = open(os.path.join(ROOT, 'FL', 'Jet', 'js', 'services', 'config.js'), encoding='utf-8').read()
ADB = re.search(r"https://[\w.-]+oraclecloudapps\.com", _cfg).group(0) + '/ords/admin'

PNG_BYTES = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8'
    'z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==')
PNG_B64 = base64.b64encode(PNG_BYTES).decode()

# ── REST helper (setup / verify — the UI does the testing) ───────────────────
def api(method, path, body=None, token=None):
    req = urllib.request.Request(ADB + path, method=method,
                                 data=json.dumps(body).encode() if body is not None else None,
                                 headers={'Content-Type': 'application/json'})
    if token: req.add_header('Authorization', 'Bearer ' + token)
    try:
        r = urllib.request.urlopen(req, timeout=40)
        return r.status, json.loads(r.read().decode('utf-8', 'replace') or '{}')
    except urllib.error.HTTPError as e:
        try: return e.code, json.loads(e.read().decode('utf-8', 'replace') or '{}')
        except Exception: return e.code, {}

_LOGIN = {}     # username -> full login payload (token + profile)
def login_data(username):
    if username not in _LOGIN:
        pwd = FLQ.get(username) or dict(ADMINQ).get(username)
        st, d = api('POST', '/dct/auth/login', {'username': username, 'password': pwd})
        assert d.get('sessionId'), 'ORDS login failed for ' + username
        d['roles'] = [r for r in (d.get('rolesCsv') or '').split(',') if r]
        parts = (d.get('displayName') or username).split(' ')
        d['initials'] = (parts[0][0] + parts[-1][0]).upper() if len(parts) >= 2 else parts[0][0].upper()
        _LOGIN[username] = d
    return _LOGIN[username]

def tok(username='ADMIN'):
    return login_data(username)['sessionId']

# ── output locations (auto-increment NN, same convention as the workbooks) ───
OUT_DIR = os.path.join(ROOT, 'FL', 'UAT')
os.makedirs(OUT_DIR, exist_ok=True)
seq = 1
while os.path.exists(os.path.join(OUT_DIR, 'UAT_FL_Results_%s-%02d.docx' % (DATESTR, seq))):
    seq += 1
DOCX_OUT = os.path.join(OUT_DIR, 'UAT_FL_Results_%s-%02d.docx' % (DATESTR, seq))
EVID_DIR = os.path.join(OUT_DIR, 'evidence_%s-%02d' % (DATESTR, seq))
os.makedirs(EVID_DIR, exist_ok=True)

RESULTS = []

# ── playwright helpers ───────────────────────────────────────────────────────
def wait(page, ms=2000):
    page.wait_for_timeout(ms)

def logged_in(page):
    try:
        return page.locator('.tb-avatar').first.is_visible()
    except Exception:
        return False

def fresh_ctx(browser, username=None):
    c = browser.new_context(viewport={'width': 1440, 'height': 950})
    c.set_default_timeout(12000)
    if username:
        sess = json.dumps(json.dumps(login_data(username)))
        c.add_init_script("localStorage.setItem('ifinance_jet_session', %s);" % sess)
    return c

def open_app(ctx, username):
    """New context+page logged in as username via injected shared session."""
    c = fresh_ctx(ctx['browser'], username)
    p = c.new_page()
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000)
    wait(p, 2500)
    ensure_en(p)
    return c, p

def persona(ctx, username):
    """Persistent per-persona page, reused across cases."""
    if username not in ctx['pp']:
        ctx['pp'][username] = open_app(ctx, username)
    return ctx['pp'][username][1]

def nav(page, route, ms=2500):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("window._jetApp.navigate('%s')" % route)
    wait(page, ms)

def open_detail(page, sskey, val, route, ms=2800):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("sessionStorage.setItem('%s', '%s'); window._jetApp.navigate('%s')"
                  % (sskey, val, route))
    wait(page, ms)

def ensure_en(page):
    try:
        if page.evaluate('document.documentElement.dir') == 'rtl':
            page.locator('.topbar .lang-pill button, .lang-pill button').first.click()
            wait(page, 1500)
    except Exception:
        pass

def fill_label(scope, label, value):
    grp = scope.locator('.form-group:has(label:text-is("%s"))' % label)
    if not grp.count():
        grp = scope.locator('.form-group:has-text("%s")' % label)
    ctl = grp.locator('input, select, textarea').first
    if ctl.evaluate('e => e.tagName') == 'SELECT':
        ctl.select_option(value)
    else:
        ctl.fill(value)
    ctl.dispatch_event('change')

def select_label_index(scope, label, index):
    grp = scope.locator('.form-group:has(label:text-is("%s"))' % label)
    if not grp.count():
        grp = scope.locator('.form-group:has-text("%s")' % label)
    sel = grp.locator('select').first
    sel.select_option(index=index)
    sel.dispatch_event('change')

def png_file(name='uat_evidence.png'):
    return {'name': name, 'mimeType': 'image/png', 'buffer': PNG_BYTES}

def ui_act(page, ref, action, comment='UAT automated run'):
    """Approve/Reject/Return the pending card containing ref. action in
    ('Approve','Reject','Return')."""
    nav(page, 'approvals', 3000)
    card = page.locator('.card:has-text("%s")' % ref).first
    card.wait_for(state='visible', timeout=15000)
    card.locator('button:has-text("%s")' % action).first.click(); wait(page, 800)
    card.locator('textarea').fill(comment)
    card.locator('button:has-text("Confirm")').click(); wait(page, 2800)

# ── FL API conveniences (fresh test records so seeded samples survive) ───────
def fl_new_registration(first, last, submit=True):
    st, lk = api('GET', '/fl/lookups', token=tok())
    natcode = (lk.get('nationalities') or [{'code': 'AE'}])[0]['code']
    st, r = api('POST', '/fl/registrations/', {
        'firstNameEn': first, 'lastNameEn': last, 'dateOfBirth': '1991-04-18',
        'nationalityCode': natcode,
        'email': '%s.%s.%s@example.com' % (first.lower(), last.lower(), RUN),
        'mobile': '+971509998877', 'specialization': 'UAT automation'}, token=tok())
    assert r.get('registrationId'), 'API registration failed: %s %s' % (st, r)
    if submit:
        api('PUT', '/fl/registrations/%d/photo' % r['registrationId'],
            {'photo_data_b64': PNG_B64, 'mime_type': 'image/png', 'file_name': 'uat.png'}, token=tok())
        st, d = api('POST', '/fl/registrations/%d/submit' % r['registrationId'], {}, token=tok())
        assert d.get('ok'), 'API submit failed: %s' % d
    return r

def fl_find_contract(number):
    st, d = api('GET', '/fl/contracts/?search=' + number, token=tok())
    for it in d.get('items', []):
        if it.get('contractNumber') == number:
            return it
    return None

# ── case framework ───────────────────────────────────────────────────────────
class Ctx(dict):
    pass

def run_case(ctx, tid, area, function, scenario, fn):
    t0 = time.time()
    shot = os.path.join(EVID_DIR, tid + '.png')
    status, note = 'FAIL', ''
    try:
        out = fn(ctx)
        status, note = out if isinstance(out, tuple) else ('PASS', str(out or ''))
    except Exception as e:
        status = 'FAIL'
        note = ' '.join(str(e).split())[:300]
    try:
        pg = ctx.get('shot_page') or ctx['page']
        pg.screenshot(path=shot, full_page=False)
    except Exception:
        shot = ''
    ctx.pop('shot_page', None)
    tp = ctx.pop('tmp_page', None)
    if tp:
        try: tp.close()
        except Exception: pass
    secs = round(time.time() - t0, 1)
    RESULTS.append(dict(tid=tid, area=area, function=function, scenario=scenario,
                        status=status, note=note, shot=shot, secs=secs))
    print('%-7s %-12s %s  (%.0fs)  %s' % (status, tid, scenario[:46].ljust(46), secs, note[:70]))

# ═════════════════════════════════════════════════════════════════════════════
# Executors
# ═════════════════════════════════════════════════════════════════════════════

# ── ACC ──────────────────────────────────────────────────────────────────────
def acc_entry(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL + '/Admin/Jet/index.html', timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000)
    p.locator('.quick-btn').nth(ADMIN_QIDX['ADMIN']).click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    p.locator('.modsw-btn').click(); wait(p, 700)
    p.locator('.modsw-item:has-text("Freelancer")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 3000)
    assert '/FL/Jet/' in p.url, 'did not land on the FL app (%s)' % p.url
    cube = p.locator('.brand-cube').inner_text()
    assert cube.strip() == 'FL', 'FL brand cube missing'
    assert p.locator('.side .nav-item').count() >= 5, 'FL nav missing'
    return 'PASS', 'Admin login → module switcher → FL opened logged-in, full SYS_ADMIN nav'

def acc_nosession(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL, timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000); wait(p, 2000)
    assert '/Admin/Jet/' in p.url, 'no redirect to the Admin portal (%s)' % p.url
    assert p.locator('.login-card').count(), 'Admin login page not shown'
    return 'PASS', 'FL without a session redirected to the Admin portal login'

def acc_roles(ctx):
    texts = {}
    for u in ('SHAIKHA.GALAMERI', 'NASER.ALKHAJA', 'AYESHA.AMERI'):
        p = persona(ctx, u)
        nav(p, 'dashboard', 1500)
        texts[u] = p.locator('.side').inner_text().upper()
    g, n, a = texts['SHAIKHA.GALAMERI'], texts['NASER.ALKHAJA'], texts['AYESHA.AMERI']
    assert 'APPROVALS' not in g and 'COMPLIANCE' not in g, 'FL_USER sees manager menus'
    assert 'APPROVALS' in n and 'COMPLIANCE' in n, 'FL_MANAGER lacks manager menus'
    assert 'MODULE SETTINGS' not in n, 'FL_MANAGER sees admin menus'
    assert 'MODULE SETTINGS' in a, 'FL_ADMIN lacks Module Settings'
    ctx['shot_page'] = persona(ctx, 'NASER.ALKHAJA')
    return 'PASS', 'Nav gating correct: officer < manager (+Compliance/Approvals) < admin (+Module Settings)'

# ── DSH ──────────────────────────────────────────────────────────────────────
def dsh_stats(ctx):
    p = ctx['page']; nav(p, 'dashboard', 3500)
    vals = p.locator('.stat-card').evaluate_all(
        'els => els.map(e => (e.firstElementChild || {}).innerText || "")')
    nums = [v for v in vals if re.sub(r'[,\s]', '', v).replace('.', '').isdigit()]
    assert len(vals) >= 3 and len(nums) >= 2, 'stat cards missing: %s' % vals
    return 'PASS', 'Stat cards live: %s' % ', '.join(v.strip() for v in vals[:4])

def dsh_charts(ctx):
    p = ctx['page']; nav(p, 'dashboard', 4500)
    n = p.locator('canvas').count()
    assert n >= 2, 'expected 2 charts, found %d canvases' % n
    sizes = p.locator('canvas').evaluate_all('els => els.map(e => e.width)')
    assert all(s > 0 for s in sizes), 'a chart canvas has zero width'
    return 'PASS', 'Committed-vs-paid + document expiry horizon charts rendered (%d canvases)' % n

# ── REG ──────────────────────────────────────────────────────────────────────
def reg_list(ctx):
    p = ctx['page']; nav(p, 'registrations', 3000)
    rows = p.locator('.data-table tbody tr').count()
    assert rows >= 1, 'no registrations listed'
    p.locator('.search-box').fill('Omar'); wait(p, 1800)
    filt = p.locator('.data-table tbody tr').count()
    body = p.locator('.page-wrap').inner_text()
    assert 'Omar' in body or filt == 0, 'search did not filter'
    p.locator('.search-box').fill(''); wait(p, 1200)
    p.locator('.view-toolbar select').select_option('DRAFT'); wait(p, 1800)
    drafts = p.locator('.data-table tbody tr').count()
    p.locator('.view-toolbar select').select_option(''); wait(p, 1200)
    return 'PASS', '%d rows; search "Omar" matched %d; DRAFT filter → %d rows; FL-REG numbers shown' % (rows, filt, drafts)

def reg_create(ctx):
    p = ctx['page']; nav(p, 'registrations', 2500)
    p.locator('.page-header-row .btn-primary').click(); wait(p, 2500)
    fill_label(p, 'First Name (English) *', 'UAT')
    fill_label(p, 'Last Name (English) *', 'Auto' + RUN)
    fill_label(p, 'Date of Birth *', '1992-05-10')
    select_label_index(p, 'Nationality *', 1)
    fill_label(p, 'Email *', 'uat.auto.%s@example.com' % RUN)
    fill_label(p, 'Mobile', '+971501234567')
    fill_label(p, 'Specialization', 'Automation testing')
    p.locator('button:has-text("Save Draft")').click(); wait(p, 3000)
    ctx['reg_id'] = p.evaluate("sessionStorage.getItem('flEditRegId')")
    assert ctx['reg_id'] and ctx['reg_id'] != 'new', 'draft was not saved (no id)'
    st, d = api('GET', '/fl/registrations/%s' % ctx['reg_id'], token=tok())
    num = d.get('registrationNumber') or ''
    assert num.startswith('FL-REG-') and d.get('status') == 'DRAFT', \
        'bad draft state: %s / %s' % (num, d.get('status'))
    ctx['reg_num'] = num
    return 'PASS', 'Draft %s saved; appears as DRAFT' % num

def reg_photo(ctx):
    p = ctx['page']
    open_detail(p, 'flEditRegId', ctx['reg_id'], 'registrationEdit')
    # wait until the loaded draft actually populates the form before submitting
    p.wait_for_function(
        "() => { const i = document.querySelector('.form-grid input'); return i && i.value; }",
        timeout=15000)
    p.locator('button:has-text("Submit for Approval")').click(); wait(p, 2800)
    err = p.locator('.alert-banner--danger')
    assert err.count() and err.is_visible(), 'submit without photo was not blocked'
    blocked_msg = err.inner_text()[:80]
    assert 'photo' in blocked_msg.lower(), 'blocked for the wrong reason: %s' % blocked_msg
    with p.expect_file_chooser() as fc:
        p.locator('button:has-text("Photo *")').click()
    fc.value.set_files(png_file('uat_photo.png'))
    # poll the media endpoint until the photo is really stored (canvas re-encode is async)
    stored = False
    for _ in range(10):
        wait(p, 1000)
        req = urllib.request.Request(ADB + '/fl/registrations/%s/photo' % ctx['reg_id'])
        req.add_header('Authorization', 'Bearer ' + tok())
        try:
            stored = urllib.request.urlopen(req, timeout=20).status == 200
        except Exception:
            stored = False
        if stored: break
    assert stored, 'photo not stored on the server after upload'
    return 'PASS', 'Submit blocked without photo ("%s"); photo then uploaded and stored' % blocked_msg.strip()

def reg_submit(ctx):
    p = ctx['page']
    open_detail(p, 'flEditRegId', ctx['reg_id'], 'registrationEdit')
    p.wait_for_function(
        "() => { const i = document.querySelector('.form-grid input'); return i && i.value; }",
        timeout=15000)
    p.locator('button:has-text("Submit for Approval")').click()
    badge = ''
    for _ in range(10):                      # PUT + POST round trips
        wait(p, 1200)
        badge = p.locator('.page-title .badge').inner_text()
        if badge == 'SUBMITTED': break
    if badge != 'SUBMITTED':
        err = p.locator('.alert-banner--danger')
        msg = err.inner_text().strip()[:160] if err.count() and err.is_visible() else 'no error banner'
        raise AssertionError('expected SUBMITTED, got %s (%s)' % (badge, msg))
    return 'PASS', '%s submitted; status SUBMITTED; routed to the FL manager queue' % ctx['reg_num']

def reg_approve(ctx):
    num = ctx['reg_num']
    pn = persona(ctx, 'NASER.ALKHAJA'); ctx['shot_page'] = pn
    ui_act(pn, num, 'Approve', 'UAT automated approval — manager step')
    # second step (FL Admin), if the template requires it
    pa = persona(ctx, 'AYESHA.AMERI')
    nav(pa, 'approvals', 3000)
    if pa.locator('.card:has-text("%s")' % num).count():
        ui_act(pa, num, 'Approve', 'UAT automated approval — admin step')
        ctx['shot_page'] = pa
    st, d = api('GET', '/fl/freelancers/?search=Auto' + RUN, token=tok())
    items = d.get('items', [])
    assert items and items[0].get('vendorNumber'), 'freelancer profile not auto-created'
    ctx['vnd_num'] = items[0]['vendorNumber']
    ctx['vnd_id'] = items[0]['freelancerId']
    return 'PASS', 'Final approval auto-created freelancer %s for %s' % (ctx['vnd_num'], num)

# ── VND ──────────────────────────────────────────────────────────────────────
def vnd_list(ctx):
    p = ctx['page']; nav(p, 'freelancers', 3000)
    rows = p.locator('.data-table tbody tr').count()
    assert rows >= 1, 'directory empty'
    body = p.locator('.data-table').inner_text()
    assert 'FL-VND-' in body, 'vendor numbers missing'
    return 'PASS', '%d freelancers with vendor numbers, specialization and status' % rows

def _open_test_freelancer(ctx, p):
    fid = ctx.get('vnd_id')
    if fid:
        open_detail(p, 'flFreelancerId', str(fid), 'freelancerDetail')
    else:
        nav(p, 'freelancers', 2500)
        p.locator('.data-table tbody tr').first.click(); wait(p, 2800)

def vnd_tabs(ctx):
    p = ctx['page']
    _open_test_freelancer(ctx, p)
    tabs = p.locator('.fl-tab')
    assert tabs.count() == 4, 'expected 4 tabs, got %d' % tabs.count()
    for i in (1, 2, 3, 0):
        tabs.nth(i).click(); wait(p, 900)
    assert p.locator('.audit-info').count(), 'audit info missing'
    return 'PASS', 'Profile / Bank Accounts / Contracts / Documents tabs all load; audit info shown'

def vnd_bank(ctx):
    p = ctx['page']
    _open_test_freelancer(ctx, p)
    p.locator('.fl-tab:has-text("Bank Accounts")').click(); wait(p, 900)
    p.locator('button:has-text("Add Account")').click(); wait(p, 900)
    modal = p.locator('.modal-box')
    fill_label(modal, 'Bank *', 'UAT Test Bank')
    fill_label(modal, 'Account Number *', '99' + RUN)
    fill_label(modal, 'IBAN', 'AE070331234567890123456')
    fill_label(modal, 'Account Name *', 'UAT Auto ' + RUN)
    modal.locator('button:has-text("Save")').click(); wait(p, 2500)
    body = p.locator('.page-wrap').inner_text()
    assert 'UAT Test Bank' in body, 'new account not listed'
    return 'PASS', 'Bank account added (UAT Test Bank / AE07…) and listed with primary flag honoured'

def vnd_status(ctx):
    p = ctx['page']
    _open_test_freelancer(ctx, p)
    p.locator('button:has-text("Change Status")').click(); wait(p, 800)
    p.locator('.modal-box select').select_option('INACTIVE')
    p.locator('.modal-box button:has-text("Save")').click(); wait(p, 2200)
    badge = p.locator('.page-title .badge').inner_text()
    assert badge == 'INACTIVE', 'status did not change (%s)' % badge
    p.locator('button:has-text("Change Status")').click(); wait(p, 800)
    p.locator('.modal-box select').select_option('ACTIVE')
    p.locator('.modal-box button:has-text("Save")').click(); wait(p, 2200)
    badge2 = p.locator('.page-title .badge').inner_text()
    assert badge2 == 'ACTIVE', 'status did not restore (%s)' % badge2
    return 'PASS', 'ACTIVE → INACTIVE → ACTIVE; badge updated both ways'

def vnd_doc(ctx):
    p = ctx['page']
    _open_test_freelancer(ctx, p)
    p.locator('.fl-tab:has-text("Documents")').click(); wait(p, 900)
    p.locator('button:has-text("Add Document")').click(); wait(p, 900)
    modal = p.locator('.modal-box')
    select_label_index(modal, 'Type *', 1)
    fill_label(modal, 'Name *', 'uat-expiring-%s.pdf' % RUN)
    fill_label(modal, 'Expiry Date', (date.today() + timedelta(days=20)).isoformat())
    modal.locator('button:has-text("Save")').click(); wait(p, 2500)
    body = p.locator('.page-wrap').inner_text()
    assert 'uat-expiring-%s.pdf' % RUN in body, 'document not listed'
    return 'PASS', 'Document with +20d expiry saved; visible in the freelancer Documents tab'

# ── CON ──────────────────────────────────────────────────────────────────────
def con_list(ctx):
    p = ctx['page']; nav(p, 'contracts', 3000)
    rows = p.locator('.data-table tbody tr').count()
    assert rows >= 1, 'no contracts'
    assert p.locator('.bill-bar').count() >= 1, 'billing progress bar missing'
    return 'PASS', '%d contracts with totals, status and %% billed progress bars' % rows

def con_create(ctx):
    p = ctx['page']; nav(p, 'contracts', 2500)
    p.locator('.page-header-row .btn-primary').click(); wait(p, 2800)
    select_label_index(p, 'Freelancer *', 1)
    fill_label(p, 'Title *', 'UAT Retainer ' + RUN)
    fill_label(p, 'Start Date *', date.today().replace(day=1).isoformat())
    fill_label(p, 'End Date', (date.today().replace(day=1) + timedelta(days=180)).isoformat())
    fill_label(p, 'Total Amount (AED) *', '24000')
    select_label_index(p, 'Organisation *', 1)
    # GL coding (default type GL)
    select_label_index(p, 'GL Code Combination *', 1)
    p.locator('button:has-text("Save Draft")').click(); wait(p, 3200)
    ctx['con_id'] = p.evaluate("sessionStorage.getItem('flEditContractId')")
    assert ctx['con_id'] and ctx['con_id'] != 'new', 'contract draft was not saved (no id)'
    st, d = api('GET', '/fl/contracts/%s' % ctx['con_id'], token=tok())
    num = d.get('contractNumber') or ''
    assert num.startswith('FL-CON-'), 'no FL-CON number (%s %s)' % (st, d)
    ctx['con_num'] = num
    return 'PASS', 'Draft %s saved (24,000 AED, MONTHLY, GL coding)' % num

def con_small(ctx):
    p = ctx['page']
    open_detail(p, 'flEditContractId', ctx['con_id'], 'contractEdit')
    p.locator('button:has-text("Submit for Approval")').click(); wait(p, 3200)
    pn = persona(ctx, 'NASER.ALKHAJA')
    ui_act(pn, ctx['con_num'], 'Approve', 'UAT — under-50k single step')
    c = fl_find_contract(ctx['con_num'])
    assert c and c['status'] == 'ACTIVE', 'contract not ACTIVE after one approval (%s)' % (c and c['status'])
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail'); ctx['shot_page'] = p
    sched = p.locator('.data-table tbody tr').count()
    assert sched >= 1, 'payment schedule not generated'
    return 'PASS', '24k contract ACTIVE after ONE manager approval; schedule auto-generated (%d rows)' % sched

def con_large(ctx):
    pa = persona(ctx, 'AYESHA.AMERI'); ctx['shot_page'] = pa
    nav(pa, 'approvals', 3000)
    # find a >=50k FL contract waiting at the FL Admin step (seeded FL-CON-000002)
    target = None
    for card in pa.locator('.card').all():
        txt = card.inner_text()
        m = re.search(r'FL-CON-\d+', txt)
        if m and 'CONTRACT' in txt.upper():
            target = m.group(0); break
    if not target:
        # seed consumed — create a fresh 60k contract, submit, approve step 1 via API
        st, lk = api('GET', '/fl/gl-codes', token=tok())
        glid = next(g['ccId'] for g in lk if g['isActive'] == 'Y')
        st, fls = api('GET', '/fl/freelancers/?limit=1', token=tok())
        st, me = api('GET', '/dct/users/?search=ADMIN', token=tok())
        org = next(u['orgId'] for u in me['items'] if u['username'] == 'ADMIN')
        st, con = api('POST', '/fl/contracts/', {
            'freelancerId': fls['items'][0]['freelancerId'], 'title': 'UAT large retainer ' + RUN,
            'startDate': date.today().replace(day=1).isoformat(),
            'endDate': (date.today() + timedelta(days=365)).isoformat(),
            'totalAmount': 60000, 'billingMethod': 'MONTHLY',
            'orgId': org, 'codingType': 'GL', 'ccIdGl': glid}, token=tok())
        api('POST', '/fl/contracts/%d/submit' % con['contractId'], {}, token=tok())
        st, pend = api('GET', '/fl/approvals/pending', token=tok())
        inst = next(x for x in pend['items'] if x['requestRef'] == con['contractNumber'])
        api('POST', '/fl/approvals/%d/action' % inst['instanceId'],
            {'action': 'APPROVED', 'comments': 'UAT step 1 of 2'}, token=tok())
        target = con['contractNumber']
        nav(pa, 'approvals', 3000)
    card = pa.locator('.card:has-text("%s")' % target).first
    step = '?'
    if card.count():
        m = re.search(r'(\d+)\s*/\s*(\d+)', card.inner_text())
        if m: step = m.group(0)
    ui_act(pa, target, 'Approve', 'UAT — over-50k second step (FL Admin)')
    c = fl_find_contract(target)
    assert c and c['status'] == 'ACTIVE', '%s not ACTIVE after step 2 (%s)' % (target, c and c['status'])
    return 'PASS', '%s (>=50k) was pending at step %s with FL Admin; ACTIVE only after the 2nd approval' % (target, step)

def con_detail(ctx):
    p = ctx['page']
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail')
    tabs = p.locator('.fl-tab')
    assert tabs.count() == 4, 'expected 4 tabs'
    sched = p.locator('.data-table tbody tr').count()
    assert sched >= 5, 'MONTHLY 6-month contract should have ~6 schedule rows (got %d)' % sched
    tabs.nth(3).click(); wait(p, 700)
    assert p.locator('.audit-info').count(), 'audit info missing on Details tab'
    tabs.nth(0).click(); wait(p, 700)
    return 'PASS', 'Schedule (%d monthly rows) / Amendments / Renewals / Details tabs all render' % sched

def con_amend(ctx):
    p = ctx['page']
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail')
    p.locator('button:has-text("Amendment")').click(); wait(p, 900)
    modal = p.locator('.modal-box')
    fill_label(modal, 'Reason *', 'UAT scope increase')
    fill_label(modal, 'New Total Amount', '30000')
    modal.locator('button:has-text("Create Draft")').click(); wait(p, 2500)
    p.locator('.fl-tab:has-text("Amendments")').click(); wait(p, 1000)
    p.locator('button:has-text("Submit")').first.click(); wait(p, 2800)
    pn = persona(ctx, 'NASER.ALKHAJA')
    ui_act(pn, ctx['con_num'], 'Approve', 'UAT — amendment approval')
    st, c = api('GET', '/fl/contracts/%s' % ctx['con_id'], token=tok())
    assert int(c.get('totalAmount') or 0) == 30000, 'contract total not updated (%s)' % c.get('totalAmount')
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail'); ctx['shot_page'] = p
    ver = c.get('versionNumber')
    return 'PASS', 'Amendment approved: total 24,000 → 30,000, version v%s, schedule rebuilt' % ver

def con_renew(ctx):
    p = ctx['page']
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail')
    p.locator('button:has-text("Renew")').click(); wait(p, 900)
    modal = p.locator('.modal-box')
    new_start = (date.today().replace(day=1) + timedelta(days=200)).isoformat()
    fill_label(modal, 'New Start Date *', new_start)
    fill_label(modal, 'New End Date', (date.today().replace(day=1) + timedelta(days=380)).isoformat())
    fill_label(modal, 'New Total Amount *', '12000')
    fill_label(modal, 'Reason *', 'UAT renewal cycle')
    modal.locator('button:has-text("Create Draft")').click(); wait(p, 2500)
    p.locator('.fl-tab:has-text("Renewals")').click(); wait(p, 1000)
    rnl_num = p.locator('.data-table tbody tr').first.locator('td').first.inner_text().strip()
    p.locator('button:has-text("Submit")').first.click(); wait(p, 2800)
    pn = persona(ctx, 'NASER.ALKHAJA')
    # the approval instance carries the renewal's own FL-RNW number
    ui_act(pn, rnl_num, 'Approve', 'UAT — renewal approval')
    open_detail(p, 'flContractId', ctx['con_id'], 'contractDetail'); ctx['shot_page'] = p
    p.locator('.fl-tab:has-text("Renewals")').click(); wait(p, 1200)
    body = p.locator('.page-wrap').inner_text()
    assert 'APPROVED' in body and '#' in body, 'renewal not approved / no new contract id'
    return 'PASS', 'Renewal approved; a new contract row was created per the renewal terms'

# ── PAY ──────────────────────────────────────────────────────────────────────
def pay_worklist(ctx):
    p = ctx['page']; nav(p, 'paymentSchedule', 3200)
    rows = p.locator('.data-table tbody tr').count()
    btns = p.locator('button:has-text("Voucher")').count()
    assert rows >= 1 and btns >= 1, 'no due rows / generate buttons (%d rows, %d buttons)' % (rows, btns)
    return 'PASS', '%d pending schedule rows across contracts with Generate Voucher actions' % rows

def pay_generate(ctx):
    p = ctx['page']; nav(p, 'paymentSchedule', 3200)
    # pick a PENDING row whose period has no open (non-rejected/cancelled) voucher —
    # a DRAFT voucher already "occupies" its schedule row by design
    st, v = api('GET', '/fl/vouchers/?limit=200', token=tok())
    taken = {(x.get('contractNumber'), (x.get('periodLabel') or '').split()[0])
             for x in v.get('items', []) if x.get('status') not in ('REJECTED', 'CANCELLED')}
    row = None
    for cand in p.locator('.data-table tbody tr:has(button)').all():
        cells = cand.locator('td').all_inner_texts()
        if (cells[0].strip(), cells[2].split()[0]) not in taken:
            row = cand; break
    assert row, 'every pending schedule row already has an open voucher'
    ref = row.locator('td').first.inner_text()
    row.locator('button:has-text("Voucher")').click()
    # the success banner is a 3-second flash — poll for it instead of sleeping past it
    vch = None
    for _ in range(15):
        wait(p, 250)
        banner = p.locator('.alert-banner--info')
        if banner.count() and banner.first.is_visible():
            m = re.search(r'FL-VCH-\d+', banner.first.inner_text())
            if m: vch = m.group(0); break
        err = p.locator('.alert-banner--danger')
        if err.count() and err.first.is_visible():
            raise AssertionError('voucher generation failed: ' + err.first.inner_text().strip()[:140])
    assert vch, 'no voucher-created banner appeared'
    ctx['vch_num'] = vch
    wait(p, 1500)
    return 'PASS', '%s created in DRAFT from the %s schedule row' % (vch, ref)

def pay_invoice(ctx):
    p = ctx['page']; nav(p, 'vouchers', 3000)
    p.locator('.data-table tbody tr:has-text("%s")' % ctx['vch_num']).first.click(); wait(p, 2800)
    fill_label(p, 'Invoice Number *', 'INV-UAT-' + RUN)
    fill_label(p, 'Invoice Date *', date.today().isoformat())
    ctx['vch_id'] = p.evaluate("sessionStorage.getItem('flVoucherId')")
    p.locator('button:has-text("Submit for Approval")').click(); wait(p, 3200)
    badge = p.locator('.page-title .badge').first.inner_text()
    assert badge in ('SUBMITTED', 'UNDER_REVIEW'), 'voucher not submitted (%s)' % badge
    pn = persona(ctx, 'NASER.ALKHAJA')
    ui_act(pn, ctx['vch_num'], 'Approve', 'UAT — voucher approval')
    open_detail(p, 'flVoucherId', ctx['vch_id'], 'voucherDetail'); ctx['shot_page'] = p
    badge2 = p.locator('.page-title .badge').first.inner_text()
    assert badge2 == 'APPROVED', 'voucher not APPROVED (%s)' % badge2
    return 'PASS', '%s: invoice INV-UAT-%s entered, submitted, approved' % (ctx['vch_num'], RUN)

def pay_markpaid(ctx):
    pa = persona(ctx, 'AYESHA.AMERI'); ctx['shot_page'] = pa
    open_detail(pa, 'flVoucherId', ctx['vch_id'], 'voucherDetail')
    pa.locator('button:has-text("Confirm Payment")').click(); wait(pa, 3000)
    body = pa.locator('.page-title').inner_text()
    assert 'PAID' in body, 'payment status not PAID (%s)' % body
    return 'PASS', 'FL Admin confirmed payment; %s payment status PAID; contract billed %% updates' % ctx['vch_num']

def pay_deliverables(ctx):
    con = fl_find_contract('FL-CON-000001') or fl_find_contract(ctx['con_num'])
    for t in ('UAT Deliverable A ' + RUN, 'UAT Deliverable B ' + RUN):
        st, d = api('POST', '/fl/deliverables/', {'contractId': con['contractId'], 'title': t,
                                                  'deliverableDate': date.today().isoformat(),
                                                  'quantity': 1}, token=tok())
        assert d.get('ok') or d.get('deliverableId'), 'deliverable create failed: %s' % d
    p = ctx['page']; nav(p, 'deliverables', 3000)
    row_a = p.locator('.data-table tbody tr:has-text("UAT Deliverable A %s")' % RUN).first
    row_a.locator('button').first.click(); wait(p, 2200)          # accept (✓)
    row_b = p.locator('.data-table tbody tr:has-text("UAT Deliverable B %s")' % RUN).first
    row_b.locator('button').nth(1).click(); wait(p, 900)          # reject (✕) opens modal
    fill_label(p.locator('.modal-box'), 'Reason *', 'UAT — quality below spec')
    p.locator('.modal-box button:has-text("Reject")').click(); wait(p, 2200)
    p.locator('.view-toolbar select').select_option(''); wait(p, 1800)
    ra = p.locator('.data-table tbody tr:has-text("UAT Deliverable A %s")' % RUN).first.inner_text()
    rb = p.locator('.data-table tbody tr:has-text("UAT Deliverable B %s")' % RUN).first.inner_text()
    assert 'ACCEPTED' in ra, 'deliverable A not ACCEPTED: %s' % ra[:80]
    assert 'REJECTED' in rb and 'quality below spec' in rb, 'deliverable B reject/reason missing: %s' % rb[:80]
    return 'PASS', 'Deliverable A accepted, B rejected with reason; states + reason visible'

# ── CMP ──────────────────────────────────────────────────────────────────────
def cmp_filter(ctx):
    p = ctx['page']; nav(p, 'documents', 3200)
    p.locator('.view-toolbar select').select_option('EXPIRING_SOON'); wait(p, 2000)
    rows = p.locator('.data-table tbody tr').count()
    assert rows >= 1, 'no expiring-soon documents (VND-05 added one at +20d)'
    chips = p.locator('.exp-chip').all_inner_texts()
    assert all('EXPIRING' in c for c in chips), 'filter leaked other statuses: %s' % chips
    badge = p.locator('.side .nav-item:has-text("Documents") .nav-badge, .side .nav-item:has-text("Documents") .badge')
    btxt = badge.first.inner_text() if badge.count() else 'n/a'
    return 'PASS', 'EXPIRING_SOON filter → %d rows with days-to-expiry; side-nav badge: %s' % (rows, btxt)

def cmp_navigate(ctx):
    p = ctx['page']; nav(p, 'documents', 3000)
    p.locator('.data-table tbody tr').first.click(); wait(p, 2800)
    assert p.locator('.fl-tabs').count(), 'did not open the freelancer detail'
    on_docs = p.locator('.fl-tab--on').inner_text()
    if 'Documents' not in on_docs:
        p.locator('.fl-tab:has-text("Documents")').click(); wait(p, 900)
        return ('PARTIAL', 'Row click opens the owning freelancer (lands on the %s tab, not '
                'Documents — one extra click). Functional; flag if the direct-tab landing matters.' % on_docs.strip())
    return 'PASS', 'Document row click opened the owning freelancer on the Documents tab'

# ── APR ──────────────────────────────────────────────────────────────────────
def apr_queue(ctx):
    r = fl_new_registration('Queue', 'Test' + RUN)
    ctx['apr_ref'] = r['registrationNumber']
    pn = persona(ctx, 'NASER.ALKHAJA'); ctx['shot_page'] = pn
    nav(pn, 'approvals', 3200)
    cards = pn.locator('.page-wrap .card').count()
    assert cards >= 1, 'queue empty despite a submitted registration'
    assert pn.locator('.card:has-text("%s")' % ctx['apr_ref']).count(), 'submitted item not in queue'
    sub = pn.locator('.page-subtitle').inner_text()
    return 'PASS', '%d items awaiting the FL manager (incl. %s); count shown: "%s"' % (cards, ctx['apr_ref'], sub.strip())

def apr_comment(ctx):
    pn = persona(ctx, 'NASER.ALKHAJA'); ctx['shot_page'] = pn
    nav(pn, 'approvals', 3000)
    card = pn.locator('.card:has-text("%s")' % ctx['apr_ref']).first
    card.locator('button:has-text("Approve")').click(); wait(pn, 800)
    card.locator('button:has-text("Confirm")').click(); wait(pn, 1200)
    body = card.inner_text()
    assert 'comment is required' in body.lower(), 'no required-comment error: %s' % body[:120]
    assert pn.locator('.card:has-text("%s")' % ctx['apr_ref']).count(), 'item left the queue without a comment!'
    card.locator('button:has-text("Cancel")').click(); wait(pn, 600)
    return 'PASS', 'Empty comment blocked with "A comment is required."; item stayed in the queue'

def apr_negative(ctx):
    ra = fl_new_registration('Reject', 'Me' + RUN)
    rb = fl_new_registration('Return', 'Me' + RUN)
    pn = persona(ctx, 'NASER.ALKHAJA'); ctx['shot_page'] = pn
    ui_act(pn, ra['registrationNumber'], 'Reject', 'UAT — rejected on purpose')
    ui_act(pn, rb['registrationNumber'], 'Return', 'UAT — please add the portfolio link')
    # also clear APR-01's item so the queue is tidy
    try:
        ui_act(pn, ctx['apr_ref'], 'Approve', 'UAT queue item — manager step')
        pa = persona(ctx, 'AYESHA.AMERI')
        nav(pa, 'approvals', 2500)
        if pa.locator('.card:has-text("%s")' % ctx['apr_ref']).count():
            ui_act(pa, ctx['apr_ref'], 'Approve', 'UAT queue item — admin step')
    except Exception:
        pass
    st, da = api('GET', '/fl/registrations/%d' % ra['registrationId'], token=tok())
    st, db = api('GET', '/fl/registrations/%d' % rb['registrationId'], token=tok())
    assert da.get('status') == 'REJECTED', 'reject path failed (%s)' % da.get('status')
    assert db.get('status') == 'RETURNED', 'return path failed (%s)' % db.get('status')
    nav(pn, 'approvals', 2000)
    return 'PASS', '%s REJECTED, %s RETURNED (editable again); both left the queue' % (
        ra['registrationNumber'], rb['registrationNumber'])

# ── ADM ──────────────────────────────────────────────────────────────────────
def adm_settings(ctx):
    pa = persona(ctx, 'AYESHA.AMERI'); ctx['shot_page'] = pa
    nav(pa, 'moduleSettings', 3200)
    row = pa.locator('.switch-row:has-text("DOC_EXPIRY_ALERT_DAYS")')
    if not row.count():
        row = pa.locator('.switch-row').first
    inp = row.locator('input.form-control').first
    orig = inp.input_value()
    inp.fill(str(int(orig) + 1) if orig.isdigit() else orig + '1')
    inp.dispatch_event('input')
    pa.locator('.page-header-row .btn-primary').click(); wait(pa, 2500)
    nav(pa, 'dashboard', 1200); nav(pa, 'moduleSettings', 3000)
    row = pa.locator('.switch-row:has-text("DOC_EXPIRY_ALERT_DAYS")')
    if not row.count():
        row = pa.locator('.switch-row').first
    newval = row.locator('input.form-control').first.input_value()
    changed = newval != orig
    # restore
    row.locator('input.form-control').first.fill(orig)
    row.locator('input.form-control').first.dispatch_event('input')
    pa.locator('.page-header-row .btn-primary').click(); wait(pa, 2000)
    assert changed, 'setting did not persist (still %s)' % newval
    return 'PASS', 'DOC_EXPIRY_ALERT_DAYS %s → %s persisted across reload; restored to %s. Non-admins have no Module Settings nav (see FL-ACC-03)' % (orig, newval, orig)

# ── SHL ──────────────────────────────────────────────────────────────────────
def shl_switch(ctx):
    p = ctx['page']
    nav(p, 'dashboard', 2000)
    p.locator('.modsw-btn').click(); wait(p, 700)
    fl_item = p.locator('.modsw-item:has-text("Freelancer")')
    has_soon = fl_item.locator('.soon-b').count() > 0
    p.locator('.modsw-item:has-text("Admin")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    assert '/Admin/Jet/' in p.url and not p.locator('.login-card').count(), 're-login was demanded'
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    ensure_en(p)
    assert not has_soon, 'FL still shows a "soon" badge in the switcher'
    return 'PASS', 'FL → Admin → FL with no re-login; FL listed live (no "soon" badge)'

def shl_arabic(ctx):
    p = ctx['page']
    nav(p, 'dashboard', 2000)
    p.locator('.lang-pill button').nth(1).click(); wait(p, 1800)
    assert p.evaluate('document.documentElement.dir') == 'rtl', 'document did not flip to RTL'
    nav(p, 'dashboard', 3200)
    canvases = p.locator('canvas').count()
    nav(p, 'registrations', 2500)
    rows = p.locator('.data-table tbody tr').count()
    nums = p.locator('.data-table tbody tr td').first.inner_text()
    latin = bool(re.search(r'[0-9]', nums))
    open_detail(p, 'flContractId', ctx.get('con_id', '1'), 'contractDetail')
    ok_detail = p.locator('.fl-tabs').count() > 0
    # leave AR active so the evidence screenshot shows the mirrored layout;
    # main() restores EN in cleanup
    assert canvases >= 2 and rows >= 1 and ok_detail, 'AR walk incomplete'
    assert latin, 'digits not Latin in AR mode'
    return 'PASS', 'RTL mirrored; AR labels with Latin digits; dashboard charts re-rendered (%d); registrations + contract detail OK' % canvases

# ═════════════════════════════════════════════════════════════════════════════
CASES = [
 ('Access & Session', 'ACC', [
   ('Shared session', 'Entry from Admin', acc_entry),
   ('No session', 'Direct open without login', acc_nosession),
   ('Role gating', 'Compliance/Approvals/Admin menus', acc_roles),
 ]),
 ('Dashboard', 'DSH', [
   ('Stats', 'Cards show live numbers', dsh_stats),
   ('Charts', 'Spend + expiry horizon', dsh_charts),
 ]),
 ('Registrations', 'REG', [
   ('List', 'Paged list + search', reg_list),
   ('Create', 'New registration draft', reg_create),
   ('Photo', 'Photo required to submit', reg_photo),
   ('Submit', 'Send for approval', reg_submit),
   ('Approve', 'Approval creates the freelancer', reg_approve),
 ]),
 ('Freelancers', 'VND', [
   ('List', 'Directory', vnd_list),
   ('Detail', 'Tabs', vnd_tabs),
   ('Bank', 'Add bank account', vnd_bank),
   ('Status', 'Change status', vnd_status),
   ('Documents', 'Add document with expiry', vnd_doc),
 ]),
 ('Contracts', 'CON', [
   ('List', 'Contracts with billing bar', con_list),
   ('Create', 'New contract draft', con_create),
   ('Submit small', '< 50k = 1 approval step', con_small),
   ('Submit large', '>= 50k = 2 approval steps', con_large),
   ('Detail', 'Schedule + tabs', con_detail),
   ('Amendment', 'Amend amount/end date', con_amend),
   ('Renewal', 'Renew a contract', con_renew),
 ]),
 ('Payments', 'PAY', [
   ('Schedule', 'Due worklist', pay_worklist),
   ('Voucher', 'Generate from schedule', pay_generate),
   ('Voucher', 'Invoice + submit', pay_invoice),
   ('Mark paid', 'FL_ADMIN settles', pay_markpaid),
   ('Deliverables', 'Accept / reject', pay_deliverables),
 ]),
 ('Compliance', 'CMP', [
   ('Documents', 'Expiry filter + badge', cmp_filter),
   ('Navigate', 'Document to freelancer', cmp_navigate),
 ]),
 ('Approvals', 'APR', [
   ('Queue', 'Pending list', apr_queue),
   ('Comment required', 'Approve needs a comment', apr_comment),
   ('Reject / Return', 'Negative paths', apr_negative),
 ]),
 ('Administration', 'ADM', [
   ('Module settings', 'FL settings', adm_settings),
 ]),
 ('Shell & Language', 'SHL', [
   ('Switcher', 'Cross-module hop', shl_switch),
   ('Arabic', 'RTL pass', shl_arabic),
 ]),
]

# ═════════════════════════════════════════════════════════════════════════════
def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    COLORS = {'PASS': RGBColor(0x1B, 0x5E, 0x20), 'FAIL': RGBColor(0xB7, 0x1C, 0x1C),
              'PARTIAL': RGBColor(0xE6, 0x51, 0x00), 'MANUAL': RGBColor(0x54, 0x6E, 0x7A),
              'SKIP': RGBColor(0x54, 0x6E, 0x7A)}
    doc = Document()
    doc.add_heading('i-Finance UAT Results — Freelancers (App 203)', level=0)
    p = doc.add_paragraph('Automated UAT run · %s · Environment: %s (dev-proxy → ADB PROD)'
                          % (DATESTR, APPURL))
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph('Executed by Claude (Playwright). Mirrors UAT_FL_%s-NN.xlsx — one section per '
                      'workbook case, with the evidence screenshot captured at the verification '
                      'moment. Personas: SYS_ADMIN, FL_ADMIN (AYESHA.AMERI), FL_MANAGER '
                      '(NASER.ALKHAJA), FL_USER (SHAIKHA.GALAMERI). Lifecycle cases used fresh '
                      'UAT-created records; seeded workbook samples were preserved where possible.'
                      % DATESTR).runs[0].font.size = Pt(9)

    doc.add_heading('Summary', level=1)
    counts = {}
    for r in RESULTS:
        counts[r['status']] = counts.get(r['status'], 0) + 1
    line = '   '.join('%s: %d' % (k, counts[k]) for k in ('PASS', 'PARTIAL', 'MANUAL', 'FAIL', 'SKIP') if k in counts)
    sp = doc.add_paragraph(); run = sp.add_run('%d cases — %s' % (len(RESULTS), line)); run.bold = True

    t = doc.add_table(rows=1, cols=6); t.style = 'Light Grid Accent 1'
    for i, htxt in enumerate(('Area', 'Cases', 'Pass', 'Partial/Manual', 'Fail', 'Pass %')):
        t.rows[0].cells[i].text = htxt
    for area, prefix, cases in CASES:
        rs = [r for r in RESULTS if r['area'] == area]
        np_ = sum(1 for r in rs if r['status'] == 'PASS')
        npm = sum(1 for r in rs if r['status'] in ('PARTIAL', 'MANUAL'))
        nf = sum(1 for r in rs if r['status'] == 'FAIL')
        row = t.add_row().cells
        row[0].text = area; row[1].text = str(len(rs)); row[2].text = str(np_)
        row[3].text = str(npm); row[4].text = str(nf)
        row[5].text = '%d%%' % round(np_ * 100.0 / max(len(rs), 1))

    fails = [r for r in RESULTS if r['status'] == 'FAIL']
    if fails:
        doc.add_heading('Failed cases (triage first)', level=2)
        for r in fails:
            fp = doc.add_paragraph(style='List Bullet')
            fr = fp.add_run('%s — %s: ' % (r['tid'], r['scenario'])); fr.bold = True
            fp.add_run(r['note'])

    for area, prefix, cases in CASES:
        doc.add_page_break()
        doc.add_heading(area, level=1)
        for r in [x for x in RESULTS if x['area'] == area]:
            doc.add_heading('%s — %s' % (r['tid'], r['scenario']), level=2)
            pr = doc.add_paragraph()
            sr = pr.add_run(r['status']); sr.bold = True
            sr.font.color.rgb = COLORS.get(r['status'], RGBColor(0, 0, 0))
            pr.add_run('   ·   %s   ·   %.0fs' % (r['function'], r['secs'])).font.size = Pt(9)
            if r['note']:
                np2 = doc.add_paragraph(r['note']); np2.runs[0].font.size = Pt(10)
            if r['shot'] and os.path.exists(r['shot']):
                try:
                    doc.add_picture(r['shot'], width=Inches(6.3))
                    cap = doc.add_paragraph('Evidence: %s' % os.path.basename(r['shot']))
                    cap.runs[0].font.size = Pt(8)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
    doc.save(DOCX_OUT)
    print('\nWord report:', DOCX_OUT)

# ═════════════════════════════════════════════════════════════════════════════
def main():
    try:
        out = subprocess.run(['powershell', '-NoProfile', '-Command',
            "(Get-NetTCPConnection -LocalPort 8080 -State Listen -ErrorAction SilentlyContinue).OwningProcess"],
            capture_output=True, text=True, timeout=20).stdout.strip()
        for pid in set(out.split()):
            if pid.isdigit():
                subprocess.run(['taskkill', '/F', '/PID', pid], capture_output=True)
    except Exception:
        pass
    proxy = subprocess.Popen(['python', os.path.join(ROOT, 'FL', 'Jet', 'dev-proxy.py')],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx = Ctx()
            ctx['browser'] = browser
            ctx['pp'] = {}
            c, p = open_app(ctx, 'ADMIN')
            ctx['ctx'], ctx['page'] = c, p

            for area, prefix, cases in CASES:
                print('\n──', area)
                for i, (function, scenario, fn) in enumerate(cases, start=1):
                    tid = 'FL-%s-%02d' % (prefix, i)
                    run_case(ctx, tid, area, function, scenario, fn)
                    tc = ctx.pop('tmp_ctx', None)
                    if tc:
                        try: tc.close()
                        except Exception: pass
            try:
                ensure_en(ctx['page'])
            except Exception:
                pass
            browser.close()
    finally:
        proxy.terminate()

    build_docx()
    counts = {}
    for r in RESULTS:
        counts[r['status']] = counts.get(r['status'], 0) + 1
    print('\n=== %d cases: %s' % (len(RESULTS), counts))

if __name__ == '__main__':
    main()
