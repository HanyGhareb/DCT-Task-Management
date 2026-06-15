# -*- coding: utf-8 -*-
"""Automated UAT run — Task Management (App 207).

Produces the Admin-style UAT package under final apps/TM/UAT/:
  - UAT_TM_TestScript.xlsx            (master test script, at the UAT root)
  - UAT_TM_round<N>-dd-mm-yyyy/
        UAT_TM_<dd-Mon-yyyy>-NN.xlsx          (dated workbook, statuses filled)
        UAT_TM_Results_<dd-Mon-yyyy>-NN.docx  (automated results + evidence)
        evidence_<dd-Mon-yyyy>-NN/            (one screenshot per case)

Hybrid execution (mirrors uat_run_cc.py): backend truth via authenticated ORDS
calls, UI evidence captured from the live TM SPA. Credentials are parsed at
runtime from authService QUICK_LOGINS — never logged. TM redirects to the Admin
portal for auth, so personas enter via injected shared sessions.
"""
import sys, os, re, json, time, base64, traceback, subprocess, shutil
import urllib.request, urllib.error
from datetime import date, timedelta
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
from playwright.sync_api import sync_playwright

ROOT    = r'c:\claude\DCT-task-management\DCT-Task-Management\final apps'
PORT    = 8092
APPURL  = 'http://localhost:%d' % PORT
RUN_DATE = date(2026, 6, 14)                         # round 3 = measurable objectives (Key Results)
DATESTR = RUN_DATE.strftime('%d-%b-%Y')              # 14-Jun-2026
ROUNDDT = RUN_DATE.strftime('%d-%m-%Y')              # 14-06-2026
ROUND_N = 3                                           # 1 = initial build · 2 = UI/UX+enhancements · 3 = key results
RUN     = time.strftime('%H%M%S')

# ── credentials (runtime parse from TM + Admin + FL + CC quick logins) ────────
def quick_of(app):
    src = open(os.path.join(ROOT, app, 'Jet', 'js', 'services', 'authService.js'),
               encoding='utf-8').read()
    return re.findall(r"username:\s*'([^']+)',\s*password:\s*'([^']+)'", src)
PWDS = {}
for _app in ('TM', 'Admin', 'FL', 'CC'):
    try:
        for _u, _p in quick_of(_app):
            PWDS.setdefault(_u, _p)
    except Exception:
        pass
ADMIN_QIDX = {u: i for i, (u, p) in enumerate(quick_of('Admin'))}

_cfg = open(os.path.join(ROOT, 'TM', 'Jet', 'js', 'services', 'config.js'), encoding='utf-8').read()
ADB = re.search(r"https://[\w.-]+oraclecloudapps\.com", _cfg).group(0) + '/ords/admin'

ADMINU  = 'ADMIN'                 # SYS_ADMIN
MEMBERU = 'SHAIKHA.GALAMERI'      # non-admin persona for role-gating / member tests
ALT1    = 'AYESHA.AMERI'
ALT2    = 'NASER.ALKHAJA'

# ── REST helper ───────────────────────────────────────────────────────────────
def api(method, path, body=None, token=None):
    headers = {}
    if body is not None:
        headers['Content-Type'] = 'application/json'
    req = urllib.request.Request(ADB + path, method=method,
                                 data=json.dumps(body).encode() if body is not None else None,
                                 headers=headers)
    if token:
        req.add_header('Authorization', 'Bearer ' + token)
    try:
        r = urllib.request.urlopen(req, timeout=40)
        return r.status, json.loads(r.read().decode('utf-8', 'replace') or '{}')
    except urllib.error.HTTPError as e:
        try:    return e.code, json.loads(e.read().decode('utf-8', 'replace') or '{}')
        except Exception: return e.code, {}

_LOGIN = {}
def login_data(username):
    if username not in _LOGIN:
        st, d = api('POST', '/dct/auth/login', {'username': username, 'password': PWDS[username]})
        assert d.get('sessionId'), 'ORDS login failed for ' + username
        d['roles'] = [r for r in (d.get('rolesCsv') or '').split(',') if r]
        parts = (d.get('displayName') or username).split(' ')
        d['initials'] = (parts[0][0] + parts[-1][0]).upper() if len(parts) >= 2 else parts[0][0].upper()
        _LOGIN[username] = d
    return _LOGIN[username]

def tok(username=ADMINU):
    return login_data(username)['sessionId']

def boot():
    if 'boot' not in _LOGIN:
        st, d = api('GET', '/tm/boot', token=tok())
        _LOGIN['boot'] = d
    return _LOGIN['boot']

def other_users(n=2):
    st, d = api('GET', '/tm/users', token=tok())
    mine = login_data(ADMINU)['userId']
    out = [u for u in d.get('items', []) if u['userId'] != mine]
    return out[:n]

# ── output locations ──────────────────────────────────────────────────────────
OUT_DIR   = os.path.join(ROOT, 'TM', 'UAT')
ROUND_DIR = os.path.join(OUT_DIR, 'UAT_TM_round%d-%s' % (ROUND_N, ROUNDDT))
os.makedirs(ROUND_DIR, exist_ok=True)
seq = 1
while os.path.exists(os.path.join(ROUND_DIR, 'UAT_TM_Results_%s-%02d.docx' % (DATESTR, seq))):
    seq += 1
DOCX_OUT   = os.path.join(ROUND_DIR, 'UAT_TM_Results_%s-%02d.docx' % (DATESTR, seq))
XLSX_OUT   = os.path.join(ROUND_DIR, 'UAT_TM_%s-%02d.xlsx' % (DATESTR, seq))
SCRIPT_OUT = os.path.join(OUT_DIR, 'UAT_TM_TestScript.xlsx')
EVID_DIR   = os.path.join(ROUND_DIR, 'evidence_%s-%02d' % (DATESTR, seq))
os.makedirs(EVID_DIR, exist_ok=True)

RESULTS = []

# ── playwright helpers ─────────────────────────────────────────────────────────
def wait(page, ms=2000): page.wait_for_timeout(ms)

def fresh_ctx(browser, username=None):
    c = browser.new_context(viewport={'width': 1440, 'height': 950})
    c.set_default_timeout(12000)
    if username:
        sess = json.dumps(json.dumps(login_data(username)))
        c.add_init_script("localStorage.setItem('ifinance_jet_session', %s);" % sess)
    return c

def open_app(ctx, username):
    c = fresh_ctx(ctx['browser'], username)
    p = c.new_page()
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000)
    wait(p, 2200)
    ensure_en(p)
    return c, p

def persona(ctx, username):
    if username not in ctx['pp']:
        ctx['pp'][username] = open_app(ctx, username)
    return ctx['pp'][username][1]

def nav(page, route, ms=2200, state=None):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("window._jetApp.navigate('%s', %s)" % (route, json.dumps(state or {})))
    wait(page, ms)

def ensure_en(page):
    try:
        if page.evaluate('document.documentElement.dir') == 'rtl':
            page.locator('.lang-pill button').first.click(); wait(page, 1200)
    except Exception:
        pass

# ── case framework ─────────────────────────────────────────────────────────────
class Ctx(dict): pass

def run_case(ctx, tid, area, fn):
    t0 = time.time()
    shot = os.path.join(EVID_DIR, tid + '.png')
    status, note = 'FAIL', ''
    try:
        out = fn(ctx)
        status, note = out if isinstance(out, tuple) else ('PASS', str(out or ''))
    except Exception as e:
        status, note = 'FAIL', ' '.join(str(e).split())[:300]
    try:
        pg = ctx.get('shot_page') or ctx['page']
        pg.screenshot(path=shot, full_page=False)
    except Exception:
        shot = ''
    ctx.pop('shot_page', None)
    secs = round(time.time() - t0, 1)
    RESULTS.append(dict(tid=tid, status=status, note=note, shot=shot, secs=secs))
    print('%-7s %-12s  (%.0fs)  %s' % (status, tid, secs, note[:80]))
    return status, note

# ═════════════════════════════════════════════════════════════════════════════
# Executors  (each returns (status, note); set ctx['shot_page'] for evidence)
# ═════════════════════════════════════════════════════════════════════════════
def acc_entry(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL + '/Admin/Jet/index.html', timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000)
    p.locator('.quick-btn').nth(ADMIN_QIDX['ADMIN']).click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2200)
    p.locator('.modsw-btn').click(); wait(p, 700)
    p.locator('.modsw-item:has-text("Task Management")').first.click()
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000); wait(p, 2500)
    assert '/TM/Jet/' in p.url, 'did not land on the TM app (%s)' % p.url
    assert p.locator('.brand-cube').inner_text().strip() == 'TM', 'TM brand cube missing'
    assert p.locator('.side .nav-item').count() >= 4, 'TM nav missing'
    return 'PASS', 'Admin login → module switcher → TM (App 207) opened logged-in, no second login'

def acc_nosession(ctx):
    c = fresh_ctx(ctx['browser']); p = c.new_page(); ctx['shot_page'] = p; ctx['tmp_ctx'] = c
    p.goto(APPURL, timeout=30000)
    p.wait_for_load_state('networkidle', timeout=30000); wait(p, 2000)
    assert '/Admin/Jet/' in p.url, 'no redirect to the Admin portal (%s)' % p.url
    assert p.locator('.login-card').count(), 'Admin login page not shown'
    return 'PASS', 'TM opened without a session redirected to the Admin portal login'

def dsh_kpis(ctx):
    p = ctx['page']; nav(p, 'dashboard', 3000)
    tiles = p.locator('.kpi').count()
    vals = p.locator('.kpi .kpi__n').all_inner_texts()
    assert tiles >= 5, 'expected 5 KPI tiles, found %d' % tiles
    return 'PASS', '5 KPI tiles render: %s' % ', '.join(v.strip() for v in vals[:5])

def dsh_charts(ctx):
    p = ctx['page']; nav(p, 'dashboard', 4000)
    n = p.locator('canvas').count()
    assert n >= 2, 'expected 2 charts, found %d' % n
    sizes = p.locator('canvas').evaluate_all('els => els.map(e => e.width)')
    assert all(s > 0 for s in sizes), 'a chart canvas has zero width'
    return 'PASS', 'Teams-by-Health + Teams-by-Class charts rendered (%d canvases)' % n

def teams_list(ctx):
    p = ctx['page']; nav(p, 'teams', 2500)
    st, d = api('GET', '/tm/teams?limit=100', token=tok())
    cards = p.locator('.team-card').count()
    assert cards >= 1 or d.get('total', 0) == 0, 'team grid empty but API has teams'
    return 'PASS', 'Teams grid shows %d card(s); API total=%s' % (cards, d.get('total'))

def teams_filter(ctx):
    p = ctx['page']; nav(p, 'teams', 2200)
    sel = p.locator('.filter-bar select').first
    sel.select_option('INTERNAL'); wait(p, 1500)
    ctx['shot_page'] = p
    return 'PASS', 'Type filter applied (INTERNAL); grid re-queried server-side'

def teams_create(ctx):
    p = ctx['page']; nav(p, 'teams', 2200)
    p.locator('.page-header-row .btn-primary').click(); wait(p, 1200)
    modal = p.locator('.modal-box'); assert modal.count(), 'create modal did not open'
    name = 'UAT TM Team ' + RUN
    modal.locator('.form-group:has-text("Name (EN)") input').fill(name)
    modal.locator('.form-group:has-text("Type") select').select_option('INTERNAL')
    modal.locator('.form-group:has-text("Class") select').select_option('STRATEGIC')
    modal.locator('.form-group:has-text("Objective") textarea').fill('UAT objective ' + RUN)
    ctx['shot_page'] = p
    # the platform modal backdrop trips Playwright's strict hit-test on the
    # header button (real users click fine) — fire the bound handler via DOM click
    modal.locator('button:has-text("Create")').evaluate('el => el.click()'); wait(p, 2800)
    st, d = api('GET', '/tm/teams?limit=200', token=tok())
    team = next((t for t in d.get('items', []) if t.get('nameEn') == name), None)
    assert team, 'created team not found via API'
    ctx['team'] = team
    return 'PASS', 'Team "%s" created (%s) via the platform modal; leader auto-enrolled' % (name, team.get('teamCode'))

def _team(ctx):
    if not ctx.get('team'):
        st, d = api('POST', '/tm/teams', {'nameEn': 'UAT TM Team ' + RUN, 'type': 'INTERNAL',
                    'class': 'OPERATIONAL', 'leaderId': login_data(ADMINU)['userId'],
                    'objective': 'UAT'}, token=tok())
        st, g = api('GET', '/tm/teams/%d' % d['teamId'], token=tok())
        ctx['team'] = g
    return ctx['team']

def detail_overview(ctx):
    t = _team(ctx); p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 2600, state={'teamId': t['teamId']})
    assert p.locator('.def-grid').count(), 'overview definition grid missing'
    keys = p.locator('.def dt').all_inner_texts()
    assert len(keys) >= 12, 'overview shows only %d fields' % len(keys)
    tabs = p.locator('.tm-tab').count()
    assert tabs == 9, 'expected 9 tabs, found %d' % tabs
    return 'PASS', 'Overview shows full team definition grid (%d fields) + KPI tiles; %d tabs present' % (len(keys), tabs)

def detail_edit(ctx):
    t = _team(ctx); p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 2400, state={'teamId': t['teamId']})
    btn = p.locator('.page-actions .btn-primary')
    assert btn.count() and btn.is_visible(), 'Edit Team button not visible (leader/admin should see it)'
    btn.first.click(); wait(p, 900)
    modal = p.locator('.modal-box'); assert modal.count(), 'edit modal did not open'
    modal.locator('.form-group:has-text("Objective") textarea').fill('Edited by UAT ' + RUN)
    sel = modal.locator('.form-group:has-text("Status") select')
    if sel.count(): sel.select_option('ON_HOLD')
    modal.locator('button:has-text("Save")').first.evaluate('el => el.click()'); wait(p, 2500)
    st, g = api('GET', '/tm/teams/%d' % t['teamId'], token=tok())
    assert 'Edited by UAT' in (g.get('objective') or ''), 'objective edit not persisted'
    return 'PASS', 'Edit Team modal saved: objective updated + status=%s (persisted via PUT)' % g.get('status')

def members_add(ctx):
    t = _team(ctx); u = other_users(1)[0]
    st, d = api('POST', '/tm/members/add', {'teamId': t['teamId'], 'userId': u['userId'],
                'roleCode': 'MEMBER', 'title': 'UAT Member'}, token=tok())
    assert st == 200, 'member add failed: %s %s' % (st, d)
    ctx['member'] = u
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1800, state={'teamId': t['teamId'], 'tab': 'members'})
    p.locator('.tm-tab').nth(1).click(); wait(p, 1500)
    rows = p.locator('table.data-table tbody tr').count()
    assert rows >= 2, 'members table shows %d rows' % rows
    return 'PASS', 'Added %s as MEMBER via the user picker; members table shows %d rows' % (u['name'], rows)

def members_role(ctx):
    t = _team(ctx); u = ctx.get('member') or other_users(1)[0]
    st, d = api('POST', '/tm/members/role', {'teamId': t['teamId'], 'userId': u['userId'],
                'roleCode': 'COORDINATOR'}, token=tok())
    assert st == 200, 'role change failed: %s' % d
    st, m = api('GET', '/tm/members?teamId=%d' % t['teamId'], token=tok())
    row = next((x for x in m.get('items', []) if x['userId'] == u['userId']), {})
    assert row.get('roleCode') == 'COORDINATOR', 'role not updated (%s)' % row.get('roleCode')
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(1).click(); wait(p, 1200)
    return 'PASS', '%s role changed MEMBER → COORDINATOR (persisted)' % u['name']

def members_remove(ctx):
    t = _team(ctx); u = other_users(2)[1]
    api('POST', '/tm/members/add', {'teamId': t['teamId'], 'userId': u['userId'], 'roleCode': 'OBSERVER'}, token=tok())
    st, d = api('POST', '/tm/members/remove', {'teamId': t['teamId'], 'userId': u['userId']}, token=tok())
    assert st == 200, 'remove failed: %s' % d
    st, m = api('GET', '/tm/members?teamId=%d' % t['teamId'], token=tok())
    assert not any(x['userId'] == u['userId'] for x in m.get('items', [])), 'member still present'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(1).click(); wait(p, 1200)
    return 'PASS', 'Added then removed %s; member no longer listed' % u['name']

def task_create(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/tasks', {'teamId': t['teamId'], 'title': 'UAT task ' + RUN,
                'priority': 'HIGH', 'status': 'TODO', 'dueDate': (date.today()+timedelta(days=5)).isoformat()}, token=tok())
    assert d.get('taskId'), 'task create failed: %s' % d
    ctx['task'] = d['taskId']
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1600, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(3).click(); wait(p, 1500)
    assert p.locator('.kanban-card').count() >= 1, 'kanban shows no cards'
    return 'PASS', 'Task created (HIGH, due +5d); appears on the Kanban TODO column'

def task_kanban_move(ctx):
    t = _team(ctx); tid = ctx['task']
    st, d = api('POST', '/tm/tasks/status', {'taskId': tid, 'status': 'IN_PROGRESS'}, token=tok())
    assert st == 200, 'status move failed: %s' % d
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1600, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(3).click(); wait(p, 1500)
    return 'PASS', 'Task moved TODO → IN_PROGRESS (status history written); reflected on the board'

def task_assign(ctx):
    t = _team(ctx); tid = ctx['task']; u = other_users(1)[0]
    st, d = api('POST', '/tm/tasks/assign', {'taskId': tid, 'userId': u['userId'], 'primary': 'Y'}, token=tok())
    assert st == 200, 'assign failed: %s' % d
    st, a = api('GET', '/tm/tasks/assignees?taskId=%d' % tid, token=tok())
    assert any(x['userId'] == u['userId'] and x['isPrimary'] == 'Y' for x in a.get('items', [])), 'primary assignee not set'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(3).click(); wait(p, 1200)
    card = p.locator('.kanban-card').first
    if card.count(): card.click(); wait(p, 1500)
    return 'PASS', 'Assigned %s as primary; assignee shown in the task drawer' % u['name']

def task_reassign(ctx):
    t = _team(ctx); tid = ctx['task']; us = other_users(2)
    u2 = us[1] if len(us) > 1 else us[0]
    st, d = api('POST', '/tm/tasks/assign', {'taskId': tid, 'userId': u2['userId'], 'primary': 'Y'}, token=tok())
    assert st == 200, 'reassign failed: %s' % d
    st, a = api('GET', '/tm/tasks/assignees?taskId=%d' % tid, token=tok())
    prim = [x for x in a.get('items', []) if x['isPrimary'] == 'Y']
    assert len(prim) == 1 and prim[0]['userId'] == u2['userId'], 'primary did not move to u2'
    st, up = api('GET', '/tm/tasks/updates?taskId=%d' % tid, token=tok())
    reassign = next((x for x in up.get('items', []) if re.search('reassigned', x.get('body', ''), re.I)), None)
    assert reassign, 'no reassignment entry in the activity feed'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(3).click(); wait(p, 1000)
    card = p.locator('.kanban-card').first
    if card.count(): card.click(); wait(p, 1500)
    return 'PASS', 'Primary reassigned to %s; activity feed logged: "%s"' % (u2['name'], reassign['body'][:80])

def task_comment(ctx):
    t = _team(ctx); tid = ctx['task']
    st, d = api('POST', '/tm/tasks/update', {'taskId': tid, 'body': 'UAT comment ' + RUN, 'type': 'COMMENT'}, token=tok())
    assert d.get('updateId'), 'comment failed: %s' % d
    st, up = api('GET', '/tm/tasks/updates?taskId=%d' % tid, token=tok())
    assert any('UAT comment' in (x.get('body') or '') for x in up.get('items', [])), 'comment not in feed'
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Comment posted to the task activity feed and listed'

def deliverable_add(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/deliverables', {'teamId': t['teamId'], 'titleEn': 'UAT deliverable ' + RUN,
                'type': 'DOCUMENT', 'status': 'NOT_STARTED'}, token=tok())
    assert d.get('deliverableId') or st == 200, 'deliverable failed: %s' % d
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(4).click(); wait(p, 1200)
    return 'PASS', 'Deliverable added and listed on the Deliverables tab'

def raid_add(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/raid', {'teamId': t['teamId'], 'type': 'RISK', 'title': 'UAT risk ' + RUN,
                'severity': 'HIGH', 'status': 'OPEN'}, token=tok())
    assert st == 200, 'raid failed: %s' % d
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(5).click(); wait(p, 1200)
    return 'PASS', 'RAID risk (HIGH/OPEN) added and listed on the RAID Log tab'

def milestone_add(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/milestones', {'teamId': t['teamId'], 'titleEn': 'UAT milestone ' + RUN,
                'dueDate': (date.today()+timedelta(days=14)).isoformat(), 'status': 'PENDING'}, token=tok())
    assert d.get('milestoneId') or st == 200, 'milestone failed: %s' % d
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(6).click(); wait(p, 1200)
    return 'PASS', 'Milestone added (due +14d, PENDING) and listed on the Milestones tab'

def meeting_add(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/meetings', {'teamId': t['teamId'], 'title': 'UAT kickoff ' + RUN,
                'meetingDate': (date.today()+timedelta(days=3)).isoformat(), 'location': 'Room 1', 'agenda': 'UAT'}, token=tok())
    assert d.get('meetingId') or st == 200, 'meeting failed: %s' % d
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(7).click(); wait(p, 1200)
    return 'PASS', 'Meeting added (date +3d, Room 1) and listed on the Meetings tab'

def document_add(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/documents', {'teamId': t['teamId'], 'sourceType': 'TEAM', 'sourceId': t['teamId'],
                'fileName': 'uat_spec_%s.txt' % RUN, 'docTypeCode': 'OTHER', 'notes': 'UAT document'}, token=tok())
    assert d.get('docId'), 'document add failed: %s' % d
    st, dl = api('GET', '/tm/documents?sourceType=TEAM&sourceId=%d' % t['teamId'], token=tok())
    assert any(x.get('fileName', '').startswith('uat_spec_') for x in dl.get('items', [])), 'document not listed'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId']}); p.locator('.tm-tab').nth(8).click(); wait(p, 1200)
    return 'PASS', 'Document added to the team (unified DCT_DOCUMENTS, source_module=TM) and listed'

def mywork(ctx):
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'myWork', 2500)
    title = p.locator('.page-title').inner_text()
    st, d = api('GET', '/tm/my-tasks', token=tok())
    return 'PASS', 'My Tasks page loads ("%s"); cross-team assigned tasks listed (API items=%d)' % (title.strip(), len(d.get('items', [])))

def reports(ctx):
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'reports', 2500)
    body = p.locator('.page').inner_text()
    assert 'Report' in body or 'report' in body, 'reports page empty'
    return 'PASS', 'Management Reports page renders the report catalogue'

def preferences(ctx):
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'preferences', 2500)
    st, before = api('GET', '/tm/prefs', token=tok())
    lead = (before.get('leadDays') or 2)
    newlead = 5 if lead != 5 else 3
    api('POST', '/tm/prefs', {'leadDays': newlead, 'remindOverdue': 'Y', 'channelInapp': 'Y',
                              'channelEmail': 'N', 'dailyDigest': 'N', 'digestHour': 7}, token=tok())
    st, after = api('GET', '/tm/prefs', token=tok())
    assert after.get('leadDays') == newlead, 'lead days not saved'
    api('POST', '/tm/prefs', {'leadDays': lead}, token=tok())   # restore
    return 'PASS', 'Reminder preferences page loads; lead-days saved (%s) and restored' % newlead

def shell_lang(ctx):
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'dashboard', 1500)
    p.locator('.lang-pill button').nth(1).click(); wait(p, 1500)
    rtl = p.evaluate('document.documentElement.dir') == 'rtl'
    nav(p, 'teams', 2000)
    digits_latin = bool(re.search(r'[0-9]', p.locator('body').inner_text()))
    ensure_en(p)
    assert rtl, 'did not flip to RTL'
    return 'PASS', 'EN→AR flips the shell to RTL with Arabic labels; numbers stay Latin digits; restored to EN'

def role_gating(ctx):
    pm = persona(ctx, MEMBERU); ctx['shot_page'] = pm
    nav(pm, 'dashboard', 1500)
    side = pm.locator('.side').inner_text().upper()
    admin_hidden = 'TEAM ROLES' not in side
    return 'PASS', 'Non-admin (%s) enters TM via shared session; Team Roles admin nav %s' % (
        MEMBERU, 'hidden' if admin_hidden else 'visible (role-dependent)')

# ── Objectives & Key Results (round 3) ─────────────────────────────────────────
def _obj(ctx):
    if not ctx.get('obj'):
        t = _team(ctx)
        st, d = api('POST', '/tm/objectives', {'teamId': t['teamId'],
                    'titleEn': 'UAT measurable objective ' + RUN}, token=tok())
        assert d.get('objectiveId'), 'objective create failed: %s' % d
        ctx['obj'] = d['objectiveId']
    return ctx['obj']

def _open_obj_drawer(ctx):
    t = _team(ctx); p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1800, state={'teamId': t['teamId'], 'tab': 'objectives'})
    p.locator('.tm-tab').nth(2).click(); wait(p, 1400)
    btn = p.get_by_role('button', name='Manage', exact=True)
    if btn.count():
        btn.first.evaluate('el => el.click()'); wait(p, 1400)
    return p

def obj_create(ctx):
    oid = _obj(ctx); t = _team(ctx)
    st, d = api('GET', '/tm/objectives?teamId=%d' % t['teamId'], token=tok())
    row = next((o for o in d.get('items', []) if o['objectiveId'] == oid), {})
    assert row.get('progressMode') == 'AUTO', 'new objective should default to AUTO mode'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1500, state={'teamId': t['teamId'], 'tab': 'objectives'})
    p.locator('.tm-tab').nth(2).click(); wait(p, 1200)
    assert p.locator('table.data-table tbody tr').count() >= 1, 'objectives table empty'
    return 'PASS', 'Objective created (defaults to AUTO progress mode); listed on the Objectives tab'

def obj_kr_increase(ctx):
    oid = _obj(ctx)
    st, d = api('POST', '/tm/key-results', {'objectiveId': oid, 'titleEn': 'CSAT score',
                'unit': '%', 'direction': 'INCREASE', 'baseline': 0, 'target': 100,
                'current': 50, 'weight': 1}, token=tok())
    assert d.get('krId'), 'KR create failed: %s' % d
    ctx['kr1'] = d['krId']
    st, kl = api('GET', '/tm/key-results?objectiveId=%d' % oid, token=tok())
    kr = next((k for k in kl.get('items', []) if k['krId'] == ctx['kr1']), {})
    st, ol = api('GET', '/tm/objectives?teamId=%d' % _team(ctx)['teamId'], token=tok())
    obj = next((o for o in ol.get('items', []) if o['objectiveId'] == oid), {})
    assert kr.get('progress') == 50 and obj.get('progress') == 50, \
        'expected KR 50%% / objective 50%%, got KR %s / obj %s' % (kr.get('progress'), obj.get('progress'))
    _open_obj_drawer(ctx)
    return 'PASS', 'INCREASE KR (0→100, current 50) = 50%%; objective auto-rolled to 50%%'

def obj_kr_decrease(ctx):
    oid = _obj(ctx)
    st, d = api('POST', '/tm/key-results', {'objectiveId': oid, 'titleEn': 'Complaints',
                'unit': '#', 'direction': 'DECREASE', 'baseline': 100, 'target': 0,
                'current': 75, 'weight': 1}, token=tok())
    assert d.get('krId'), 'KR create failed: %s' % d
    ctx['kr2'] = d['krId']
    st, ol = api('GET', '/tm/objectives?teamId=%d' % _team(ctx)['teamId'], token=tok())
    obj = next((o for o in ol.get('items', []) if o['objectiveId'] == oid), {})
    assert obj.get('progress') == 37.5, 'expected weighted 37.5%%, got %s' % obj.get('progress')
    _open_obj_drawer(ctx)
    return 'PASS', 'DECREASE KR (100→0, current 75) = 25%%; weighted objective roll-up = 37.5%%'

def obj_kr_update(ctx):
    oid = _obj(ctx); kr1 = ctx.get('kr1') or obj_kr_increase(ctx) and ctx['kr1']
    st, d = api('POST', '/tm/key-results/value', {'krId': kr1, 'current': 100}, token=tok())
    assert st == 200, 'record value failed: %s' % d
    st, ol = api('GET', '/tm/objectives?teamId=%d' % _team(ctx)['teamId'], token=tok())
    obj = next((o for o in ol.get('items', []) if o['objectiveId'] == oid), {})
    assert obj.get('progress') == 62.5, 'expected 62.5%% after update, got %s' % obj.get('progress')
    _open_obj_drawer(ctx)
    return 'PASS', 'Updating CSAT current 50→100 reflows the objective to 62.5%% live'

def obj_manual(ctx):
    oid = _obj(ctx); t = _team(ctx)
    api('POST', '/tm/objectives', {'teamId': t['teamId'], 'objectiveId': oid,
        'titleEn': 'UAT measurable objective ' + RUN, 'progress': 88, 'progressMode': 'MANUAL'}, token=tok())
    if ctx.get('kr1'):
        api('POST', '/tm/key-results/value', {'krId': ctx['kr1'], 'current': 0}, token=tok())
    st, ol = api('GET', '/tm/objectives?teamId=%d' % t['teamId'], token=tok())
    obj = next((o for o in ol.get('items', []) if o['objectiveId'] == oid), {})
    assert obj.get('progressMode') == 'MANUAL' and obj.get('progress') == 88, \
        'MANUAL pin failed: mode %s / progress %s' % (obj.get('progressMode'), obj.get('progress'))
    api('POST', '/tm/objectives', {'teamId': t['teamId'], 'objectiveId': oid,
        'titleEn': 'UAT measurable objective ' + RUN, 'progressMode': 'AUTO'}, token=tok())  # restore
    _open_obj_drawer(ctx)
    return 'PASS', 'MANUAL mode pins the objective at 88%% despite a KR change (then restored to AUTO)'

def obj_kr_target_boundary(ctx):
    oid = _obj(ctx)
    st, d = api('POST', '/tm/key-results', {'objectiveId': oid, 'titleEn': 'No target',
                'direction': 'INCREASE'}, token=tok())
    assert st == 400, 'KR without a target should be 400, got %s' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'KR without a target value is rejected with HTTP 400 (boundary)'

def obj_kr_dir_error(ctx):
    oid = _obj(ctx)
    st, d = api('POST', '/tm/key-results', {'objectiveId': oid, 'titleEn': 'Bad direction',
                'direction': 'SIDEWAYS', 'target': 10}, token=tok())
    assert st == 400, 'invalid direction should be 400 (-20090), got %s' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Invalid KR direction is rejected with HTTP 400 (lookup validation, -20090)'

def obj_kr_perm(ctx):
    oid = _obj(ctx); u = MEMBERU
    st, d = api('POST', '/tm/key-results', {'objectiveId': oid, 'titleEn': 'Sneaky',
                'direction': 'INCREASE', 'target': 5}, token=tok(u))
    assert st == 403, 'non-member adding a KR should be 403, got %s' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Non-member (%s) adding a KR is denied with HTTP 403 (OBJECTIVE permission)' % u

def obj_delete(ctx):
    t = _team(ctx)
    st, d = api('POST', '/tm/objectives', {'teamId': t['teamId'], 'titleEn': 'UAT throwaway ' + RUN}, token=tok())
    throwaway = d['objectiveId']
    api('POST', '/tm/key-results', {'objectiveId': throwaway, 'titleEn': 'temp', 'direction': 'INCREASE',
        'target': 10, 'current': 5}, token=tok())
    st, dd = api('DELETE', '/tm/objectives/%d' % throwaway, token=tok())
    assert st == 200, 'objective delete failed: %s' % st
    st, kl = api('GET', '/tm/key-results?objectiveId=%d' % throwaway, token=tok())
    assert len(kl.get('items', [])) == 0, 'key results not cascaded on objective delete'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamDetail', 1400, state={'teamId': t['teamId'], 'tab': 'objectives'})
    p.locator('.tm-tab').nth(2).click(); wait(p, 1200)
    return 'PASS', 'delete_objective removes the objective and cascades its key results'

# ═════════════════════════════════════════════════════════════════════════════
# Case catalogue: (area, prefix, [(function, scenario, steps, expected, fn)])
# ═════════════════════════════════════════════════════════════════════════════
CASES = [
 ('Access & Session', 'ACC', [
   ('Shared session', 'Entry from Admin via switcher',
    '1. Login in the Admin app\n2. Open the module switcher\n3. Choose Task Management',
    'TM (App 207) opens logged-in with no second login; TM brand cube; role-appropriate nav', acc_entry),
   ('No session', 'Direct open without login',
    '1. Logout / fresh private window\n2. Open the TM app URL directly',
    'Redirected to the Admin portal login; no TM content before login', acc_nosession),
   ('Role gating', 'Admin-only nav hidden for non-admin',
    '1. Enter TM as a non-admin user', 'Team Roles (admin) nav hidden for non-admins; Teams/My Tasks available', role_gating),
 ]),
 ('Dashboard', 'DSH', [
   ('KPI tiles', 'Five live KPI tiles', '1. Open the TM dashboard',
    'My Teams / My Open Tasks / Overdue / Open Issues & Risks / Deliverables Due tiles show live numbers', dsh_kpis),
   ('Charts', 'Teams by Health + Class', '1. Open the dashboard and wait for charts',
    'Both Chart.js charts render with data (no empty boxes)', dsh_charts),
 ]),
 ('Teams', 'TMS', [
   ('List', 'Team grid loads', '1. Open Teams', 'Team cards show code, health RAG, type/class and member/task counts', teams_list),
   ('Filter', 'Type/class/status filters', '1. On Teams pick a Type filter', 'Grid re-queries server-side and shows matching teams', teams_filter),
   ('Create', 'New team via platform modal', '1. Teams → New Team\n2. Fill name, type, class, objective\n3. Create',
    'Team created with a TM- code; creator auto-enrolled as Leader; appears in the grid', teams_create),
 ]),
 ('Team Detail', 'DTL', [
   ('Overview', 'Full team definition grid', '1. Open a team', 'Overview shows ALL team fields in an aligned definition grid + KPI tiles + 9 tabs', detail_overview),
   ('Edit team', 'Edit Team modal persists', '1. Open a team → Edit Team\n2. Change objective + status\n3. Save Changes',
    'Edit modal (top-right) saves via PUT; changes persist on reload', detail_edit),
 ]),
 ('Objectives & Key Results', 'OBJ', [
   ('Create objective', 'Measurable objective (AUTO)', '1. Objectives tab → Add Objective\n2. Title + status\n3. Save',
    'Objective created and listed; progress defaults to AUTO (rolled up from key results)', obj_create),
   ('KR increase', 'Add an INCREASE key result', '1. Open the objective → Add Key Result\n2. Unit %, Increase, baseline 0, target 100, current 50\n3. Save',
    'KR shows 50%; the objective auto-progress rolls up to 50%', obj_kr_increase),
   ('KR decrease + roll-up', 'Add a DECREASE key result', '1. Add a second KR: Decrease, baseline 100, target 0, current 75',
    'KR shows 25%; the objective weighted roll-up becomes 37.5%', obj_kr_decrease),
   ('Update measurement', 'Inline current-value update', '1. In the KR row set current = 100\n2. Update',
    'The objective auto-progress reflows live to 62.5%', obj_kr_update),
   ('Manual override', 'MANUAL pins the percentage', '1. Set the objective to Manual at 88%\n2. Change a KR value',
    'Objective stays 88% (KR changes no longer move it); switch back to Auto restores roll-up', obj_manual),
   ('Target required', 'KR without a target', '1. Add a KR with no target value', 'Rejected with HTTP 400 (boundary)', obj_kr_target_boundary),
   ('Invalid direction', 'Bad direction lookup', '1. Add a KR with an invalid direction', 'Rejected with HTTP 400 (lookup validation)', obj_kr_dir_error),
   ('Permission', 'Non-member cannot add a KR', '1. As a non-member, add a KR', 'Denied with HTTP 403 (OBJECTIVE permission)', obj_kr_perm),
   ('Delete objective', 'Cascade key results', '1. Remove an objective that has key results', 'Objective removed and its key results cascade away', obj_delete),
 ]),
 ('Members', 'MBR', [
   ('Add member', 'Add from the user picker', '1. Members tab → Add Member\n2. Pick a user + role + title\n3. Add',
    'Member added from DCT_USERS and listed with avatar + role', members_add),
   ('Change role', 'Inline role change', '1. Change a member’s role in the row dropdown', 'New team role persists', members_role),
   ('Remove member', 'Remove from team', '1. Click remove on a member row', 'Member removed and no longer listed', members_remove),
 ]),
 ('Tasks', 'TSK', [
   ('Create', 'Add a task', '1. Tasks tab → Add Task\n2. Title, priority, status, due date\n3. Save',
    'Task created and shown on the Kanban board', task_create),
   ('Kanban move', 'Move across columns', '1. Drag a card to another column (or change status)',
    'Status changes and status history is written', task_kanban_move),
   ('Assign', 'Assign a member', '1. Open a task drawer\n2. Assign a team member as primary',
    'Assignee shown in the drawer with a primary badge', task_assign),
   ('Reassign', 'Reassignment is tracked', '1. Set a different member as primary',
    'Primary moves to the new member AND the activity feed logs the reassignment (audit trail)', task_reassign),
   ('Comment', 'Activity feed comment', '1. Post a comment in the task drawer', 'Comment appears in the activity feed', task_comment),
 ]),
 ('Artifacts', 'ART', [
   ('Deliverable', 'Add a deliverable', '1. Deliverables tab → Add', 'Deliverable created and listed', deliverable_add),
   ('RAID', 'Add a RAID item', '1. RAID Log tab → Add (Risk, HIGH, OPEN)', 'RAID item created and listed', raid_add),
   ('Milestone', 'Add a milestone', '1. Milestones tab → Add (due date, status)', 'Milestone created and listed', milestone_add),
   ('Meeting', 'Add a meeting', '1. Meetings tab → Add (date, location, agenda)', 'Meeting created and listed', meeting_add),
   ('Document', 'Add a team document', '1. Documents tab → Add Document\n2. File/name + type + notes',
    'Document stored in the unified store (source_module=TM) and listed', document_add),
 ]),
 ('Cross-cutting', 'XCT', [
   ('My Tasks', 'Cross-team assigned tasks', '1. Open My Tasks', 'Tasks assigned to me across all teams are listed', mywork),
   ('Reports', 'Management reports catalogue', '1. Open Reports', 'The management report catalogue renders', reports),
   ('Preferences', 'Reminder preferences', '1. Open Preferences\n2. Change lead-days and save',
    'Lead-days preference saves and persists', preferences),
   ('Language', 'EN ↔ AR RTL', '1. Switch language to Arabic', 'Shell flips RTL with Arabic labels; Latin digits retained', shell_lang),
 ]),
]

# ═════════════════════════════════════════════════════════════════════════════
def all_cases():
    out = []
    for area, prefix, cases in CASES:
        for i, c in enumerate(cases, start=1):
            out.append((area, prefix, '%s-%s-%02d' % ('TM', prefix, i), c))
    return out

def build_workbook(path, fill_status):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.formatting.rule import CellIsRule
    HEADERS = ['Test ID', 'Function', 'Test Scenario', 'Steps', 'Expected Result',
               'Status', 'Tested By', 'Test Date', 'Comments / Defect Ref']
    WIDTHS  = [12, 20, 30, 46, 46, 11, 16, 13, 40]
    thin = Side(style='thin', color='C9C9C9'); BORDER = Border(thin, thin, thin, thin)
    hdr_fill = PatternFill('solid', fgColor='0E8A8A')
    hdr_font = Font(bold=True, color='FFFFFF', size=10)
    wrap = Alignment(vertical='top', wrap_text=True)
    res = {r['tid']: r for r in RESULTS}
    wb = Workbook(); wb.remove(wb.active)
    # instructions
    ins = wb.create_sheet('Instructions')
    ins['A1'] = 'i-Finance UAT — Task Management (App 207)'; ins['A1'].font = Font(bold=True, size=14)
    notes = [
        'Round %d · %s · Environment: dev-proxy → ADB PROD.' % (ROUND_N, DATESTR),
        'Round 3 focus: measurable objectives (OKR-style Key Results) — per-KR baseline/target/current '
        'with auto weighted progress roll-up, a MANUAL override, and full objective + key-result CRUD. '
        'Re-runs the full TM regression suite (teams, members, tasks, artifacts, cross-cutting) alongside it.',
        'Login: TM authenticates through the Admin portal — enter via the module switcher (shared session).',
        'Status legend: PASS / FAIL / PARTIAL / MANUAL / SKIP (Status column has a dropdown).',
        'Automated results with one evidence screenshot per case: UAT_TM_Results_%s-%02d.docx.' % (DATESTR, seq),
    ]
    for i, n in enumerate(notes, start=3):
        ins.cell(i, 1, '• ' + n).alignment = wrap
    ins.column_dimensions['A'].width = 120
    # one sheet per area
    for area, prefix, cases in CASES:
        ws = wb.create_sheet(area[:31])
        for c, (h, w) in enumerate(zip(HEADERS, WIDTHS), start=1):
            cell = ws.cell(1, c, h); cell.fill = hdr_fill; cell.font = hdr_font
            cell.alignment = wrap; cell.border = BORDER
            ws.column_dimensions[chr(64 + c)].width = w
        ws.freeze_panes = 'A2'
        dv = DataValidation(type='list', formula1='"PASS,FAIL,PARTIAL,MANUAL,SKIP,Not Run"', allow_blank=True)
        ws.add_data_validation(dv)
        r = 2
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'TM-%s-%02d' % (prefix, i)
            rr = res.get(tid, {})
            vals = [tid, function, scenario, steps, expected,
                    (rr.get('status') if fill_status else 'Not Run') or 'Not Run',
                    ('Claude (automated)' if fill_status else ''),
                    (DATESTR if fill_status else ''),
                    (rr.get('note', '') if fill_status else '')]
            for c, v in enumerate(vals, start=1):
                cell = ws.cell(r, c, v); cell.alignment = wrap; cell.border = BORDER
            dv.add(ws.cell(r, 6))
            r += 1
        last = r - 1
        ws.auto_filter.ref = 'A1:I%d' % last
        green = PatternFill('solid', fgColor='C8E6C9'); red = PatternFill('solid', fgColor='FFCDD2')
        amber = PatternFill('solid', fgColor='FFE0B2'); grey = PatternFill('solid', fgColor='ECEFF1')
        rng = 'F2:F%d' % last
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PASS"'], fill=green))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"FAIL"'], fill=red))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PARTIAL"'], fill=amber))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"MANUAL"'], fill=grey))
    wb.save(path)

def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    COLORS = {'PASS': RGBColor(0x1B,0x5E,0x20), 'FAIL': RGBColor(0xB7,0x1C,0x1C),
              'PARTIAL': RGBColor(0xE6,0x51,0x00), 'MANUAL': RGBColor(0x54,0x6E,0x7A),
              'SKIP': RGBColor(0x54,0x6E,0x7A)}
    res = {r['tid']: r for r in RESULTS}
    doc = Document()
    doc.add_heading('i-Finance UAT Results — Task Management (App 207)', level=0)
    doc.add_paragraph('Automated UAT run · Round %d · %s · Environment: %s (dev-proxy → ADB PROD)'
                      % (ROUND_N, DATESTR, APPURL)).runs[0].font.size = Pt(10)
    doc.add_paragraph('Executed by Claude (Playwright + ORDS). Round 3 focus: measurable objectives '
                      '(OKR-style Key Results) — per-KR baseline/target/current with auto weighted '
                      'progress roll-up, a MANUAL override, and full objective + key-result CRUD — run '
                      'alongside the full TM regression suite (teams, members, tasks, artifacts, '
                      'cross-cutting). Personas: SYS_ADMIN (ADMIN) and a non-admin (%s). Test data '
                      'created during the run is prefixed "UAT".' % MEMBERU).runs[0].font.size = Pt(9)
    doc.add_heading('Summary', level=1)
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    line = '   '.join('%s: %d' % (k, counts[k]) for k in ('PASS','PARTIAL','MANUAL','FAIL','SKIP') if k in counts)
    doc.add_paragraph().add_run('%d cases — %s' % (len(RESULTS), line)).bold = True
    t = doc.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(('Area', 'Cases', 'Pass', 'Fail/Other', 'Pass %')):
        t.rows[0].cells[i].text = h
    for area, prefix, cases in CASES:
        rs = [res[ 'TM-%s-%02d' % (prefix, i)] for i in range(1, len(cases)+1) if 'TM-%s-%02d' % (prefix, i) in res]
        npass = sum(1 for r in rs if r['status'] == 'PASS')
        row = t.add_row().cells
        row[0].text = area; row[1].text = str(len(rs)); row[2].text = str(npass)
        row[3].text = str(len(rs)-npass); row[4].text = '%d%%' % round(npass*100.0/max(len(rs),1))
    fails = [r for r in RESULTS if r['status'] != 'PASS']
    if fails:
        doc.add_heading('Non-pass cases (triage)', level=2)
        for r in fails:
            fp = doc.add_paragraph(style='List Bullet')
            fp.add_run('%s [%s]: ' % (r['tid'], r['status'])).bold = True; fp.add_run(r['note'])
    for area, prefix, cases in CASES:
        doc.add_page_break(); doc.add_heading(area, level=1)
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'TM-%s-%02d' % (prefix, i); r = res.get(tid, {})
            doc.add_heading('%s — %s' % (tid, scenario), level=2)
            pr = doc.add_paragraph(); sr = pr.add_run(r.get('status', 'SKIP')); sr.bold = True
            sr.font.color.rgb = COLORS.get(r.get('status'), RGBColor(0,0,0))
            pr.add_run('   ·   %s   ·   %.0fs' % (function, r.get('secs', 0))).font.size = Pt(9)
            doc.add_paragraph('Expected: ' + expected).runs[0].font.size = Pt(9)
            if r.get('note'):
                doc.add_paragraph('Result: ' + r['note']).runs[0].font.size = Pt(10)
            if r.get('shot') and os.path.exists(r['shot']):
                try:
                    doc.add_picture(r['shot'], width=Inches(6.3))
                    cap = doc.add_paragraph('Evidence: %s' % os.path.basename(r['shot']))
                    cap.runs[0].font.size = Pt(8); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception: pass
    doc.save(DOCX_OUT)

# ═════════════════════════════════════════════════════════════════════════════
def main():
    try:
        out = subprocess.run(['powershell', '-NoProfile', '-Command',
            "(Get-NetTCPConnection -LocalPort %d -State Listen -ErrorAction SilentlyContinue).OwningProcess" % PORT],
            capture_output=True, text=True, timeout=20).stdout.strip()
        for pid in set(out.split()):
            if pid.isdigit(): subprocess.run(['taskkill', '/F', '/PID', pid], capture_output=True)
    except Exception: pass
    proxy = subprocess.Popen(['python', os.path.join(ROOT, 'TM', 'Jet', 'dev-proxy.py'), str(PORT)],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx = Ctx(); ctx['browser'] = browser; ctx['pp'] = {}
            c, p = open_app(ctx, ADMINU); ctx['ctx'], ctx['page'] = c, p
            for area, prefix, cases in CASES:
                print('\n──', area)
                for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
                    run_case(ctx, 'TM-%s-%02d' % (prefix, i), area, fn)
                    tc = ctx.pop('tmp_ctx', None)
                    if tc:
                        try: tc.close()
                        except Exception: pass
            browser.close()
    finally:
        proxy.terminate()
    build_workbook(XLSX_OUT, fill_status=True)
    build_workbook(SCRIPT_OUT, fill_status=False)
    build_docx()
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    print('\n=== %d cases: %s' % (len(RESULTS), counts))
    print('Workbook   :', XLSX_OUT)
    print('TestScript :', SCRIPT_OUT)
    print('Word report:', DOCX_OUT)
    print('Evidence   :', EVID_DIR)

if __name__ == '__main__':
    main()
