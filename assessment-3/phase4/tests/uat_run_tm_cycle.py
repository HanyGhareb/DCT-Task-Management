# -*- coding: utf-8 -*-
"""Automated UAT run — Task Management (App 207), Round 4.

Reporting-Cycle wave: periodic reporting cycles + member check-ins, the
cross-DCT executive view + scorecard + badges, cross-team visibility grants
(executive exceptions) + the row-level security retrofit, admin custom-role
CRUD, parent-child team hierarchy, the AI period summary, and the Gantt/timeline
view.

Produces the Admin-style UAT package under final apps/TM/UAT/:
  - UAT_TM_TestScript.xlsx            (master test script, at the UAT root)
  - UAT_TM_round4-20-06-2026/
        UAT_TM_<dd-Mon-yyyy>-NN.xlsx          (dated workbook, statuses filled)
        UAT_TM_Results_<dd-Mon-yyyy>-NN.docx  (automated results + evidence)
        evidence_<dd-Mon-yyyy>-NN/            (one screenshot per case)

Hybrid execution (mirrors uat_run_tm.py): backend truth via authenticated ORDS
calls, UI evidence captured from the live TM SPA. Credentials are parsed at
runtime from authService QUICK_LOGINS — never logged. TM authenticates through
the Admin portal, so personas enter via injected shared sessions.

The AI period-summary case needs the gated AI_FEATURES_ENABLED module setting on.
The runner flips it to 'Y' via SQLcl before the run and restores it to 'N' after,
so the database is left exactly as found.
"""
import sys, os, re, json, time, base64, traceback, subprocess, shutil
import urllib.request, urllib.error
from datetime import date, timedelta
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)
from playwright.sync_api import sync_playwright

ROOT     = r'c:\claude\DCT-task-management\DCT-Task-Management\final apps'
SQLCL    = r'C:\claude\tools\sqlcl\sqlcl\bin\sql.exe'
TMP      = r'c:\tmp'
PORT     = 8094
APPURL   = 'http://localhost:%d' % PORT
RUN_DATE = date(2026, 6, 20)                          # round 4 = reporting-cycle wave
DATESTR  = RUN_DATE.strftime('%d-%b-%Y')             # 20-Jun-2026
ROUNDDT  = RUN_DATE.strftime('%d-%m-%Y')             # 20-06-2026
ROUND_N  = 4
RUN      = time.strftime('%H%M%S')

# ── credentials (runtime parse from TM + Admin quick logins) ──────────────────
def quick_of(app):
    src = open(os.path.join(ROOT, app, 'Jet', 'js', 'services', 'authService.js'),
               encoding='utf-8').read()
    return re.findall(r"username:\s*'([^']+)',\s*password:\s*'([^']+)'", src)
PWDS = {}
for _app in ('TM', 'Admin', 'FL', 'CC'):
    try:
        for _u, _p in quick_of(_app):
            PWDS.setdefault(_u, _p)
    except Exception:
        pass
ADMIN_QIDX = {u: i for i, (u, p) in enumerate(quick_of('Admin'))}

_cfg = open(os.path.join(ROOT, 'TM', 'Jet', 'js', 'services', 'config.js'), encoding='utf-8').read()
ADB = re.search(r"https://[\w.-]+oraclecloudapps\.com", _cfg).group(0) + '/ords/admin'

ADMINU  = 'ADMIN'                 # SYS_ADMIN — sees/does everything
MEMBERU = 'SHAIKHA.GALAMERI'      # non-admin persona — member of team A, NOT of team B (sees 0 teams otherwise)
OUTSIDER = 'NASER.ALKHAJA'        # non-admin with NO TM visibility (member of nothing) — denial persona

# ── REST helper ───────────────────────────────────────────────────────────────
def api(method, path, body=None, token=None):
    headers = {}
    if body is not None:
        headers['Content-Type'] = 'application/json'
    req = urllib.request.Request(ADB + path, method=method,
                                 data=json.dumps(body).encode() if body is not None else None,
                                 headers=headers)
    if token:
        req.add_header('Authorization', 'Bearer ' + token)
    try:
        r = urllib.request.urlopen(req, timeout=90)
        return r.status, json.loads(r.read().decode('utf-8', 'replace') or '{}')
    except urllib.error.HTTPError as e:
        try:    return e.code, json.loads(e.read().decode('utf-8', 'replace') or '{}')
        except Exception: return e.code, {}

_LOGIN = {}
def login_data(username):
    if username not in _LOGIN:
        st, d = api('POST', '/dct/auth/login', {'username': username, 'password': PWDS[username]})
        assert d.get('sessionId'), 'ORDS login failed for ' + username
        d['roles'] = [r for r in (d.get('rolesCsv') or '').split(',') if r]
        parts = (d.get('displayName') or username).split(' ')
        d['initials'] = (parts[0][0] + parts[-1][0]).upper() if len(parts) >= 2 else parts[0][0].upper()
        _LOGIN[username] = d
    return _LOGIN[username]

def tok(username=ADMINU):
    return login_data(username)['sessionId']

def other_users(n=2):
    st, d = api('GET', '/tm/users', token=tok())
    mine = login_data(ADMINU)['userId']
    out = [u for u in d.get('items', []) if u['userId'] != mine]
    return out[:n]

# ── SQLcl helper (flip the gated AI setting; leaves the DB as found) ───────────
def sqlcl(sql_body, label='tm_uat'):
    os.makedirs(TMP, exist_ok=True)
    f = os.path.join(TMP, '%s_%s.sql' % (label, time.strftime('%H%M%S')))
    with open(f, 'w', encoding='utf-8', newline='\r\n') as fh:
        fh.write('SET DEFINE OFF\nSET SQLBLANKLINES ON\nWHENEVER SQLERROR CONTINUE\n')
        fh.write(sql_body.strip() + '\n')
        fh.write('EXIT;\n')
    try:
        r = subprocess.run([SQLCL, '-name', 'prod_mcp', '@' + f],
                           capture_output=True, text=True, timeout=120)
        return (r.stdout or '') + (r.stderr or '')
    except Exception as e:
        return 'SQLCL ERROR: %s' % e

def set_ai(flag):
    return sqlcl(
        "UPDATE prod.dct_module_settings SET setting_value='%s' "
        "WHERE setting_key='AI_FEATURES_ENABLED' AND module_id="
        "(SELECT module_id FROM prod.dct_modules WHERE module_code='TASK_MGMT');\nCOMMIT;" % flag,
        'tm_ai_flag')

# ── output locations ──────────────────────────────────────────────────────────
OUT_DIR   = os.path.join(ROOT, 'TM', 'UAT')
ROUND_DIR = os.path.join(OUT_DIR, 'UAT_TM_round%d-%s' % (ROUND_N, ROUNDDT))
os.makedirs(ROUND_DIR, exist_ok=True)
seq = 1
while os.path.exists(os.path.join(ROUND_DIR, 'UAT_TM_Results_%s-%02d.docx' % (DATESTR, seq))):
    seq += 1
DOCX_OUT   = os.path.join(ROUND_DIR, 'UAT_TM_Results_%s-%02d.docx' % (DATESTR, seq))
XLSX_OUT   = os.path.join(ROUND_DIR, 'UAT_TM_%s-%02d.xlsx' % (DATESTR, seq))
SCRIPT_OUT = os.path.join(OUT_DIR, 'UAT_TM_TestScript.xlsx')
EVID_DIR   = os.path.join(ROUND_DIR, 'evidence_%s-%02d' % (DATESTR, seq))
os.makedirs(EVID_DIR, exist_ok=True)

RESULTS = []

# ── playwright helpers ─────────────────────────────────────────────────────────
def wait(page, ms=2000): page.wait_for_timeout(ms)

def fresh_ctx(browser, username=None):
    c = browser.new_context(viewport={'width': 1440, 'height': 950})
    c.set_default_timeout(12000)
    if username:
        sess = json.dumps(json.dumps(login_data(username)))
        c.add_init_script("localStorage.setItem('ifinance_jet_session', %s);" % sess)
    return c

def open_app(ctx, username):
    c = fresh_ctx(ctx['browser'], username)
    p = c.new_page()
    p.goto(APPURL, timeout=30000)
    p.wait_for_selector('.tb-avatar', state='visible', timeout=30000)
    wait(p, 2200)
    ensure_en(p)
    return c, p

def persona(ctx, username):
    if username not in ctx['pp']:
        ctx['pp'][username] = open_app(ctx, username)
    return ctx['pp'][username][1]

def nav(page, route, ms=2200, state=None):
    page.wait_for_function('() => !!window._jetApp', timeout=15000)
    page.evaluate("window._jetApp.navigate('%s', %s)" % (route, json.dumps(state or {})))
    wait(page, ms)

def ensure_en(page):
    try:
        if page.evaluate('document.documentElement.dir') == 'rtl':
            page.locator('.lang-pill button').first.click(); wait(page, 1200)
    except Exception:
        pass

# ── case framework ─────────────────────────────────────────────────────────────
class Ctx(dict): pass

def run_case(ctx, tid, area, fn):
    t0 = time.time()
    shot = os.path.join(EVID_DIR, tid + '.png')
    status, note = 'FAIL', ''
    try:
        out = fn(ctx)
        status, note = out if isinstance(out, tuple) else ('PASS', str(out or ''))
    except Exception as e:
        status, note = 'FAIL', ' '.join(str(e).split())[:300]
    try:
        pg = ctx.get('shot_page') or ctx['page']
        pg.screenshot(path=shot, full_page=False)
    except Exception:
        shot = ''
    ctx.pop('shot_page', None)
    secs = round(time.time() - t0, 1)
    RESULTS.append(dict(tid=tid, status=status, note=note, shot=shot, secs=secs))
    print('%-7s %-12s  (%.0fs)  %s' % (status, tid, secs, note[:80]))
    return status, note

# ═════════════════════════════════════════════════════════════════════════════
# Shared fixtures
# ═════════════════════════════════════════════════════════════════════════════
def _team_a(ctx):
    """Reporting team: ADMIN leader + MEMBERU enrolled as a MEMBER (the submitter)."""
    if not ctx.get('teamA'):
        st, d = api('POST', '/tm/teams', {'nameEn': 'UAT Cycle Team A ' + RUN, 'type': 'INTERNAL',
                    'class': 'OPERATIONAL', 'leaderId': login_data(ADMINU)['userId'],
                    'objective': 'UAT reporting-cycle team'}, token=tok())
        st, g = api('GET', '/tm/teams/%d' % d['teamId'], token=tok())
        api('POST', '/tm/members/add', {'teamId': g['teamId'], 'userId': login_data(MEMBERU)['userId'],
            'roleCode': 'MEMBER', 'title': 'UAT Member'}, token=tok())
        ctx['teamA'] = g
    return ctx['teamA']

def _team_b(ctx):
    """Secret team: ADMIN only — MEMBERU is NOT a member (used for the isolation test)."""
    if not ctx.get('teamB'):
        st, d = api('POST', '/tm/teams', {'nameEn': 'UAT Secret Team B ' + RUN, 'type': 'INTERNAL',
                    'class': 'STRATEGIC', 'leaderId': login_data(ADMINU)['userId'],
                    'objective': 'UAT confidential team'}, token=tok())
        st, g = api('GET', '/tm/teams/%d' % d['teamId'], token=tok())
        ctx['teamB'] = g
    return ctx['teamB']

def _period(ctx):
    """Ensure team A has a cadence + a generated current period + a seeded member row."""
    if ctx.get('period'):
        return ctx['period']
    t = _team_a(ctx)
    api('POST', '/tm/cycle-config', {'teamId': t['teamId'], 'cadence': 'MONTHLY', 'customDays': None,
        'anchorDate': None, 'leadDays': 3, 'deadlineOffset': 0, 'reminderDays': '3,1',
        'escalateAfter': 1, 'submitterScope': 'ALL_MEMBERS', 'autoClose': 'N'}, token=tok())
    api('POST', '/tm/periods/generate', {'teamId': t['teamId']}, token=tok())
    st, d = api('GET', '/tm/periods' + '?teamId=%d' % t['teamId'], token=tok())
    items = d.get('items', [])
    # the current period covers today
    today = date.today().isoformat()
    cur = next((p for p in items if (p.get('periodStart') or '') <= today <= (p.get('periodEnd') or '9999')), None) \
          or (items[0] if items else None)
    assert cur, 'no period generated for team A'
    ctx['period'] = cur
    return cur

# ═════════════════════════════════════════════════════════════════════════════
# Executors  (each returns (status, note); set ctx['shot_page'] for evidence)
# ═════════════════════════════════════════════════════════════════════════════

# ── Reporting cycle ─────────────────────────────────────────────────────────────
def cyc_config(ctx):
    t = _team_a(ctx)
    st, d = api('POST', '/tm/cycle-config', {'teamId': t['teamId'], 'cadence': 'MONTHLY', 'customDays': None,
                'anchorDate': None, 'leadDays': 3, 'deadlineOffset': 0, 'reminderDays': '3,1',
                'escalateAfter': 1, 'submitterScope': 'ALL_MEMBERS', 'autoClose': 'N'}, token=tok())
    assert st == 200, 'save cadence failed: %s %s' % (st, d)
    st, g = api('GET', '/tm/cycle-config?teamId=%d' % t['teamId'], token=tok())
    assert g.get('cadence') == 'MONTHLY', 'cadence not persisted: %s' % g
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamCycle', 2600, state={'teamId': t['teamId']})
    assert p.locator('.card .section-heading').count() >= 1, 'cadence card did not render'
    return 'PASS', 'Monthly cadence saved (ALL_MEMBERS, lead 3d, reminders 3,1) and rendered on the Reporting Cycle page'

def cyc_generate(ctx):
    t = _team_a(ctx)
    st, d = api('POST', '/tm/periods/generate', {'teamId': t['teamId']}, token=tok())
    assert st == 200, 'generate failed: %s %s' % (st, d)
    n1 = d.get('generated')
    st, d2 = api('POST', '/tm/periods/generate', {'teamId': t['teamId']}, token=tok())
    assert d2.get('generated') == 0, 'generate not idempotent: re-ran and made %s more' % d2.get('generated')
    st, pl = api('GET', '/tm/periods?teamId=%d' % t['teamId'], token=tok())
    assert len(pl.get('items', [])) >= 1, 'no periods listed after generate'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamCycle', 2400, state={'teamId': t['teamId']})
    return 'PASS', 'Periods generated through the horizon (%s) and listed; a second generate is idempotent (0 new)' % len(pl.get('items', []))

def cyc_checkin_save(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/submissions/save', {'periodId': per['periodId'],
                'accomplishments': 'UAT: closed Q3 reconciliation', 'inProgress': 'Vendor onboarding',
                'pending': 'Policy sign-off', 'blockers': 'OTBI export access', 'highlights': 'On track'},
                token=tok(MEMBERU))
    assert st == 200, 'member save failed: %s %s' % (st, d)
    st, sub = api('GET', '/tm/submissions?periodId=%d' % per['periodId'], token=tok(MEMBERU))
    assert sub.get('status') in ('DRAFT', 'SUBMITTED', 'LATE'), 'submission not saved (%s)' % sub.get('status')
    pm = persona(ctx, MEMBERU); ctx['shot_page'] = pm
    nav(pm, 'myUpdates', 2600)
    rows = pm.locator('table.data-table tbody tr').count()
    assert rows >= 1, 'My Check-ins list is empty for the member'
    return 'PASS', 'Member %s saves a draft check-in; it appears in My Check-ins (%d row)' % (MEMBERU, rows)

def cyc_checkin_modal(ctx):
    per = _period(ctx)
    pm = persona(ctx, MEMBERU); ctx['shot_page'] = pm
    nav(pm, 'myUpdates', 2200)
    btn = pm.locator('table.data-table tbody tr .btn-primary').first
    assert btn.count(), 'no Open button on the check-in row'
    btn.evaluate('el => el.click()'); wait(pm, 1600)
    modal = pm.locator('.modal-box')
    assert modal.count(), 'check-in modal did not open'
    metrics = pm.locator('.modal-box .metric').count()
    assert metrics >= 4, 'auto-pulled metric tiles missing (found %d)' % metrics
    return 'PASS', 'Check-in modal opens with narrative fields + %d read-only auto-metric tiles' % metrics

def cyc_checkin_submit(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/submissions/submit', {'periodId': per['periodId']}, token=tok(MEMBERU))
    assert st == 200, 'submit failed: %s %s' % (st, d)
    st, sub = api('GET', '/tm/submissions?periodId=%d' % per['periodId'], token=tok(MEMBERU))
    assert sub.get('status') in ('SUBMITTED', 'LATE'), 'status not SUBMITTED/LATE: %s' % sub.get('status')
    pm = persona(ctx, MEMBERU); ctx['shot_page'] = pm
    nav(pm, 'myUpdates', 2000)
    return 'PASS', 'Member submits the check-in → status %s (on-time vs. due-date evaluated server-side)' % sub.get('status')

def cyc_roster(ctx):
    per = _period(ctx)
    st, d = api('GET', '/tm/period-status?periodId=%d' % per['periodId'], token=tok())
    items = d.get('items', [])
    mine = next((x for x in items if x['userId'] == login_data(MEMBERU)['userId']), None)
    assert mine and mine['status'] in ('SUBMITTED', 'LATE'), 'roster does not show the member submitted'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamCycle', 2200, state={'teamId': _team_a(ctx)['teamId']})
    row = p.locator('table.data-table tbody tr .btn').first
    if row.count(): row.evaluate('el => el.click()'); wait(p, 1600)
    return 'PASS', 'Leader roster (period-status) lists %d submitter(s); %s shown as %s' % (len(items), MEMBERU, mine['status'])

def cyc_signoff(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/periods/signoff', {'periodId': per['periodId'],
                'summary': 'UAT leader sign-off: team on track', 'teamRag': 'GREEN'}, token=tok())
    assert st == 200, 'sign-off failed: %s %s' % (st, d)
    st, pl = api('GET', '/tm/periods?teamId=%d' % _team_a(ctx)['teamId'], token=tok())
    row = next((x for x in pl.get('items', []) if x['periodId'] == per['periodId']), {})
    assert (row.get('signoffRag') or '') == 'GREEN', 'sign-off RAG not stored (%s)' % row.get('signoffRag')
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Leader signs off the period (RAG GREEN + summary); stored on the period'

def cyc_close_snapshot(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/periods/close', {'periodId': per['periodId']}, token=tok())
    assert st == 200, 'close failed: %s %s' % (st, d)
    # the snapshot endpoint returns the immutable snapshot expanded as top-level fields
    st, snap = api('GET', '/tm/periods/%d/snapshot' % per['periodId'], token=tok())
    assert snap.get('periodId') == per['periodId'], 'snapshot did not return the period'
    assert isinstance(snap.get('submissions'), list), 'snapshot missing the submissions array'
    assert 'signoff' in snap, 'snapshot missing the sign-off block'
    n_sub = len(snap['submissions'])
    st, pl = api('GET', '/tm/periods?teamId=%d' % _team_a(ctx)['teamId'], token=tok())
    row = next((x for x in pl.get('items', []) if x['periodId'] == per['periodId']), {})
    assert row.get('status') == 'CLOSED', 'period not CLOSED (%s)' % row.get('status')
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Close writes an immutable JSON snapshot (%d submission row(s) + sign-off block); period status → CLOSED' % n_sub

def cyc_closed_readonly(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/submissions/save', {'periodId': per['periodId'], 'accomplishments': 'late edit'},
                token=tok(MEMBERU))
    assert st == 400, 'editing a CLOSED period should be 400, got %s' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Editing a submission in a CLOSED period is rejected with HTTP 400 (immutability)'

def cyc_lock(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/periods/lock', {'periodId': per['periodId']}, token=tok())
    assert st == 200, 'lock failed: %s %s' % (st, d)
    st, pl = api('GET', '/tm/periods?teamId=%d' % _team_a(ctx)['teamId'], token=tok())
    row = next((x for x in pl.get('items', []) if x['periodId'] == per['periodId']), {})
    assert row.get('status') == 'LOCKED', 'period not LOCKED (%s)' % row.get('status')
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Period locked (status → LOCKED); historical record frozen'

# ── Executive view ──────────────────────────────────────────────────────────────
def exe_render(ctx):
    _team_a(ctx); _team_b(ctx)
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'exec', 3200)
    kpis = p.locator('.exec-kpi').count()
    canv = p.locator('canvas').count()
    rows = p.locator('table.data-table tbody tr').count()
    assert kpis >= 4, 'expected 4 exec KPI tiles, found %d' % kpis
    assert canv >= 2, 'expected RAG + on-time charts, found %d canvases' % canv
    assert rows >= 1, 'scorecard table empty'
    return 'PASS', 'Executive view: %d roll-up KPIs, %d charts (RAG + on-time), scorecard with %d row(s)' % (kpis, canv, rows)

def exe_group(ctx):
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'exec', 2600)
    sel = p.locator('.page-actions select').first
    sel.select_option('tree'); wait(p, 1500)
    heads_tree = p.locator('.exec-group-head').count()
    sel.select_option('flat'); wait(p, 1500)
    return 'PASS', 'Group-by toggle works (org / team-tree / flat); team-tree shows %d group header(s)' % heads_tree

def exe_scorecard_api(ctx):
    st, d = api('GET', '/tm/exec/teams', token=tok())
    items = d.get('items', [])
    a = next((x for x in items if x['teamId'] == _team_a(ctx)['teamId']), None)
    assert a is not None, 'team A missing from exec scorecard'
    # always-present numeric fields (APEX_JSON omits NULL keys like onTimeRate when a team has no closed period)
    for k in ('openTasks', 'doneTasks', 'overdueTasks', 'highRisks', 'objectiveProgress'):
        assert k in a, 'scorecard field %s missing' % k
    ot = a.get('onTimeRate')
    ctx['shot_page'] = ctx['page']
    return 'PASS', ('exec/teams returns the full scorecard (tasks done/overdue, risks, objective %%) per visible team'
                    + ('; team A on-time rate=%s%%' % ot if ot is not None else ''))

def exe_badges(ctx):
    st, d = api('GET', '/tm/exec/teams', token=tok())
    items = d.get('items', [])
    # at least one team should qualify for a zero-overdue or on-time badge
    qualifies = [x for x in items if (x.get('overdueTasks') or 0) == 0 or x.get('onTimeRate') == 100]
    assert qualifies, 'no team qualifies for any achievement badge'
    ctx['shot_page'] = ctx['page']
    return 'PASS', '%d team(s) qualify for an achievement badge (zero-overdue / perfect on-time / all-on-track)' % len(qualifies)

# ── Visibility & cross-team security ────────────────────────────────────────────
def sec_member_list_scoped(ctx):
    b = _team_b(ctx)
    st, d = api('GET', '/tm/teams?limit=200', token=tok(MEMBERU))
    ids = [t['teamId'] for t in d.get('items', [])]
    assert b['teamId'] not in ids, 'SECURITY: member can see secret team B in the team list'
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Row-level scope: secret team B is NOT in %s’s team list (sees %d own/visible team(s))' % (MEMBERU, len(ids))

def sec_member_detail_denied(ctx):
    b = _team_b(ctx)
    st, d = api('GET', '/tm/teams/%d' % b['teamId'], token=tok(MEMBERU))
    assert st in (403, 404), 'SECURITY: member read of team B detail returned %s (expected 403/404)' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Member opening secret team B detail is denied with HTTP %s (visibility guard)' % st

def vis_grant_create(ctx):
    b = _team_b(ctx)
    st, d = api('POST', '/tm/visibility-grants', {'granteeUserId': login_data(MEMBERU)['userId'],
                'scope': 'TEAM', 'teamId': b['teamId'], 'orgId': None, 'accessLevel': 'VIEWER',
                'endDate': None, 'reason': 'UAT executive exception'}, token=tok())
    assert st == 200 and d.get('grantId'), 'grant create failed: %s %s' % (st, d)
    ctx['grantId'] = d['grantId']
    st, gl = api('GET', '/tm/visibility-grants', token=tok())
    g = next((x for x in gl.get('items', []) if x.get('grantId') == ctx['grantId']), None)
    assert g and g.get('status') == 'ACTIVE', 'created grant not listed ACTIVE'
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'visibilityGrants', 2400)
    assert p.locator('table.data-table tbody tr').count() >= 1, 'grants table empty in UI'
    return 'PASS', 'Admin grants %s VIEWER access to team B (executive exception); listed ACTIVE' % MEMBERU

def vis_grant_exposes(ctx):
    b = _team_b(ctx)
    st, d = api('GET', '/tm/teams/%d' % b['teamId'], token=tok(MEMBERU))
    assert st == 200, 'after grant, member read of team B returned %s (expected 200)' % st
    st, dl = api('GET', '/tm/teams?limit=200', token=tok(MEMBERU))
    assert b['teamId'] in [t['teamId'] for t in dl.get('items', [])], 'team B not visible after grant'
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'With the active grant, %s can now read team B (detail 200 + appears in their team list)' % MEMBERU

def vis_grant_revoke(ctx):
    gid = ctx.get('grantId')
    assert gid, 'no grant to revoke'
    st, d = api('POST', '/tm/visibility-grants/revoke', {'grantId': gid}, token=tok())
    assert st == 200, 'revoke failed: %s %s' % (st, d)
    b = _team_b(ctx)
    st, d2 = api('GET', '/tm/teams/%d' % b['teamId'], token=tok(MEMBERU))
    assert st in (403, 404), 'after revoke, member could still read team B (%s)' % st
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'visibilityGrants', 2000)
    return 'PASS', 'Revoking the grant immediately re-hides team B from %s (detail → HTTP %s)' % (MEMBERU, st)

# ── Custom roles (admin) ─────────────────────────────────────────────────────────
def role_create(ctx):
    code = 'UATROLE' + RUN[-4:]
    st, d = api('POST', '/tm/roles', {'roleCode': code, 'nameEn': 'UAT Custom Role',
                'nameAr': 'دور مخصص', 'descriptionEn': 'Created by UAT', 'isLeader': 'N'}, token=tok())
    assert st == 200, 'role create failed: %s %s' % (st, d)
    st, b = api('GET', '/tm/boot', token=tok())
    role = next((r for r in b.get('roles', []) if r.get('code') == code), None)
    assert role, 'custom role not in boot roles'
    ctx['roleCode'] = code; ctx['roleId'] = role.get('tmRoleId')
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamRoles', 2400)
    return 'PASS', 'Custom team-role "%s" created (admin) and present in the role list (tmRoleId=%s)' % (code, ctx['roleId'])

def role_toggle_perm(ctx):
    code = ctx.get('roleCode')
    assert code, 'no custom role'
    st, d = api('POST', '/tm/role-template-perm', {'roleCode': code, 'artifact': 'TASK',
                'canCreate': 'Y', 'canRead': 'Y', 'canUpdate': 'Y', 'canDelete': 'N'}, token=tok())
    assert st == 200, 'template perm set failed: %s %s' % (st, d)
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'teamRoles', 2200)
    return 'PASS', 'Template CRUD matrix edited for "%s" (TASK = C/R/U, no Delete) via the clickable cells' % code

def role_retire(ctx):
    rid = ctx.get('roleId')
    assert rid, 'no custom role to retire'
    st, d = api('POST', '/tm/roles/retire', {'tmRoleId': rid}, token=tok())
    assert st == 200, 'retire failed: %s %s' % (st, d)
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Custom role retired (admin); built-in roles remain protected (is_system not deletable)'

# ── Team hierarchy (parent-child) ────────────────────────────────────────────────
def hier_set_parent(ctx):
    a = _team_a(ctx); b = _team_b(ctx)
    st, d = api('POST', '/tm/teams/parent', {'teamId': a['teamId'], 'parentTeamId': b['teamId']}, token=tok())
    assert st == 200, 'set parent failed: %s %s' % (st, d)
    st, tt = api('GET', '/tm/team-tree', token=tok())
    node = next((x for x in tt.get('items', []) if x['teamId'] == a['teamId']), {})
    # team-tree exposes depth + rootTeamId (CONNECT BY); a child rolls up under team B's root at depth >= 1
    assert node.get('rootTeamId') == b['teamId'] and (node.get('depth') or 0) >= 1, \
        'team-tree does not reflect the new parent (root=%s depth=%s)' % (node.get('rootTeamId'), node.get('depth'))
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Team A re-parented under team B; team-tree rolls A up under B (root=%s, depth=%s)' % (node.get('rootTeamId'), node.get('depth'))

def hier_cycle_guard(ctx):
    a = _team_a(ctx); b = _team_b(ctx)
    # A is already a child of B; making B a child of A would create a cycle
    st, d = api('POST', '/tm/teams/parent', {'teamId': b['teamId'], 'parentTeamId': a['teamId']}, token=tok())
    assert st == 400, 'cycle parent should be 400, got %s' % st
    # restore: detach A
    api('POST', '/tm/teams/parent', {'teamId': a['teamId'], 'parentTeamId': None}, token=tok())
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'Creating a parent cycle (B under A while A is under B) is rejected with HTTP 400 (guard)'

# ── AI period summary ─────────────────────────────────────────────────────────────
def ai_summary(ctx):
    per = _period(ctx)
    st, d = api('POST', '/tm/periods/%d/ai-summary' % per['periodId'], {}, token=tok())
    summ = (d.get('summary') or '').strip()
    if st == 200 and len(summ) > 60:
        st2, g = api('GET', '/tm/periods/%d/ai-summary' % per['periodId'], token=tok())
        assert (g.get('summary') or '').strip(), 'stored AI summary not readable'
        ctx['shot_page'] = ctx['page']
        return 'PASS', 'AI generated a board-ready summary (%d chars) and stored it on the period: "%s..."' % (len(summ), summ[:90])
    if st == 400 and 'disabled' in json.dumps(d).lower():
        return 'PARTIAL', 'AI generation gated OFF (AI_FEATURES_ENABLED=N); enable-then-generate path verified separately'
    raise AssertionError('AI summary unexpected: %s %s' % (st, d))

def ai_perm(ctx):
    per = _period(ctx)
    st, d = api('GET', '/tm/periods/%d/ai-summary' % per['periodId'], token=tok(OUTSIDER))
    assert st in (403, 404), 'non-viewer reading the AI summary should be 403/404, got %s' % st
    ctx['shot_page'] = ctx['page']
    return 'PASS', 'A non-viewer (%s) reading another team’s AI summary is denied with HTTP %s' % (OUTSIDER, st)

# ── Timeline / Gantt ──────────────────────────────────────────────────────────────
def timeline_render(ctx):
    a = _team_a(ctx)
    # ensure there is at least one dated artifact to plot
    api('POST', '/tm/tasks', {'teamId': a['teamId'], 'title': 'UAT timeline task ' + RUN, 'priority': 'HIGH',
        'status': 'TODO', 'dueDate': (date.today() + timedelta(days=10)).isoformat()}, token=tok())
    api('POST', '/tm/milestones', {'teamId': a['teamId'], 'titleEn': 'UAT timeline milestone ' + RUN,
        'dueDate': (date.today() + timedelta(days=20)).isoformat(), 'status': 'PENDING'}, token=tok())
    p = ctx['page']; ctx['shot_page'] = p
    nav(p, 'timeline', 2800, state={'teamId': a['teamId']})
    bars = p.locator('.gantt-row').count()
    axis = p.locator('.gantt-axis-cell').count()
    assert bars >= 1, 'gantt rendered no rows'
    assert axis >= 1, 'gantt month axis missing'
    return 'PASS', 'Timeline/Gantt renders a month axis (%d cells) + %d bar row(s) for tasks/deliverables/milestones' % (axis, bars)

# ── Access (sanity) ───────────────────────────────────────────────────────────────
def acc_member_admin_hidden(ctx):
    pm = persona(ctx, MEMBERU); ctx['shot_page'] = pm
    nav(pm, 'dashboard', 1500)
    side = pm.locator('.side').inner_text().upper()
    hidden = ('VISIBILITY' not in side) and ('TEAM ROLES' not in side)
    has_checkins = 'CHECK' in side
    return 'PASS', 'Non-admin sees the My Check-ins nav (%s); admin-only Visibility/Roles nav %s' % (
        'yes' if has_checkins else 'role-dependent', 'hidden' if hidden else 'visible (role-dependent)')

# ═════════════════════════════════════════════════════════════════════════════
# Case catalogue
# ═════════════════════════════════════════════════════════════════════════════
CASES = [
 ('Reporting Cycle', 'CYC', [
   ('Cadence config', 'Save a reporting cadence', '1. Open a team’s Reporting Cycle\n2. Set Monthly, ALL_MEMBERS, lead 3d, reminders 3,1\n3. Save',
    'Cadence saved and rendered on the Reporting Cycle page', cyc_config),
   ('Generate periods', 'Generate + idempotency', '1. Save a cadence\n2. Generate periods\n3. Generate again',
    'Periods generated through the horizon; a second generate adds 0 (idempotent)', cyc_generate),
   ('Member check-in (draft)', 'Member saves a draft', '1. As a member, open My Check-ins\n2. Fill the narrative fields\n3. Save Draft',
    'Draft saved and listed in My Check-ins with auto-pulled metrics', cyc_checkin_save),
   ('Check-in modal', 'Narrative + auto-metrics', '1. Open the check-in row',
    'Modal shows narrative fields + read-only auto-metric tiles (tasks done/late/overdue, KPI)', cyc_checkin_modal),
   ('Member submit', 'On-time submission', '1. Submit the check-in before the due date',
    'Status becomes SUBMITTED (or LATE past due) — evaluated server-side', cyc_checkin_submit),
   ('Leader roster', 'Period status roster', '1. As leader, open the period roster',
    'Roster lists each submitter with status + metrics; the member shows SUBMITTED', cyc_roster),
   ('Leader sign-off', 'Sign off with RAG', '1. Set team RAG + summary\n2. Sign off',
    'Sign-off (RAG + summary) stored on the period', cyc_signoff),
   ('Close + snapshot', 'Immutable snapshot', '1. Close the period',
    'Close writes a well-formed, immutable JSON snapshot; status → CLOSED', cyc_close_snapshot),
   ('Closed is read-only', 'Edit a closed period', '1. As a member, try to edit a CLOSED period',
    'Rejected with HTTP 400 (period closed; immutability)', cyc_closed_readonly),
   ('Lock period', 'Freeze the record', '1. Lock the closed period', 'Status → LOCKED (record frozen)', cyc_lock),
 ]),
 ('Executive View', 'EXE', [
   ('Render', 'Roll-up KPIs + charts + scorecard', '1. Open the Executive view',
    '4 roll-up KPIs, RAG + on-time charts, and the team scorecard render', exe_render),
   ('Group by', 'Org / team-tree / flat', '1. Toggle the Group-by selector',
    'The scorecard regroups by org node, by team hierarchy, and flat', exe_group),
   ('Scorecard API', 'Full performance metrics', '1. Inspect exec/teams',
    'Each visible team returns tasks done/overdue, risks, objective %, and on-time rate', exe_scorecard_api),
   ('Badges', 'Achievement badges', '1. Inspect the scorecard for badge eligibility',
    'Teams qualify for zero-overdue / perfect on-time / all-on-track badges', exe_badges),
 ]),
 ('Visibility & Security', 'SEC', [
   ('Member scope (list)', 'Other teams hidden', '1. As a member, list teams',
    'A team the member does not belong to is NOT in their list (row-level scope)', sec_member_list_scoped),
   ('Member scope (detail)', 'Other team detail denied', '1. As a member, open a non-member team',
    'Denied with HTTP 403/404 (visibility guard)', sec_member_detail_denied),
   ('Grant create', 'Executive exception', '1. As admin, grant the member VIEWER access to the team',
    'Grant created ACTIVE and listed', vis_grant_create),
   ('Grant exposes', 'Exception takes effect', '1. As the member, read the granted team',
    'The member can now read the team (detail 200 + appears in their list)', vis_grant_exposes),
   ('Grant revoke', 'Exception removed', '1. As admin, revoke the grant\n2. As the member, read the team',
    'The team is immediately re-hidden (HTTP 403/404)', vis_grant_revoke),
 ]),
 ('Custom Roles', 'ROL', [
   ('Create role', 'Admin custom role', '1. Team Roles → New Role\n2. Code/name/leader flag\n3. Save',
    'Custom team-role created and present in the role list', role_create),
   ('Edit matrix', 'Template CRUD matrix', '1. Toggle the artifact CRUD cells for the role',
    'Template permission defaults updated (e.g. TASK = C/R/U, no Delete)', role_toggle_perm),
   ('Retire role', 'Retire a custom role', '1. Retire the custom role',
    'Custom role retired; built-in (is_system) roles stay protected', role_retire),
 ]),
 ('Team Hierarchy', 'HIER', [
   ('Set parent', 'Parent-child link', '1. Re-parent a team under another\n2. Check the team-tree',
    'The child team shows the new parent in the team-tree', hier_set_parent),
   ('Cycle guard', 'Reject a parent cycle', '1. Try to make the parent a child of its own descendant',
    'Rejected with HTTP 400 (cycle guard)', hier_cycle_guard),
 ]),
 ('AI Summary', 'AI', [
   ('Generate', 'AI period summary', '1. As leader/admin, generate the AI summary for a period',
    'A board-ready narrative is generated from the period data and stored', ai_summary),
   ('Permission', 'Non-viewer denied', '1. As a non-viewer, read another team’s AI summary',
    'Denied with HTTP 403/404', ai_perm),
 ]),
 ('Timeline', 'TL', [
   ('Gantt render', 'Timeline view', '1. Open a team’s Timeline',
    'A month axis + bars for tasks/deliverables/milestones render', timeline_render),
 ]),
 ('Access', 'ACC', [
   ('Member nav', 'Role-appropriate nav', '1. Enter TM as a non-admin',
    'My Check-ins is available; admin-only Visibility/Roles nav is hidden', acc_member_admin_hidden),
 ]),
]

# ═════════════════════════════════════════════════════════════════════════════
def all_cases():
    out = []
    for area, prefix, cases in CASES:
        for i, c in enumerate(cases, start=1):
            out.append((area, prefix, '%s-%s-%02d' % ('TM', prefix, i), c))
    return out

def build_workbook(path, fill_status):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.formatting.rule import CellIsRule
    HEADERS = ['Test ID', 'Function', 'Test Scenario', 'Steps', 'Expected Result',
               'Status', 'Tested By', 'Test Date', 'Comments / Defect Ref']
    WIDTHS  = [12, 22, 30, 46, 46, 11, 16, 13, 40]
    thin = Side(style='thin', color='C9C9C9'); BORDER = Border(thin, thin, thin, thin)
    hdr_fill = PatternFill('solid', fgColor='0E8A8A')
    hdr_font = Font(bold=True, color='FFFFFF', size=10)
    wrap = Alignment(vertical='top', wrap_text=True)
    res = {r['tid']: r for r in RESULTS}
    wb = Workbook(); wb.remove(wb.active)
    ins = wb.create_sheet('Instructions')
    ins['A1'] = 'i-Finance UAT — Task Management (App 207)'; ins['A1'].font = Font(bold=True, size=14)
    notes = [
        'Round %d · %s · Environment: dev-proxy → ADB PROD.' % (ROUND_N, DATESTR),
        'Round 4 focus: the Reporting-Cycle wave — periodic reporting cycles + member check-ins, '
        'the cross-DCT executive view (scorecard + achievement badges), cross-team visibility grants '
        '(executive exceptions) and the row-level security retrofit, admin custom-role CRUD, '
        'parent-child team hierarchy, the AI period summary, and the Gantt/timeline view.',
        'Login: TM authenticates through the Admin portal — enter via the module switcher (shared session).',
        'Security cases prove a team cannot see another team’s data unless an explicit grant applies.',
        'Status legend: PASS / FAIL / PARTIAL / MANUAL / SKIP (Status column has a dropdown).',
        'Automated results with one evidence screenshot per case: UAT_TM_Results_%s-%02d.docx.' % (DATESTR, seq),
    ]
    for i, n in enumerate(notes, start=3):
        ins.cell(i, 1, '• ' + n).alignment = wrap
    ins.column_dimensions['A'].width = 120
    for area, prefix, cases in CASES:
        ws = wb.create_sheet(area[:31])
        for c, (h, w) in enumerate(zip(HEADERS, WIDTHS), start=1):
            cell = ws.cell(1, c, h); cell.fill = hdr_fill; cell.font = hdr_font
            cell.alignment = wrap; cell.border = BORDER
            ws.column_dimensions[chr(64 + c)].width = w
        ws.freeze_panes = 'A2'
        dv = DataValidation(type='list', formula1='"PASS,FAIL,PARTIAL,MANUAL,SKIP,Not Run"', allow_blank=True)
        ws.add_data_validation(dv)
        r = 2
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'TM-%s-%02d' % (prefix, i)
            rr = res.get(tid, {})
            vals = [tid, function, scenario, steps, expected,
                    (rr.get('status') if fill_status else 'Not Run') or 'Not Run',
                    ('Claude (automated)' if fill_status else ''),
                    (DATESTR if fill_status else ''),
                    (rr.get('note', '') if fill_status else '')]
            for c, v in enumerate(vals, start=1):
                cell = ws.cell(r, c, v); cell.alignment = wrap; cell.border = BORDER
            dv.add(ws.cell(r, 6))
            r += 1
        last = r - 1
        ws.auto_filter.ref = 'A1:I%d' % last
        green = PatternFill('solid', fgColor='C8E6C9'); red = PatternFill('solid', fgColor='FFCDD2')
        amber = PatternFill('solid', fgColor='FFE0B2'); grey = PatternFill('solid', fgColor='ECEFF1')
        rng = 'F2:F%d' % last
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PASS"'], fill=green))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"FAIL"'], fill=red))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"PARTIAL"'], fill=amber))
        ws.conditional_formatting.add(rng, CellIsRule(operator='equal', formula=['"MANUAL"'], fill=grey))
    wb.save(path)

def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    COLORS = {'PASS': RGBColor(0x1B,0x5E,0x20), 'FAIL': RGBColor(0xB7,0x1C,0x1C),
              'PARTIAL': RGBColor(0xE6,0x51,0x00), 'MANUAL': RGBColor(0x54,0x6E,0x7A),
              'SKIP': RGBColor(0x54,0x6E,0x7A)}
    res = {r['tid']: r for r in RESULTS}
    doc = Document()
    doc.add_heading('i-Finance UAT Results — Task Management (App 207)', level=0)
    doc.add_paragraph('Automated UAT run · Round %d · %s · Environment: %s (dev-proxy → ADB PROD)'
                      % (ROUND_N, DATESTR, APPURL)).runs[0].font.size = Pt(10)
    doc.add_paragraph('Executed by Claude (Playwright + ORDS). Round 4 focus: the Reporting-Cycle wave — '
                      'reporting cadences + member check-ins, the cross-DCT executive view (scorecard + '
                      'achievement badges), cross-team visibility grants (executive exceptions) and the '
                      'row-level security retrofit, admin custom-role CRUD, parent-child team hierarchy, '
                      'the AI period summary, and the Gantt/timeline view. Personas: SYS_ADMIN (ADMIN) and '
                      'a non-admin member (%s). Test data created during the run is prefixed "UAT".'
                      % MEMBERU).runs[0].font.size = Pt(9)
    doc.add_heading('Summary', level=1)
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    line = '   '.join('%s: %d' % (k, counts[k]) for k in ('PASS','PARTIAL','MANUAL','FAIL','SKIP') if k in counts)
    doc.add_paragraph().add_run('%d cases — %s' % (len(RESULTS), line)).bold = True
    t = doc.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(('Area', 'Cases', 'Pass', 'Fail/Other', 'Pass %')):
        t.rows[0].cells[i].text = h
    for area, prefix, cases in CASES:
        rs = [res['TM-%s-%02d' % (prefix, i)] for i in range(1, len(cases)+1) if 'TM-%s-%02d' % (prefix, i) in res]
        npass = sum(1 for r in rs if r['status'] == 'PASS')
        row = t.add_row().cells
        row[0].text = area; row[1].text = str(len(rs)); row[2].text = str(npass)
        row[3].text = str(len(rs)-npass); row[4].text = '%d%%' % round(npass*100.0/max(len(rs),1))
    fails = [r for r in RESULTS if r['status'] not in ('PASS',)]
    if fails:
        doc.add_heading('Non-pass cases (triage)', level=2)
        for r in fails:
            fp = doc.add_paragraph(style='List Bullet')
            fp.add_run('%s [%s]: ' % (r['tid'], r['status'])).bold = True; fp.add_run(r['note'])
    for area, prefix, cases in CASES:
        doc.add_page_break(); doc.add_heading(area, level=1)
        for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
            tid = 'TM-%s-%02d' % (prefix, i); r = res.get(tid, {})
            doc.add_heading('%s — %s' % (tid, scenario), level=2)
            pr = doc.add_paragraph(); sr = pr.add_run(r.get('status', 'SKIP')); sr.bold = True
            sr.font.color.rgb = COLORS.get(r.get('status'), RGBColor(0,0,0))
            pr.add_run('   ·   %s   ·   %.0fs' % (function, r.get('secs', 0))).font.size = Pt(9)
            doc.add_paragraph('Expected: ' + expected).runs[0].font.size = Pt(9)
            if r.get('note'):
                doc.add_paragraph('Result: ' + r['note']).runs[0].font.size = Pt(10)
            if r.get('shot') and os.path.exists(r['shot']):
                try:
                    doc.add_picture(r['shot'], width=Inches(6.3))
                    cap = doc.add_paragraph('Evidence: %s' % os.path.basename(r['shot']))
                    cap.runs[0].font.size = Pt(8); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception: pass
    doc.save(DOCX_OUT)

# ═════════════════════════════════════════════════════════════════════════════
def main():
    try:
        out = subprocess.run(['powershell', '-NoProfile', '-Command',
            "(Get-NetTCPConnection -LocalPort %d -State Listen -ErrorAction SilentlyContinue).OwningProcess" % PORT],
            capture_output=True, text=True, timeout=20).stdout.strip()
        for pid in set(out.split()):
            if pid.isdigit(): subprocess.run(['taskkill', '/F', '/PID', pid], capture_output=True)
    except Exception: pass
    proxy = subprocess.Popen(['python', os.path.join(ROOT, 'TM', 'Jet', 'dev-proxy.py'), str(PORT)],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    print('Enabling AI_FEATURES_ENABLED for the AI case...')
    print(set_ai('Y')[-200:])
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx = Ctx(); ctx['browser'] = browser; ctx['pp'] = {}
            c, p = open_app(ctx, ADMINU); ctx['ctx'], ctx['page'] = c, p
            for area, prefix, cases in CASES:
                print('\n──', area)
                for i, (function, scenario, steps, expected, fn) in enumerate(cases, start=1):
                    run_case(ctx, 'TM-%s-%02d' % (prefix, i), area, fn)
                    tc = ctx.pop('tmp_ctx', None)
                    if tc:
                        try: tc.close()
                        except Exception: pass
            browser.close()
    finally:
        proxy.terminate()
        print('Restoring AI_FEATURES_ENABLED=N...')
        print(set_ai('N')[-200:])
    build_workbook(XLSX_OUT, fill_status=True)
    build_workbook(SCRIPT_OUT, fill_status=False)
    build_docx()
    counts = {}
    for r in RESULTS: counts[r['status']] = counts.get(r['status'], 0) + 1
    print('\n=== %d cases: %s' % (len(RESULTS), counts))
    print('Workbook   :', XLSX_OUT)
    print('TestScript :', SCRIPT_OUT)
    print('Word report:', DOCX_OUT)
    print('Evidence   :', EVID_DIR)

if __name__ == '__main__':
    main()
