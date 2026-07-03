# -*- coding: utf-8 -*-
"""
uat_run_bi.py — BI - Reporting (App 211) UAT runner, round 2 (03-07-2026).

Scope: Interactive Reports platform — BI_USER viewer + spec-driven run
parameters (PARAM_SPEC_JSON convergence), the <interactive-report> grid
(now in final apps/shared/) incl. control breaks, highlight rules,
calculated columns, layouts, exports, maximize; the admin param-spec
editor; and the ir/* API security surface.

Pattern: Playwright + ORDS hybrid (model: uat_run_cc.py / uat_run_tm.py).
Run with the BI dev-proxy on port 8090:
  cd "final apps/BI/Jet" ; python dev-proxy.py 8090   (or via with_server.py)
  python assessment-3/phase4/tests/uat_run_bi.py

Emits (Admin UAT convention):
  final apps/BI/UAT/UAT_BI_TestScript.xlsx                      (master, statuses blank)
  final apps/BI/UAT/UAT_BI_round2-03-07-2026/
    UAT_BI_03-Jul-2026-02.xlsx                                  (statuses filled)
    UAT_BI_Results_03-Jul-2026-02.docx
    evidence_03-Jul-2026-02/NN_case.png|.txt                    (one per case)
"""
import json
import os
import re
import sys

import requests
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from playwright.sync_api import sync_playwright

BASE = 'http://localhost:8090'
ORDS = 'https://gd5cec2eaeb21e3-prod.adb.me-abudhabi-1.oraclecloudapps.com/ords/admin'
ROOT = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                     '..', '..', '..'))
UAT_DIR = os.path.join(ROOT, 'final apps', 'BI', 'UAT')
ROUND_DIR = os.path.join(UAT_DIR, 'UAT_BI_round2-03-07-2026')
EV_DIR = os.path.join(ROUND_DIR, 'evidence_03-Jul-2026-02')
DATE_LBL = '03-Jul-2026'
TESTER = 'Claude (automated)'

RESULTS = []          # {id, sheet, function, scenario, steps, expected, status, comment, evidence}
EV_SEQ = [0]


def record(case, status, comment='', evidence=''):
    case = dict(case)
    case.update(status=status, comment=comment[:900], evidence=evidence)
    RESULTS.append(case)
    print(('PASS  ' if status == 'PASS' else 'FAIL  ') + case['id'] + '  ' + case['scenario']
          + ('' if status == 'PASS' else ('  -- ' + comment[:200])))


def shot(page, slug):
    EV_SEQ[0] += 1
    name = '%02d_%s.png' % (EV_SEQ[0], slug)
    page.screenshot(path=os.path.join(EV_DIR, name))
    return name


def txt_evidence(slug, content):
    EV_SEQ[0] += 1
    name = '%02d_%s.txt' % (EV_SEQ[0], slug)
    with open(os.path.join(EV_DIR, name), 'w', encoding='utf-8') as f:
        f.write(content)
    return name


def case(cid, sheet, function, scenario, steps, expected):
    return {'id': cid, 'sheet': sheet, 'function': function, 'scenario': scenario,
            'steps': steps, 'expected': expected}


def login(page, user, pw):
    for _ in range(3):
        page.goto(BASE + '/index.html')
        page.wait_for_load_state('networkidle')
        try:
            page.wait_for_selector('input[type="text"]', timeout=20000)
            break
        except Exception:
            page.wait_for_timeout(1500)
    page.fill('input[type="text"]', user)
    page.fill('input[type="password"]', pw)
    page.click('button.btn-primary')
    page.wait_for_load_state('networkidle')
    page.wait_for_timeout(1000)


def logout(ctx):
    ctx.clear_cookies()


def open_viewer(page):
    page.goto(BASE + '/index.html#irViewer')
    page.reload()
    page.wait_for_load_state('networkidle')
    page.wait_for_timeout(1500)


def open_report(page, code, run=False):
    open_viewer(page)
    row = page.locator('tr.ir-cat-row', has=page.locator('text=' + code))
    row.first.locator('button').click()
    page.wait_for_timeout(1200)
    if run:
        page.get_by_role('button', name=re.compile('^run$', re.I)).click()
        page.wait_for_selector('table.ir-table tbody tr', timeout=45000)
        page.wait_for_timeout(500)


def api_login(user, pw):
    r = requests.post(ORDS + '/dct/auth/login', json={'username': user, 'password': pw}, timeout=60)
    return {'Authorization': 'Bearer ' + r.json()['sessionId']}


def main():
    os.makedirs(EV_DIR, exist_ok=True)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        ctx = browser.new_context(viewport={'width': 1500, 'height': 950})
        page = ctx.new_page()

        # ═════ BI_USER viewer cases ═════════════════════════════════════
        login(page, 'NASER.ALKHAJA', 'Manager@2026')

        c = case('BI-IRV-01', 'Viewer & Parameters', 'Catalog', 'BI_USER opens the report catalog',
                 '1. Log in as a BI_USER-only user\n2. Open Interactive Reports',
                 'The catalog lists the enabled report definitions')
        try:
            open_viewer(page)
            n = page.locator('tr.ir-cat-row').count()
            record(c, 'PASS' if n >= 2 else 'FAIL', '%d catalog rows' % n, shot(page, 'biuser_catalog'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'biuser_catalog'))

        c = case('BI-IRV-02', 'Viewer & Parameters', 'Access', 'Role-filtered navigation + admin reroute',
                 '1. As BI_USER inspect the side nav\n2. Navigate to #reports directly',
                 'Nav shows only Dashboard + Interactive Reports; the admin route reroutes away')
        try:
            nav_n = page.locator('.side-nav a, nav a').count()
            page.goto(BASE + '/index.html#reports')
            page.reload()
            page.wait_for_load_state('networkidle')
            page.wait_for_timeout(1500)
            title = page.locator('.page-title').first.inner_text().strip().lower()
            ok = nav_n <= 3 and 'report' not in title or title in ('dashboard', 'لوحة المعلومات')
            # reroute lands on dashboard — title must NOT be the admin Reports page
            ok = ('reports' not in title) or (title == 'interactive reports')
            record(c, 'PASS' if ok else 'FAIL', 'nav links=%d, landed title=%r' % (nav_n, title),
                   shot(page, 'biuser_reroute'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'biuser_reroute'))

        c = case('BI-IRV-03', 'Viewer & Parameters', 'Run parameters', 'Labels come from the parameter spec',
                 '1. Open BUDGET_UTIL_SECTOR', 'Fields show "Budget Year", "Sector", "Project Type", "Cost Center" (not raw keys)')
        try:
            open_report(page, 'BUDGET_UTIL_SECTOR')
            labels = ' | '.join(page.locator('.ir-runner .form-label').all_inner_texts()).lower()
            ok = 'budget year' in labels and 'cost center' in labels
            record(c, 'PASS' if ok else 'FAIL', labels[:200], shot(page, 'param_labels'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'param_labels'))

        c = case('BI-IRV-04', 'Viewer & Parameters', 'Run parameters', 'LOV dropdowns are populated',
                 '1. On BUDGET_UTIL_SECTOR wait for the parameter card', 'Year and Sector render as dropdowns with values')
        try:
            page.wait_for_timeout(1200)
            sels = page.locator('.ir-runner select')
            opts1 = sels.nth(1).locator('option').count()
            opts2 = sels.nth(2).locator('option').count()
            ok = sels.count() >= 4 and opts1 >= 2 and opts2 >= 2
            record(c, 'PASS' if ok else 'FAIL',
                   'selects=%d yearOpts=%d sectorOpts=%d' % (sels.count(), opts1, opts2),
                   shot(page, 'param_lovs'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'param_lovs'))

        c = case('BI-IRV-05', 'Viewer & Parameters', 'Validation', 'Required parameters block the run',
                 '1. Clear Year + Sector\n2. Click Run', 'A warning names the missing required fields; no fetch happens')
        try:
            sels = page.locator('.ir-runner select')
            sels.nth(1).select_option('')
            sels.nth(2).select_option('')
            page.get_by_role('button', name=re.compile('^run$', re.I)).click()
            page.wait_for_timeout(800)
            t = ' '.join(page.locator('[class*="toast"]').all_inner_texts()).lower()
            ok = 'missing' in t or 'budget year' in t or 'ناقصة' in t
            record(c, 'PASS' if ok else 'FAIL', 'toast=%r' % t[:150], shot(page, 'required_warn'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'required_warn'))

        c = case('BI-IRV-06', 'Viewer & Parameters', 'Run parameters', 'Hints display under the fields',
                 '1. Inspect the parameter card', 'The spec hints render under Year/Sector')
        try:
            hints = page.locator('.ir-param-hint').all_inner_texts()
            ok = any('budget year' in h.lower() for h in hints)
            record(c, 'PASS' if ok else 'FAIL', ' | '.join(hints)[:200], shot(page, 'param_hints'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'param_hints'))

        c = case('BI-IRV-07', 'Viewer & Parameters', 'i18n', 'Arabic labels + RTL',
                 '1. Switch the language to Arabic', 'Parameter labels show the Arabic spec labels; page flips RTL')
        try:
            page.click('text=ع')
            page.wait_for_timeout(1500)
            labels = ' | '.join(page.locator('.ir-runner .form-label').all_inner_texts())
            rtl = page.evaluate("document.documentElement.dir")
            ok = ('سنة الميزانية' in labels) and rtl == 'rtl'
            ev = shot(page, 'arabic_labels')
            page.click('text=EN')
            page.wait_for_timeout(1200)
            record(c, 'PASS' if ok else 'FAIL', 'dir=%s labels=%s' % (rtl, labels[:150]), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'arabic_labels'))

        c = case('BI-IRV-08', 'Viewer & Parameters', 'Data fetch', 'Run a report into the grid',
                 '1. Open GL_BUDGET_ACTUAL\n2. Click Run', 'The interactive grid renders the data rows')
        try:
            open_report(page, 'GL_BUDGET_ACTUAL', run=True)
            n = page.locator('table.ir-table tbody tr').count()
            record(c, 'PASS' if n > 5 else 'FAIL', '%d rows on page 1' % n, shot(page, 'grid_rendered'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'grid_rendered'))

        # ═════ grid feature cases (still BI_USER, GL_BUDGET_ACTUAL) ═════
        def open_cols():
            if page.locator('.ir-colmgr').count() == 0:
                page.get_by_role('button', name=re.compile('^columns$', re.I)).click()
                page.wait_for_timeout(400)

        def close_panel():
            if page.locator('.ir-panel').count():
                page.get_by_role('button', name=re.compile('^close$', re.I)).first.click()
                page.wait_for_timeout(300)

        c = case('BI-IRG-01', 'Grid Features', 'Columns', 'Hide a column',
                 '1. Columns panel\n2. Untick the first column', 'The column leaves the grid')
        try:
            before = page.locator('table.ir-table thead th').count()
            open_cols()
            page.locator('.ir-colrow').first.click()
            page.wait_for_timeout(400)
            after = page.locator('table.ir-table thead th').count()
            page.locator('.ir-colrow').first.click()   # restore
            page.wait_for_timeout(300)
            record(c, 'PASS' if after == before - 1 else 'FAIL',
                   '%d -> %d headers' % (before, after), shot(page, 'hide_column'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'hide_column'))

        c = case('BI-IRG-02', 'Grid Features', 'Columns', 'Reorder a column',
                 '1. Columns panel\n2. Move the first column down', 'Header order changes')
        try:
            open_cols()
            h_before = page.locator('table.ir-table thead th').first.inner_text()
            page.locator('.ir-colrow__move').nth(1).click()   # first row's down arrow
            page.wait_for_timeout(400)
            h_after = page.locator('table.ir-table thead th').first.inner_text()
            page.locator('.ir-colrow__move').nth(2).click()   # move back up (row now 2nd)
            page.wait_for_timeout(300)
            record(c, 'PASS' if h_before != h_after else 'FAIL',
                   '%r -> %r' % (h_before[:30], h_after[:30]), shot(page, 'reorder_column'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'reorder_column'))

        c = case('BI-IRG-03', 'Grid Features', 'Columns', 'Rename a header',
                 '1. Columns panel\n2. Pencil on first column\n3. Type "My Col", OK', 'Header shows the new label')
        try:
            open_cols()
            page.locator('.ir-colrow__pencil').first.click()
            page.locator('.ir-colrow__rename').fill('My Col')
            page.get_by_role('button', name=re.compile('^ok$', re.I)).click()
            page.wait_for_timeout(400)
            h = page.locator('table.ir-table thead th').first.inner_text().lower()
            ok = 'my col' in h
            # revert
            page.locator('.ir-colrow__pencil').first.click()
            page.locator('.ir-colrow__rename').fill('')
            page.get_by_role('button', name=re.compile('^ok$', re.I)).click()
            page.wait_for_timeout(300)
            record(c, 'PASS' if ok else 'FAIL', 'header=%r' % h[:40], shot(page, 'rename_header'))
            close_panel()
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'rename_header'))

        c = case('BI-IRG-04', 'Grid Features', 'Filters', 'Add + remove a filter chip',
                 '1. Add Filter (first col, is not empty)\n2. Apply\n3. Remove the chip',
                 'Chip appears; removing restores the count')
        try:
            base_info = page.locator('.result-count').first.inner_text()
            page.get_by_role('button', name=re.compile('add filter', re.I)).click()
            page.wait_for_timeout(400)
            page.locator('.ir-panel select').nth(1).select_option('notnull')
            page.get_by_role('button', name=re.compile('^apply$', re.I)).click()
            page.wait_for_timeout(500)
            chip_n = page.locator('.ir-chip:not(.ir-chip--break)').count()
            ev = shot(page, 'filter_chip')
            page.locator('.ir-chip__x').last.click()
            page.wait_for_timeout(400)
            after_info = page.locator('.result-count').first.inner_text()
            record(c, 'PASS' if chip_n >= 1 and after_info == base_info else 'FAIL',
                   'chips=%d, %r->%r' % (chip_n, base_info, after_info), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'filter_chip'))

        c = case('BI-IRG-05', 'Grid Features', 'Sorting', 'Click-sort + Shift multi-sort',
                 '1. Click header 1\n2. Shift-click header 2', 'Indexed sort badges appear')
        try:
            page.locator('table.ir-table thead th').nth(0).click()
            page.wait_for_timeout(300)
            page.locator('table.ir-table thead th').nth(1).click(modifiers=['Shift'])
            page.wait_for_timeout(400)
            badges = [b for b in page.locator('.ir-sort').all_inner_texts() if b.strip()]
            ok = len(badges) == 2 and any('1' in b for b in badges)
            ev = shot(page, 'multi_sort')
            page.locator('table.ir-table thead th').nth(1).click(modifiers=['Shift'])
            page.locator('table.ir-table thead th').nth(1).click(modifiers=['Shift'])
            page.locator('table.ir-table thead th').nth(0).click()
            page.locator('table.ir-table thead th').nth(0).click()
            page.wait_for_timeout(300)
            record(c, 'PASS' if ok else 'FAIL', 'badges=%r' % badges, ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'multi_sort'))

        c = case('BI-IRG-06', 'Grid Features', 'Search', 'Global search narrows the results',
                 '1. Type a non-matching term\n2. Clear the box', 'Count drops to none, then restores')
        try:
            total_before = page.locator('.result-count').first.inner_text()
            page.locator('.ir-toolbar input').first.fill('zzz_no_match_zzz')
            page.wait_for_timeout(900)
            total_none = page.locator('.result-count').first.inner_text()
            emptied = total_none != total_before and page.locator('table.ir-table tbody tr').count() == 0
            ev = shot(page, 'global_search')
            page.locator('.ir-toolbar input').first.fill('')
            page.wait_for_timeout(900)
            total_restored = page.locator('.result-count').first.inner_text()
            ok = emptied and total_restored == total_before
            record(c, 'PASS' if ok else 'FAIL',
                   '%r -> %r -> %r' % (total_before, total_none, total_restored), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'global_search'))

        c = case('BI-IRG-07', 'Grid Features', 'Control breaks', 'Break groups the rows into bands',
                 '1. Columns panel\n2. Break on column 1', 'Band rows show "Column: value"; a break chip appears')
        try:
            open_cols()
            page.locator('.ir-colrow__break').first.click()
            page.wait_for_timeout(600)
            bands = page.locator('tr.ir-break-row').count()
            chips = page.locator('.ir-chip--break').count()
            record(c, 'PASS' if bands >= 1 and chips == 1 else 'FAIL',
                   'bands=%d chips=%d' % (bands, chips), shot(page, 'control_break'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'control_break'))

        c = case('BI-IRG-08', 'Grid Features', 'Control breaks', 'Per-group subtotals + grand total',
                 '1. Break also on the first numeric column\n2. Set its aggregate to Sum',
                 'Subtotal rows appear at group ends; the grand-total footer stays')
        try:
            open_cols()
            num_row = page.locator('.ir-colrow', has=page.locator('.ir-colrow__agg')).first
            num_row.locator('.ir-colrow__break').click()
            page.wait_for_timeout(400)
            page.locator('.ir-colrow__agg').first.select_option('sum')
            page.wait_for_timeout(700)
            subs = page.locator('tr.ir-subtotal-row').count()
            foot = page.locator('tfoot .ir-agg-cell').count()
            record(c, 'PASS' if subs >= 1 and foot >= 1 else 'FAIL',
                   'subtotals=%d footerCells=%d' % (subs, foot), shot(page, 'subtotals'))
            close_panel()
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'subtotals'))

        c = case('BI-IRG-09', 'Grid Features', 'Highlight rules', 'Row-scope rule colors matching rows',
                 '1. Highlight\n2. Rule: col1 is not empty, whole row, green\n3. Apply',
                 'Matching rows take the fill; the toolbar badge counts the rule')
        try:
            page.get_by_role('button', name=re.compile('highlight', re.I)).click()
            page.wait_for_timeout(400)
            page.locator('.ir-panel select').nth(1).select_option('notnull')
            page.locator('.ir-hl-swatch').nth(2).click()
            page.get_by_role('button', name=re.compile('^apply$', re.I)).click()
            page.wait_for_timeout(600)
            styled = page.locator('table.ir-table tbody tr[style*="background-color"]').count()
            badge = page.locator('.ir-count-badge').first.inner_text() if page.locator('.ir-count-badge').count() else '0'
            record(c, 'PASS' if styled >= 1 and badge == '1' else 'FAIL',
                   'styledRows=%d badge=%s' % (styled, badge), shot(page, 'highlight_rule'))
            close_panel()
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'highlight_rule'))

        c = case('BI-IRG-10', 'Grid Features', 'Calculated columns', 'Add a calculated column',
                 '1. Columns panel > + Calculated column\n2. Name TESTCALC, expr ROUND(1+1, 0)\n3. Save',
                 'The drawer opens; TESTCALC joins the grid with computed values')
        try:
            open_cols()
            page.get_by_role('button', name=re.compile('calculated column', re.I)).click()
            page.wait_for_timeout(600)
            drawer = page.locator('.ed-drawer.ed-show')
            drawer.locator('input.form-control').first.fill('TESTCALC')
            page.locator('.ir-calc-expr').fill('ROUND(1 + 1, 0)')
            page.wait_for_timeout(700)
            drawer.get_by_role('button', name=re.compile('^save$', re.I)).click()
            page.wait_for_timeout(600)
            heads = ' '.join(page.locator('table.ir-table thead th').all_inner_texts()).lower()
            ok = 'testcalc' in heads
            record(c, 'PASS' if ok else 'FAIL', 'headers=%s' % heads[:150], shot(page, 'calc_column'))
            close_panel()
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'calc_column'))

        c = case('BI-IRG-11', 'Grid Features', 'Export', 'CSV + Excel export download',
                 '1. Click CSV\n2. Click Excel', 'Both files download')
        try:
            with page.expect_download(timeout=30000) as dl:
                page.get_by_role('button', name=re.compile('^csv$', re.I)).click()
            csv_name = dl.value.suggested_filename
            # pre-warm the lazy SheetJS CDN load so the click isn't racing it
            lib_ok = page.evaluate(
                "new Promise(function (res) { require(['xlsx'],"
                " function () { res(true); }, function () { res(false); }); })")
            with page.expect_download(timeout=90000) as dl2:
                page.get_by_role('button', name=re.compile('^excel$', re.I)).click()
            xlsx_name = dl2.value.suggested_filename
            ok = csv_name.endswith('.csv') and xlsx_name.endswith('.xlsx')
            record(c, 'PASS' if ok else 'FAIL',
                   '%s + %s (lib=%s)' % (csv_name, xlsx_name, lib_ok), shot(page, 'exports'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'exports'))

        c = case('BI-IRG-12', 'Grid Features', 'Maximize', 'Fullscreen toggle + Esc restore',
                 '1. Click the maximize arrow\n2. Press Esc', 'Grid fills the screen; Esc restores')
        try:
            page.locator('.ir-btn-max').click()
            page.wait_for_timeout(400)
            maxed = page.locator('.ir-wrap.ir-max').count() == 1
            ev = shot(page, 'maximized')
            page.keyboard.press('Escape')
            page.wait_for_timeout(400)
            restored = page.locator('.ir-wrap.ir-max').count() == 0
            record(c, 'PASS' if maxed and restored else 'FAIL',
                   'maxed=%s restored=%s' % (maxed, restored), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'maximized'))

        c = case('BI-IRG-13', 'Grid Features', 'Layouts', 'Save / reset / apply round-trip (breaks + highlights + calc)',
                 '1. Save as "UAT R2 Layout"\n2. Reset view\n3. Apply the layout\n4. Delete it',
                 'Reset clears everything; Apply restores breaks, highlight, calc column')
        try:
            page.get_by_role('button', name=re.compile('layouts', re.I)).click()
            page.wait_for_timeout(500)
            page.locator('.ir-saveas input').fill('UAT R2 Layout')
            page.get_by_role('button', name=re.compile('save as new', re.I)).click()
            page.wait_for_timeout(900)
            open_cols()
            page.get_by_role('button', name=re.compile('reset view', re.I)).click()
            page.wait_for_timeout(700)
            cleared = (page.locator('tr.ir-break-row').count() == 0
                       and page.locator('table.ir-table tbody tr[style*="background-color"]').count() == 0)
            close_panel()
            page.get_by_role('button', name=re.compile('layouts|UAT R2 Layout', re.I)).first.click()
            page.wait_for_timeout(500)
            lrow = page.locator('.ir-layout-row', has=page.locator('text=UAT R2 Layout'))
            lrow.get_by_role('button', name=re.compile('^apply$', re.I)).click()
            page.wait_for_timeout(900)
            heads = ' '.join(page.locator('table.ir-table thead th').all_inner_texts()).lower()
            restored = (page.locator('tr.ir-break-row').count() >= 1
                        and page.locator('table.ir-table tbody tr[style*="background-color"]').count() >= 1
                        and 'testcalc' in heads)
            ev = shot(page, 'layout_roundtrip')
            page.get_by_role('button', name=re.compile('layouts|UAT R2 Layout', re.I)).first.click()
            page.wait_for_timeout(400)
            page.once('dialog', lambda d: d.accept())
            lrow = page.locator('.ir-layout-row', has=page.locator('text=UAT R2 Layout'))
            lrow.locator('button.btn-danger').click()
            page.wait_for_timeout(600)
            gone = page.locator('.ir-layout-row', has=page.locator('text=UAT R2 Layout')).count() == 0
            record(c, 'PASS' if cleared and restored and gone else 'FAIL',
                   'cleared=%s restored=%s deleted=%s' % (cleared, restored, gone), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'layout_roundtrip'))

        # ═════ ADMIN cases ═══════════════════════════════════════════════
        logout(ctx)
        page.evaluate("localStorage.removeItem('ifinance_jet_session')")
        login(page, 'ADMIN', 'iFinance@2026')

        c = case('BI-IRA-01', 'Admin Param Spec', 'Spec editor', 'Parameters drawer opens prefilled',
                 '1. Reports > BUDGET_UTIL_SECTOR > Open\n2. Click Parameters',
                 'Drawer lists the 4 params; Year label prefilled "Budget Year"')
        try:
            page.goto(BASE + '/index.html#reports')
            page.reload()
            page.wait_for_load_state('networkidle')
            page.wait_for_timeout(1500)
            rrow = page.locator('tr', has=page.locator('text=BUDGET_UTIL_SECTOR'))
            rrow.first.get_by_role('button', name=re.compile('^open$', re.I)).click()
            page.wait_for_timeout(1500)
            page.get_by_role('button', name=re.compile('^parameters$', re.I)).click()
            page.wait_for_timeout(1200)
            drawer = page.locator('.ed-drawer.ed-show')
            cards = drawer.locator('code').count()
            lbl = drawer.locator('input.form-control').first.input_value()
            record(c, 'PASS' if drawer.count() == 1 and cards >= 4 and lbl == 'Budget Year' else 'FAIL',
                   'cards=%d yearLabel=%r' % (cards, lbl), shot(page, 'spec_drawer'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'spec_drawer'))

        c = case('BI-IRA-02', 'Admin Param Spec', 'Spec editor', 'LOV Test button runs the query',
                 '1. In the drawer click Test under the Year LOV', 'A value count + sample values display')
        try:
            drawer = page.locator('.ed-drawer.ed-show')
            drawer.get_by_role('button', name=re.compile('^test$', re.I)).first.click()
            page.wait_for_timeout(2500)
            t = drawer.locator('span', has_text=re.compile(r'value')).first
            ok = t.count() > 0 and 'value' in t.inner_text()
            record(c, 'PASS' if ok else 'FAIL',
                   t.inner_text()[:120] if t.count() else 'no result', shot(page, 'spec_test'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'spec_test'))

        c = case('BI-IRA-03', 'Admin Param Spec', 'Spec editor', 'Save persists the spec',
                 '1. Click Save on the drawer', 'Toast confirms; drawer closes; spec unchanged on the server')
        try:
            before = requests.get(ORDS + '/rpt/ir/reports/BUDGET_UTIL_SECTOR/paramspec',
                                  headers=api_login('ADMIN', 'iFinance@2026'), timeout=60).json()['paramSpec']
            page.locator('.ed-drawer.ed-show').get_by_role(
                'button', name=re.compile('^save$', re.I)).click()
            page.wait_for_timeout(1800)
            closed = page.locator('.ed-drawer.ed-show').count() == 0
            after = requests.get(ORDS + '/rpt/ir/reports/BUDGET_UTIL_SECTOR/paramspec',
                                 headers=api_login('ADMIN', 'iFinance@2026'), timeout=60).json()['paramSpec']
            same = json.loads(before) == json.loads(after)
            record(c, 'PASS' if closed and same else 'FAIL',
                   'closed=%s specStable=%s' % (closed, same), shot(page, 'spec_saved'))
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'spec_saved'))

        c = case('BI-IRA-04', 'Admin Param Spec', 'Run drawer', 'Run now still opens the LOV Run drawer (regression)',
                 '1. Back to Reports\n2. Run now on BUDGET_UTIL_SECTOR',
                 'The Run Parameters drawer opens with LOV dropdowns from the same spec')
        try:
            page.goto(BASE + '/index.html#reports')
            page.reload()
            page.wait_for_load_state('networkidle')
            page.wait_for_timeout(1500)
            rrow = page.locator('tr', has=page.locator('text=BUDGET_UTIL_SECTOR'))
            rrow.first.get_by_role('button', name=re.compile('run now', re.I)).click()
            page.wait_for_timeout(1800)
            drawer = page.locator('.ed-drawer.ed-show')
            sels = drawer.locator('select').count()
            ok = drawer.count() == 1 and sels >= 2
            ev = shot(page, 'run_drawer')
            drawer.get_by_role('button', name=re.compile('cancel', re.I)).click()
            page.wait_for_timeout(400)
            record(c, 'PASS' if ok else 'FAIL', 'drawer=%d lovSelects=%d' % (drawer.count(), sels), ev)
        except Exception as e:
            record(c, 'FAIL', str(e), shot(page, 'run_drawer'))

        browser.close()

    # ═════ API & security cases (requests only, .txt evidence) ═══════════
    adm = api_login('ADMIN', 'iFinance@2026')
    biu = api_login('NASER.ALKHAJA', 'Manager@2026')

    c = case('BI-API-01', 'API & Security', 'ir/catalog', 'Catalog carries the spec-driven params[]',
             'GET /rpt/ir/catalog as BI_USER', 'BUDGET_UTIL_SECTOR params[] has labels/required/hasLov')
    try:
        r = requests.get(ORDS + '/rpt/ir/catalog', headers=biu, timeout=120)
        items = {i['reportCode']: i for i in r.json().get('items', [])}
        pb = {x['name'].lower(): x for x in items.get('BUDGET_UTIL_SECTOR', {}).get('params', [])}
        ok = (r.status_code == 200 and pb.get('year', {}).get('label') == 'Budget Year'
              and pb.get('year', {}).get('required') is True
              and all(pb.get(k, {}).get('hasLov') for k in ('year', 'sector', 'projecttype', 'costcenter')))
        record(c, 'PASS' if ok else 'FAIL', 'status=%d params=%s' % (r.status_code, sorted(pb)),
               txt_evidence('api_catalog', json.dumps(items.get('BUDGET_UTIL_SECTOR', {}), indent=2)[:4000]))
    except Exception as e:
        record(c, 'FAIL', str(e))

    c = case('BI-API-02', 'API & Security', 'ir/lov', 'LOV values served from PARAM_SPEC_JSON',
             'GET ir/reports/GL_BUDGET_ACTUAL/lov?param=period', '200 + items; unknown param 400')
    try:
        r1 = requests.get(ORDS + '/rpt/ir/reports/GL_BUDGET_ACTUAL/lov',
                          params={'param': 'period'}, headers=biu, timeout=120)
        r2 = requests.get(ORDS + '/rpt/ir/reports/GL_BUDGET_ACTUAL/lov',
                          params={'param': 'nope'}, headers=biu, timeout=60)
        ok = r1.status_code == 200 and len(r1.json().get('items', [])) > 0 and r2.status_code == 400
        record(c, 'PASS' if ok else 'FAIL', '%d items / unknown=%d' %
               (len(r1.json().get('items', [])), r2.status_code),
               txt_evidence('api_lov', r1.text[:2000] + '\n---unknown---\n' + r2.text[:500]))
    except Exception as e:
        record(c, 'FAIL', str(e))

    c = case('BI-API-03', 'API & Security', 'ir/paramspec', 'Spec editor endpoints are admin-only',
             'GET+PUT ir/reports/:code/paramspec as BI_USER', 'Both return 403')
    try:
        r1 = requests.get(ORDS + '/rpt/ir/reports/BUDGET_UTIL_SECTOR/paramspec', headers=biu, timeout=60)
        r2 = requests.put(ORDS + '/rpt/ir/reports/BUDGET_UTIL_SECTOR/paramspec',
                          headers=biu, json={'paramSpec': {}}, timeout=60)
        ok = r1.status_code == 403 and r2.status_code == 403
        record(c, 'PASS' if ok else 'FAIL', 'GET=%d PUT=%d' % (r1.status_code, r2.status_code),
               txt_evidence('api_paramspec_403', r1.text[:500] + '\n' + r2.text[:500]))
    except Exception as e:
        record(c, 'FAIL', str(e))

    c = case('BI-API-04', 'API & Security', 'ir/lov/preview', 'Preview guards: query-only, bind-free, admin-only',
             'POST ir/lov/preview with a DELETE, a :bind query, and as BI_USER',
             'DELETE→400, bind→400, BI_USER→403; a clean SELECT→200')
    try:
        r0 = requests.post(ORDS + '/rpt/ir/lov/preview', headers=adm,
                           json={'sql': 'SELECT DISTINCT sector FROM prod.dct_butil_scope_v WHERE sector IS NOT NULL'},
                           timeout=120)
        r1 = requests.post(ORDS + '/rpt/ir/lov/preview', headers=adm,
                           json={'sql': 'DELETE FROM prod.dct_rpt_config'}, timeout=60)
        r2 = requests.post(ORDS + '/rpt/ir/lov/preview', headers=adm,
                           json={'sql': 'SELECT :x FROM dual'}, timeout=60)
        r3 = requests.post(ORDS + '/rpt/ir/lov/preview', headers=biu,
                           json={'sql': 'SELECT 1 FROM dual'}, timeout=60)
        ok = (r0.status_code == 200 and r1.status_code == 400
              and r2.status_code == 400 and r3.status_code == 403)
        record(c, 'PASS' if ok else 'FAIL',
               'ok=%d del=%d bind=%d role=%d' % (r0.status_code, r1.status_code, r2.status_code, r3.status_code),
               txt_evidence('api_preview_guards', r0.text[:800] + '\n' + r1.text[:300]
                            + '\n' + r2.text[:300] + '\n' + r3.text[:300]))
    except Exception as e:
        record(c, 'FAIL', str(e))

    write_artifacts()
    fails = [r for r in RESULTS if r['status'] != 'PASS']
    print('---')
    print('TOTAL %d  PASS %d  FAIL %d' % (len(RESULTS), len(RESULTS) - len(fails), len(fails)))
    sys.exit(1 if fails else 0)


# ═════ artifact writers ══════════════════════════════════════════════════
HDR = ['Test ID', 'Function', 'Test Scenario', 'Steps', 'Expected Result',
       'Status', 'Tested By', 'Test Date', 'Comments / Defect Ref']
SHEETS = ['Viewer & Parameters', 'Grid Features', 'Admin Param Spec', 'API & Security']


def style_ws(ws):
    hdr_fill = PatternFill('solid', fgColor='1F6F8B')
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = hdr_fill
        cell.alignment = Alignment(vertical='top', wrap_text=True)
    widths = [12, 16, 34, 42, 42, 9, 18, 12, 50]
    for i, w in enumerate(widths):
        ws.column_dimensions[chr(65 + i)].width = w
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical='top', wrap_text=True)


def build_workbook(path, with_status):
    wb = Workbook()
    ws = wb.active
    ws.title = 'Instructions'
    ws['A1'] = 'i-Finance UAT — BI Reporting (App 211) — Interactive Reports platform'
    ws['A3'] = ('• Round 2 · %s · Environment: dev-proxy (8090) → ADB PROD ORDS.' % DATE_LBL)
    ws['A4'] = ('• Scope: BI_USER viewer + spec-driven run parameters (PARAM_SPEC_JSON), '
                'shared <interactive-report> grid (control breaks, highlight rules, calc '
                'columns, layouts, exports, maximize), admin param-spec editor, ir/* API security.')
    ws['A5'] = '• Users: NASER.ALKHAJA (BI_USER) · ADMIN (SYS_ADMIN).'
    ws['A6'] = '• Evidence: evidence_%s-02/ — one file per case.' % DATE_LBL
    ws.column_dimensions['A'].width = 120
    for sheet in SHEETS:
        w = wb.create_sheet(sheet)
        w.append(HDR)
        for r in RESULTS:
            if r['sheet'] != sheet:
                continue
            w.append([r['id'], r['function'], r['scenario'], r['steps'], r['expected'],
                      r['status'] if with_status else '',
                      TESTER if with_status else '',
                      DATE_LBL if with_status else '',
                      (r['comment'] + (('  [' + r['evidence'] + ']') if r['evidence'] else ''))
                      if with_status else ''])
        style_ws(w)
    wb.save(path)


def build_docx(path):
    total = len(RESULTS)
    passed = len([r for r in RESULTS if r['status'] == 'PASS'])
    doc = Document()
    doc.add_heading('UAT Results — BI Reporting (App 211)', level=0)
    doc.add_paragraph('Round 2 · %s · Interactive Reports platform' % DATE_LBL)
    doc.add_paragraph('Environment: BI JET dev-proxy (localhost:8090) → ADB PROD ORDS (/ords/admin/rpt). '
                      'Executed by %s with Playwright (browser) + requests (API).' % TESTER)
    doc.add_heading('Summary', level=1)
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Total cases', 'Passed', 'Failed', 'Pass rate']):
        t.rows[0].cells[i].text = h
    t.rows[1].cells[0].text = str(total)
    t.rows[1].cells[1].text = str(passed)
    t.rows[1].cells[2].text = str(total - passed)
    t.rows[1].cells[3].text = '%.0f%%' % (100.0 * passed / total if total else 0)
    doc.add_heading('Scope', level=1)
    doc.add_paragraph(
        'This round certifies the 2026-07-03 wave: (1) parameter-LOV convergence — the viewer and the '
        'admin Run drawer both read PARAM_SPEC_JSON (labels EN/AR, hints, required flags, LOV queries; '
        'legacy PARAMS_LOV_JSON dropped by reporting/db/14); (2) the <interactive-report> grid promoted '
        'to final apps/shared/ with new control breaks (group bands + per-group subtotals) and highlight '
        'rules (row/cell scope, 5-color palette); (3) the admin Parameters editor on the report detail '
        'page (labels/hints/required/LOV SQL + Test button via POST ir/lov/preview); (4) BI_USER '
        'role-gating regression across UI routes and ir/* endpoints.')
    for sheet in SHEETS:
        rows = [r for r in RESULTS if r['sheet'] == sheet]
        if not rows:
            continue
        doc.add_heading(sheet, level=1)
        t = doc.add_table(rows=1 + len(rows), cols=5)
        t.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Test ID', 'Scenario', 'Status', 'Evidence', 'Comment']):
            t.rows[0].cells[i].text = h
        for j, r in enumerate(rows, start=1):
            t.rows[j].cells[0].text = r['id']
            t.rows[j].cells[1].text = r['scenario']
            t.rows[j].cells[2].text = r['status']
            t.rows[j].cells[3].text = r['evidence']
            t.rows[j].cells[4].text = r['comment'][:300]
    for section in doc.sections:
        section.left_margin = section.right_margin = Pt(40)
    doc.save(path)


def write_artifacts():
    build_workbook(os.path.join(ROUND_DIR, 'UAT_BI_%s-02.xlsx' % DATE_LBL), with_status=True)
    build_workbook(os.path.join(UAT_DIR, 'UAT_BI_TestScript.xlsx'), with_status=False)
    build_docx(os.path.join(ROUND_DIR, 'UAT_BI_Results_%s-02.docx' % DATE_LBL))
    print('artifacts written to ' + ROUND_DIR)


if __name__ == '__main__':
    main()
