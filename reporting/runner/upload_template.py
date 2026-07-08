"""Seed / upload report templates into DCT_RPT_TEMPLATE.

Two transports: direct DB (default; RPT_DB_*/ATD_DB_* wallet env, like the
runner) or --ords (PUT /rpt/templates/:name; needs RPT_ORDS_BASE +
RPT_ORDS_USER + RPT_ORDS_PW of a SYS_ADMIN).

Usage:
  python upload_template.py --seed              # bundled .j2 files + starter .docx
  python upload_template.py --seed --ords       # same, via the ORDS endpoints
  python upload_template.py path/to/My.docx     # upload/replace specific file(s)
  python upload_template.py --starter out.docx  # just write the starter Word file

The starter .docx is the Word twin of report.html.j2: title, meta line, params
chips and a dynamic full-width data table (docxtpl {%tr%}/{%tc%} loops over
`headers` + `rows`). Authors open it in Word, restyle freely, keep the tags.
"""
import argparse
import os
import sys

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TPL_DIR = os.path.join(os.path.dirname(__file__), "templates")
BUNDLED_J2 = {
    "report.html.j2": "Default single-table HTML report (Chromium PDF)",
    "budget_util_sector.html.j2": "Budget Utilization by Sector 6-section pack (landscape)",
}
STARTER_NAME = "report_starter.docx"
STARTER_DESC = "Starter Word template -- open in Word, restyle, keep the {{ }} tags"


def make_starter_docx():
    """Build the starter Word template as bytes (python-docx ships with docxtpl)."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    brand = RGBColor(0x1F, 0x6F, 0x8B)

    h = doc.add_heading(level=0)
    r = h.add_run("{{ report_name }}")
    r.font.color.rgb = brand

    meta = doc.add_paragraph()
    r = meta.add_run("{{ meta }}")
    r.font.size = Pt(9)

    p = doc.add_paragraph()
    r = p.add_run("{% for k, v in params.items() %}{{ k }}: {{ v }}   {% endfor %}")
    r.font.size = Pt(9)
    r.italic = True

    # docxtpl dynamic table: each {%tr%} tag needs its OWN row and each {%tc%}
    # tag its OWN cell (the tag's row/cell is consumed; the middle cell repeats
    # per loop iteration, giving one real column per header)
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "{%tc for h in headers %}"
    hdr = tbl.rows[0].cells[1].paragraphs[0].add_run("{{ h }}")
    hdr.bold = True
    tbl.rows[0].cells[2].text = "{%tc endfor %}"
    tbl.rows[1].cells[0].text = "{%tr for r in rows %}"
    tbl.rows[2].cells[0].text = "{%tc for v in r %}"
    tbl.rows[2].cells[1].text = "{{ v|fmt }}"
    tbl.rows[2].cells[2].text = "{%tc endfor %}"
    tbl.rows[3].cells[0].text = "{%tr endfor %}"

    foot = doc.add_paragraph()
    r = foot.add_run("Generated {{ generated_at }} (Asia/Dubai) - {{ report_code }}")
    r.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def upsert_ords(session_headers, name, data, mime, desc):
    import urllib.parse
    import requests
    base = os.environ["RPT_ORDS_BASE"].rstrip("/")
    q = "?mime=" + urllib.parse.quote(mime, safe="")
    if desc:
        q += "&descr=" + urllib.parse.quote(desc, safe="")
    r = requests.put(f"{base}/rpt/templates/{name}{q}", data=data,
                     headers={**session_headers,
                              "Content-Type": "application/octet-stream"},
                     timeout=120)
    r.raise_for_status()
    print(f"  {name}: {len(data):,} bytes upserted (ORDS, created={r.json()['created']})")


def ords_login():
    import requests
    base = os.environ["RPT_ORDS_BASE"].rstrip("/")
    r = requests.post(base + "/dct/auth/login",
                      json={"username": os.environ["RPT_ORDS_USER"],
                            "password": os.environ["RPT_ORDS_PW"]}, timeout=60)
    r.raise_for_status()
    return {"Authorization": "Bearer " + r.json()["sessionId"]}


def upsert(conn, name, data, mime, desc):
    cur = conn.cursor()
    cur.execute(
        """update prod.dct_rpt_template
              set content = :b, mime_type = :m, file_size = :s,
                  description = nvl(:d, description),
                  updated_by = 'SEED', updated_at = systimestamp
            where template_name = :n""",
        b=data, m=mime, s=len(data), d=desc, n=name)
    if cur.rowcount == 0:
        cur.execute(
            """insert into prod.dct_rpt_template
                 (template_name, description, mime_type, content, file_size,
                  created_by, updated_by)
               values (:n, :d, :m, :b, :s, 'SEED', 'SEED')""",
            n=name, d=desc, m=mime, b=data, s=len(data))
    conn.commit()
    print(f"  {name}: {len(data):,} bytes upserted")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("files", nargs="*", help=".docx/.j2 files to upload")
    ap.add_argument("--seed", action="store_true",
                    help="seed bundled .j2 templates + starter .docx")
    ap.add_argument("--starter", metavar="OUT",
                    help="write the starter .docx locally and exit (no DB)")
    ap.add_argument("--ords", action="store_true",
                    help="upload via PUT /rpt/templates/ instead of direct DB")
    args = ap.parse_args()

    if args.starter:
        with open(args.starter, "wb") as f:
            f.write(make_starter_docx())
        print(f"starter written: {args.starter}")
        return

    if not args.seed and not args.files:
        ap.error("nothing to do: pass --seed and/or files")

    if args.ords:
        hdrs = ords_login()
        put = lambda name, data, mime, desc: upsert_ords(hdrs, name, data, mime, desc)  # noqa: E731
    else:
        import config
        conn = config.connect()
        put = lambda name, data, mime, desc: upsert(conn, name, data, mime, desc)  # noqa: E731
    if args.seed:
        for name, desc in BUNDLED_J2.items():
            with open(os.path.join(TPL_DIR, name), "rb") as f:
                put(name, f.read(), "text/html", desc)
        put(STARTER_NAME, make_starter_docx(), DOCX_MIME, STARTER_DESC)
    for path in args.files:
        name = os.path.basename(path)
        if not (name.lower().endswith(".docx") or name.lower().endswith(".j2")):
            sys.exit(f"refusing {name}: template names must end .docx or .j2")
        mime = DOCX_MIME if name.lower().endswith(".docx") else "text/html"
        with open(path, "rb") as f:
            put(name, f.read(), mime, None)


if __name__ == "__main__":
    main()
